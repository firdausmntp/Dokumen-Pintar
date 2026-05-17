"""Tests for :mod:`dokumen_pintar.tools.toc`."""

from __future__ import annotations

from pathlib import Path
from typing import Callable

import pytest
from docx import Document
from mcp.server.fastmcp import FastMCP

from dokumen_pintar.config import AppConfig
from dokumen_pintar.context import build_context
from dokumen_pintar.errors import HandlerError, UnsupportedFormatError, ValidationError
from dokumen_pintar.tools import toc as toc_tool


def _setup(cfg: AppConfig) -> tuple[FastMCP, ...]:
    ctx = build_context(cfg)
    mcp = FastMCP(name="t-toc")
    toc_tool.register(mcp, ctx)
    return mcp, ctx


def _tool(mcp: FastMCP, name: str):
    return mcp._tool_manager._tools[name].fn


def _make_chaptered_docx(target: Path) -> None:
    doc = Document()
    doc.add_paragraph("Laporan KP", style="Title")
    doc.add_heading("BAB I PENDAHULUAN", level=1)
    doc.add_paragraph("body")
    doc.add_heading("1.1 Latar Belakang", level=2)
    doc.add_heading("1.2 Tujuan", level=2)
    doc.add_heading("BAB II", level=1)
    doc.add_heading("2.1 Section", level=2)
    doc.add_heading("2.1.1 Subsection", level=3)
    doc.save(str(target))


# ── happy path ──


def test_toc_generate_default(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    target = docs_dir / "kp.docx"
    _make_chaptered_docx(target)
    mcp, _ctx = _setup(make_config())
    result = _tool(mcp, "toc_generate")(path=str(target))
    assert result["entries"] >= 6  # title + 2 BAB + sub-headings within max_depth
    out = Document(str(target))
    paragraphs = [p.text for p in out.paragraphs]
    assert any("DAFTAR ISI" in t for t in paragraphs)
    assert any("BAB I PENDAHULUAN" in t for t in paragraphs)


def test_toc_generate_max_depth(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    target = docs_dir / "depth.docx"
    _make_chaptered_docx(target)
    mcp, _ctx = _setup(make_config())
    result = _tool(mcp, "toc_generate")(path=str(target), max_depth=1)
    # Title + 2 H1 = 3 entries.
    assert result["entries"] == 3


def test_toc_generate_indented_style(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    target = docs_dir / "indented.docx"
    _make_chaptered_docx(target)
    mcp, _ctx = _setup(make_config())
    result = _tool(mcp, "toc_generate")(path=str(target), style="indented")
    assert result["style"] == "indented"


def test_toc_generate_with_page_numbers_dotted(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    """dotted_leader + page_numbers=True emits trailing dash placeholders."""
    docs_dir, _ = tmp_roots
    target = docs_dir / "pn.docx"
    _make_chaptered_docx(target)
    mcp, _ctx = _setup(make_config())
    _tool(mcp, "toc_generate")(path=str(target), page_numbers=True)
    out = Document(str(target))
    text = "\n".join(p.text for p in out.paragraphs)
    assert "..." in text and " -" in text


def test_toc_generate_with_page_numbers_indented(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    """indented + page_numbers=True emits tab-separated placeholder."""
    docs_dir, _ = tmp_roots
    target = docs_dir / "pni.docx"
    _make_chaptered_docx(target)
    mcp, _ctx = _setup(make_config())
    _tool(mcp, "toc_generate")(path=str(target), style="indented", page_numbers=True)
    out = Document(str(target))
    # Find any line ending with \t-
    assert any(p.text.endswith("\t-") for p in out.paragraphs)


def test_toc_generate_exclude_patterns(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    target = docs_dir / "ex.docx"
    _make_chaptered_docx(target)
    mcp, _ctx = _setup(make_config())
    result = _tool(mcp, "toc_generate")(
        path=str(target),
        exclude_patterns=[r"BAB II"],
    )
    out = Document(str(target))
    text = "\n".join(p.text for p in out.paragraphs)
    # BAB II heading content was excluded from TOC body, but the original
    # heading paragraph still exists below the TOC.
    daftar_idx = next(i for i, p in enumerate(out.paragraphs) if "DAFTAR ISI" in p.text)
    body_after = out.paragraphs[daftar_idx : daftar_idx + result["entries"] + 1]
    toc_text = "\n".join(p.text for p in body_after)
    assert "BAB II" not in toc_text


def test_toc_generate_insert_at_paragraph(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    target = docs_dir / "ins_p.docx"
    _make_chaptered_docx(target)
    mcp, _ctx = _setup(make_config())
    _tool(mcp, "toc_generate")(path=str(target), insert_at="paragraph:0")
    out = Document(str(target))
    # Title is index 0; TOC marker should land at index 1 or later.
    assert "DAFTAR ISI" in [p.text for p in out.paragraphs][1]


def test_toc_generate_insert_at_after_heading(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    target = docs_dir / "ins_after.docx"
    _make_chaptered_docx(target)
    mcp, _ctx = _setup(make_config())
    _tool(mcp, "toc_generate")(path=str(target), insert_at="after:Laporan KP")
    out = Document(str(target))
    paragraphs = [p.text for p in out.paragraphs]
    title_idx = paragraphs.index("Laporan KP")
    assert "DAFTAR ISI" in paragraphs[title_idx + 1]


def test_toc_generate_insert_at_paragraph_out_of_range(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    """An out-of-range paragraph index falls through to insertion at end."""
    docs_dir, _ = tmp_roots
    target = docs_dir / "oor.docx"
    _make_chaptered_docx(target)
    mcp, _ctx = _setup(make_config())
    # Should not raise even if the index is past the end.
    _tool(mcp, "toc_generate")(path=str(target), insert_at="paragraph:9999")


# ── validation errors ──


def test_toc_generate_invalid_style(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    target = docs_dir / "x.docx"
    _make_chaptered_docx(target)
    mcp, _ctx = _setup(make_config())
    with pytest.raises(ValidationError, match="style must be"):
        _tool(mcp, "toc_generate")(path=str(target), style="weird")


def test_toc_generate_invalid_max_depth(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    target = docs_dir / "x.docx"
    _make_chaptered_docx(target)
    mcp, _ctx = _setup(make_config())
    with pytest.raises(ValidationError, match="max_depth"):
        _tool(mcp, "toc_generate")(path=str(target), max_depth=99)
    with pytest.raises(ValidationError, match="max_depth"):
        _tool(mcp, "toc_generate")(path=str(target), max_depth=-1)


def test_toc_generate_invalid_extension(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    target = docs_dir / "wrong.txt"
    target.write_text("text", encoding="utf-8")
    mcp, _ctx = _setup(make_config())
    with pytest.raises(UnsupportedFormatError, match="must be .docx"):
        _tool(mcp, "toc_generate")(path=str(target))


def test_toc_generate_file_missing(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    missing = docs_dir / "no.docx"
    mcp, _ctx = _setup(make_config())
    with pytest.raises(ValidationError, match="file not found"):
        _tool(mcp, "toc_generate")(path=str(missing))


def test_toc_generate_corrupted_docx(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    target = docs_dir / "broken.docx"
    target.write_bytes(b"PK\x03\x04not-a-docx")
    mcp, _ctx = _setup(make_config())
    with pytest.raises(HandlerError, match="failed to open"):
        _tool(mcp, "toc_generate")(path=str(target))


def test_toc_generate_no_headings(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    target = docs_dir / "no_h.docx"
    doc = Document()
    doc.add_paragraph("plain body only")
    doc.save(str(target))
    mcp, _ctx = _setup(make_config())
    with pytest.raises(ValidationError, match="no headings"):
        _tool(mcp, "toc_generate")(path=str(target))


def test_toc_generate_insert_at_invalid_paragraph_value(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    target = docs_dir / "iav.docx"
    _make_chaptered_docx(target)
    mcp, _ctx = _setup(make_config())
    with pytest.raises(ValidationError, match="must be an integer"):
        _tool(mcp, "toc_generate")(path=str(target), insert_at="paragraph:abc")


def test_toc_generate_insert_at_negative_paragraph(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    target = docs_dir / "neg.docx"
    _make_chaptered_docx(target)
    mcp, _ctx = _setup(make_config())
    with pytest.raises(ValidationError, match=r">= 0"):
        _tool(mcp, "toc_generate")(path=str(target), insert_at="paragraph:-1")


def test_toc_generate_insert_at_empty_marker(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    target = docs_dir / "em.docx"
    _make_chaptered_docx(target)
    mcp, _ctx = _setup(make_config())
    with pytest.raises(ValidationError, match="cannot be empty"):
        _tool(mcp, "toc_generate")(path=str(target), insert_at="after:")


def test_toc_generate_insert_at_marker_not_found(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    target = docs_dir / "nf.docx"
    _make_chaptered_docx(target)
    mcp, _ctx = _setup(make_config())
    with pytest.raises(ValidationError, match="marker not found"):
        _tool(mcp, "toc_generate")(path=str(target), insert_at="after:NONEXISTENT_HEADING")


def test_toc_generate_insert_at_unknown_format(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    target = docs_dir / "uk.docx"
    _make_chaptered_docx(target)
    mcp, _ctx = _setup(make_config())
    with pytest.raises(ValidationError, match="paragraph:N"):
        _tool(mcp, "toc_generate")(path=str(target), insert_at="garbage")


def test_toc_generate_save_failure_raises_handler_error(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    """Save errors during write surface as HandlerError."""
    from unittest.mock import patch

    docs_dir, _ = tmp_roots
    target = docs_dir / "fs.docx"
    _make_chaptered_docx(target)
    mcp, _ctx = _setup(make_config())
    with patch("docx.document.Document.save", side_effect=OSError("no disk")):
        with pytest.raises(HandlerError, match="failed to save"):
            _tool(mcp, "toc_generate")(path=str(target))
