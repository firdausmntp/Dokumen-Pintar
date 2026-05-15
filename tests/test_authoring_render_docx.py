"""Tests for :mod:`dokumen_pintar.authoring.render_docx`."""

from __future__ import annotations

import struct
import zlib
from pathlib import Path

import pytest
from docx import Document

from dokumen_pintar.authoring.render_docx import render_docx
from dokumen_pintar.authoring.spec import validate_spec
from dokumen_pintar.errors import HandlerError


def _png_bytes(width: int = 4, height: int = 4) -> bytes:
    """Tiny 4x4 white PNG for image-embedding tests."""

    def chunk(tag: bytes, data: bytes) -> bytes:
        return (
            struct.pack(">I", len(data))
            + tag
            + data
            + struct.pack(">I", zlib.crc32(tag + data) & 0xFFFFFFFF)
        )

    sig = b"\x89PNG\r\n\x1a\n"
    ihdr = chunk(
        b"IHDR",
        struct.pack(">IIBBBBB", width, height, 8, 2, 0, 0, 0),
    )
    raw = b""
    for _ in range(height):
        raw += b"\x00" + b"\xff\xff\xff" * width
    idat = chunk(b"IDAT", zlib.compress(raw))
    iend = chunk(b"IEND", b"")
    return sig + ihdr + idat + iend


def test_render_full_doc(tmp_path: Path) -> None:
    img_path = tmp_path / "tiny.png"
    img_path.write_bytes(_png_bytes())
    spec = validate_spec(
        {
            "meta": {"title": "T", "author": "A"},
            "blocks": [
                {"type": "heading", "level": 1, "text": "Title"},
                {
                    "type": "paragraph",
                    "runs": [
                        {"text": "Hello "},
                        {"text": "bold", "bold": True},
                        {"text": " "},
                        {"text": "italic", "italic": True},
                        {"text": " "},
                        {"text": "underlined", "underline": True},
                        {"text": " "},
                        {"text": "code", "code": True},
                        {"text": " "},
                        {"text": "big", "font_size": 16},
                        {"text": " "},
                        {"text": "red", "color": "#ff0000"},
                    ],
                },
                {"type": "list", "items": ["a", "b"]},
                {"type": "list", "ordered": True, "items": ["one", "two"]},
                {
                    "type": "table",
                    "header": ["A", "B"],
                    "rows": [["1", "2"], ["3", "4"]],
                },
                {"type": "image", "path": str(img_path), "width_cm": 3, "caption": "c"},
                {"type": "page_break"},
                {"type": "code", "language": "python", "text": "print('x')"},
                {"type": "math", "latex": "E=mc^2"},
                {"type": "hr"},
                {"type": "blockquote", "text": "wisdom"},
            ],
        }
    )
    out = tmp_path / "out.docx"
    render_docx(spec, out)
    assert out.exists()

    # Re-open and assert structure.
    doc = Document(str(out))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "Title" in text
    assert "bold" in text
    assert "wisdom" in text
    assert len(doc.tables) == 1
    assert doc.core_properties.title == "T"


def test_render_invalid_color_falls_back(tmp_path: Path) -> None:
    spec = validate_spec(
        {
            "blocks": [
                {
                    "type": "paragraph",
                    "runs": [{"text": "x", "color": "not-a-color"}],
                }
            ]
        }
    )
    out = tmp_path / "x.docx"
    # Must not raise — invalid color is silently ignored.
    render_docx(spec, out)
    assert out.exists()


def test_image_not_found_raises(tmp_path: Path) -> None:
    spec = validate_spec(
        {"blocks": [{"type": "image", "path": str(tmp_path / "nope.png")}]}
    )
    with pytest.raises(HandlerError, match="image not found"):
        render_docx(spec, tmp_path / "x.docx")


def test_path_resolver_used(tmp_path: Path) -> None:
    img_path = tmp_path / "real.png"
    img_path.write_bytes(_png_bytes())

    def resolver(uri: str) -> Path:
        if uri == "alias:/img.png":
            return img_path
        return Path(uri)

    spec = validate_spec(
        {"blocks": [{"type": "image", "path": "alias:/img.png", "width_cm": 2}]}
    )
    out = tmp_path / "with_img.docx"
    render_docx(spec, out, path_resolver=resolver)
    assert out.exists()


def test_save_failure_wraps_handlererror(tmp_path: Path, monkeypatch) -> None:
    spec = validate_spec({"blocks": [{"type": "heading", "text": "x"}]})
    target = tmp_path / "out.docx"

    from docx.document import Document as DocxDoc

    def boom(self, *args, **kwargs):  # type: ignore[no-untyped-def]
        raise RuntimeError("disk full")

    monkeypatch.setattr(DocxDoc, "save", boom)
    with pytest.raises(HandlerError, match="failed to write docx"):
        render_docx(spec, target)


def test_empty_table_skipped(tmp_path: Path) -> None:
    spec = validate_spec({"blocks": [{"type": "table", "rows": []}]})
    out = tmp_path / "empty_table.docx"
    render_docx(spec, out)
    doc = Document(str(out))
    assert len(doc.tables) == 0


def test_unsupported_block_after_validate_is_unreachable(tmp_path: Path) -> None:
    """Render-level guard: if a block somehow bypasses validation, render
    must raise rather than silently drop it."""
    from dokumen_pintar.authoring.render_docx import _render_block
    from docx import Document as Doc

    doc = Doc()
    with pytest.raises(HandlerError, match="unsupported block type"):
        _render_block(doc, {"type": "marquee"}, path_resolver=None)
