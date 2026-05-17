"""Tests for :mod:`dokumen_pintar.tools.compare`."""

from __future__ import annotations

from pathlib import Path
from typing import Callable

import pytest
from docx import Document
from mcp.server.fastmcp import FastMCP

from dokumen_pintar.config import AppConfig
from dokumen_pintar.context import build_context
from dokumen_pintar.errors import HandlerError, UnsupportedFormatError, ValidationError
from dokumen_pintar.tools import compare


def _setup(cfg: AppConfig) -> tuple[FastMCP, ...]:
    ctx = build_context(cfg)
    mcp = FastMCP(name="t-compare")
    compare.register(mcp, ctx)
    return mcp, ctx


def _tool(mcp: FastMCP, name: str):
    return mcp._tool_manager._tools[name].fn


def _make_docx(target: Path, lines: list[str]) -> None:
    doc = Document()
    for line in lines:
        doc.add_paragraph(line)
    doc.save(str(target))


# ── happy paths ──


def test_document_compare_track_changes(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    a = docs_dir / "a.docx"
    b = docs_dir / "b.docx"
    dst = docs_dir / "diff.docx"
    _make_docx(a, ["alpha", "beta", "gamma"])
    _make_docx(b, ["alpha", "BETA", "gamma", "delta"])
    mcp, _ctx = _setup(make_config())
    result = _tool(mcp, "document_compare")(src_a=str(a), src_b=str(b), dst=str(dst))
    assert result["style"] == "track_changes"
    assert result["lines_a"] >= 3
    assert result["lines_b"] >= 4
    out = Document(str(dst))
    text = "\n".join(p.text for p in out.paragraphs)
    assert "[+ delta +]" in text or "[+ BETA +]" in text


def test_document_compare_side_by_side(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    a = docs_dir / "sa.docx"
    b = docs_dir / "sb.docx"
    dst = docs_dir / "sbs.docx"
    _make_docx(a, ["foo", "bar"])
    _make_docx(b, ["foo", "BAZ", "extra"])
    mcp, _ctx = _setup(make_config())
    _tool(mcp, "document_compare")(src_a=str(a), src_b=str(b), dst=str(dst), style="side_by_side")
    out = Document(str(dst))
    assert out.tables  # the comparison creates a 2-column table
    assert out.tables[0].cell(0, 0).text == "sa.docx"
    assert out.tables[0].cell(0, 1).text == "sb.docx"


def test_document_compare_diff_doc(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    a = docs_dir / "da.docx"
    b = docs_dir / "db.docx"
    dst = docs_dir / "diff_out.docx"
    _make_docx(a, ["alpha\n", "removed line\n", "shared\n"])
    _make_docx(b, ["alpha\n", "added line\n", "shared\n"])
    mcp, _ctx = _setup(make_config())
    _tool(mcp, "document_compare")(src_a=str(a), src_b=str(b), dst=str(dst), style="diff_doc")
    out = Document(str(dst))
    text = "\n".join(p.text for p in out.paragraphs)
    # Diff output contains diff markers + the differing line content.
    assert "@@" in text or "Diff:" in text


def test_document_compare_identical_files(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    a = docs_dir / "id_a.docx"
    b = docs_dir / "id_b.docx"
    dst = docs_dir / "identical.docx"
    _make_docx(a, ["same"])
    _make_docx(b, ["same"])
    mcp, _ctx = _setup(make_config())
    _tool(mcp, "document_compare")(src_a=str(a), src_b=str(b), dst=str(dst))
    # Same content, output should still be created (even if minimal).
    assert dst.exists()


# ── error paths ──


def test_document_compare_invalid_style(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    a = docs_dir / "a.docx"
    b = docs_dir / "b.docx"
    _make_docx(a, ["x"])
    _make_docx(b, ["y"])
    dst = docs_dir / "out.docx"
    mcp, _ctx = _setup(make_config())
    with pytest.raises(ValidationError, match="style must be"):
        _tool(mcp, "document_compare")(src_a=str(a), src_b=str(b), dst=str(dst), style="weird")


def test_document_compare_dst_wrong_extension(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    a = docs_dir / "a.docx"
    b = docs_dir / "b.docx"
    _make_docx(a, ["x"])
    _make_docx(b, ["y"])
    dst = docs_dir / "wrong.txt"
    mcp, _ctx = _setup(make_config())
    with pytest.raises(UnsupportedFormatError, match="dst must be .docx"):
        _tool(mcp, "document_compare")(src_a=str(a), src_b=str(b), dst=str(dst))


def test_document_compare_refuses_overwrite(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    a = docs_dir / "ow_a.docx"
    b = docs_dir / "ow_b.docx"
    _make_docx(a, ["x"])
    _make_docx(b, ["y"])
    dst = docs_dir / "exists.docx"
    Document().save(str(dst))
    mcp, _ctx = _setup(make_config())
    with pytest.raises(ValidationError, match="overwrite"):
        _tool(mcp, "document_compare")(src_a=str(a), src_b=str(b), dst=str(dst))
    _tool(mcp, "document_compare")(src_a=str(a), src_b=str(b), dst=str(dst), overwrite=True)


def test_document_compare_unsupported_source_format(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    """Source whose extension has no registered handler raises."""
    docs_dir, _ = tmp_roots
    bad = docs_dir / "bad.unknownext"
    bad.write_text("?", encoding="utf-8")
    a = docs_dir / "ok.docx"
    _make_docx(a, ["x"])
    dst = docs_dir / "out.docx"
    mcp, _ctx = _setup(make_config())
    with pytest.raises(UnsupportedFormatError, match="no handler"):
        _tool(mcp, "document_compare")(src_a=str(bad), src_b=str(a), dst=str(dst))


def test_document_compare_xlsx_falls_back_to_extract(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    """xlsx has no read_text -> falls back to extract_for_search."""
    import openpyxl

    docs_dir, _ = tmp_roots
    a = docs_dir / "a.xlsx"
    b = docs_dir / "b.xlsx"
    wb = openpyxl.Workbook()
    wb.active["A1"] = "alpha"
    wb.save(a)
    wb2 = openpyxl.Workbook()
    wb2.active["A1"] = "beta"
    wb2.save(b)
    dst = docs_dir / "xx.docx"
    mcp, _ctx = _setup(make_config())
    _tool(mcp, "document_compare")(src_a=str(a), src_b=str(b), dst=str(dst), style="diff_doc")
    assert dst.exists()


def test_document_compare_save_failure(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    """Document.save errors during write surface as HandlerError."""
    from unittest.mock import patch

    docs_dir, _ = tmp_roots
    a = docs_dir / "fa.docx"
    b = docs_dir / "fb.docx"
    _make_docx(a, ["x"])
    _make_docx(b, ["y"])
    dst = docs_dir / "fail.docx"
    mcp, _ctx = _setup(make_config())
    with patch("docx.document.Document.save", side_effect=OSError("disk")):
        with pytest.raises(HandlerError, match="failed to write comparison"):
            _tool(mcp, "document_compare")(src_a=str(a), src_b=str(b), dst=str(dst))


def test_document_compare_track_changes_replace_block(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    """Cover the 'replace' opcode branch: line-for-line substitution."""
    docs_dir, _ = tmp_roots
    a = docs_dir / "ra.docx"
    b = docs_dir / "rb.docx"
    _make_docx(a, ["line one", "line two", "line three"])
    _make_docx(b, ["LINE ONE", "LINE TWO", "line three"])
    dst = docs_dir / "rep.docx"
    mcp, _ctx = _setup(make_config())
    _tool(mcp, "document_compare")(src_a=str(a), src_b=str(b), dst=str(dst), style="track_changes")
    out = Document(str(dst))
    text = "\n".join(p.text for p in out.paragraphs)
    # Both removed and added paragraphs render in the same opcode block.
    assert "[- line one -]" in text
    assert "[+ LINE ONE +]" in text


def test_document_compare_overwrite_takes_pre_snapshot(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    a = docs_dir / "snap_a.docx"
    b = docs_dir / "snap_b.docx"
    _make_docx(a, ["x"])
    _make_docx(b, ["y"])
    dst = docs_dir / "snap.docx"
    Document().save(str(dst))
    cfg = make_config()
    ctx = build_context(cfg)
    mcp = FastMCP(name="t-snap-compare")
    compare.register(mcp, ctx)
    _tool(mcp, "document_compare")(src_a=str(a), src_b=str(b), dst=str(dst), overwrite=True)
    versions = ctx.versions.list_versions(root_name="documents", rel_path="snap.docx")
    actions = {v["action"] for v in versions}
    assert "document_compare_pre" in actions
    assert "document_compare_post" in actions



def test_document_compare_track_changes_pure_deletion_block(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    """Cover the 'delete' opcode path: lines removed without insertion."""
    docs_dir, _ = tmp_roots
    a = docs_dir / "del_a.docx"
    b = docs_dir / "del_b.docx"
    _make_docx(a, ["keep", "remove this", "keep too"])
    _make_docx(b, ["keep", "keep too"])
    dst = docs_dir / "del.docx"
    mcp, _ctx = _setup(make_config())
    _tool(mcp, "document_compare")(
        src_a=str(a), src_b=str(b), dst=str(dst), style="track_changes"
    )
    out = Document(str(dst))
    text = "\n".join(p.text for p in out.paragraphs)
    assert "[- remove this -]" in text


def test_document_compare_track_changes_pure_insertion_block(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    """Cover the 'insert' opcode path: lines added without removal."""
    docs_dir, _ = tmp_roots
    a = docs_dir / "ins_a.docx"
    b = docs_dir / "ins_b.docx"
    _make_docx(a, ["base"])
    _make_docx(b, ["base", "added 1", "added 2"])
    dst = docs_dir / "ins.docx"
    mcp, _ctx = _setup(make_config())
    _tool(mcp, "document_compare")(
        src_a=str(a), src_b=str(b), dst=str(dst), style="track_changes"
    )
    out = Document(str(dst))
    text = "\n".join(p.text for p in out.paragraphs)
    assert "[+ added 1 +]" in text
    assert "[+ added 2 +]" in text



def test_document_compare_track_changes_mixed_opcodes(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    """Ensure all four opcode tags (equal, delete, insert, replace) flow back
    through the loop body so the loop's continuation branch is covered."""
    docs_dir, _ = tmp_roots
    a = docs_dir / "mix_a.docx"
    b = docs_dir / "mix_b.docx"
    # Crafted so the diff produces equal -> replace -> equal -> insert.
    _make_docx(a, ["common1", "to_change", "common2"])
    _make_docx(b, ["common1", "CHANGED", "common2", "extra"])
    dst = docs_dir / "mix.docx"
    mcp, _ctx = _setup(make_config())
    _tool(mcp, "document_compare")(
        src_a=str(a), src_b=str(b), dst=str(dst), style="track_changes"
    )
    text = "\n".join(p.text for p in Document(str(dst)).paragraphs)
    assert "common1" in text
    assert "CHANGED" in text
    assert "extra" in text