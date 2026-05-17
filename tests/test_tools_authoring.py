"""Tests for :mod:`dokumen_pintar.tools.authoring`."""

from __future__ import annotations

from pathlib import Path
from typing import Callable

import pytest
from mcp.server.fastmcp import FastMCP

from dokumen_pintar.config import AppConfig
from dokumen_pintar.context import build_context
from dokumen_pintar.errors import HandlerError, UnsupportedFormatError, ValidationError
from dokumen_pintar.tools import authoring


def _setup(cfg: AppConfig):  # type: ignore[no-untyped-def]
    ctx = build_context(cfg)
    mcp = FastMCP(name="test")
    authoring.register(mcp, ctx)
    return mcp, ctx


def _tool(mcp: FastMCP, name: str):  # type: ignore[no-untyped-def]
    return mcp._tool_manager._tools[name].fn


def test_validate_spec_tool_ok(make_config: Callable[..., AppConfig]) -> None:
    mcp, _ = _setup(make_config())
    res = _tool(mcp, "validate_spec")(spec={"blocks": [{"type": "heading", "text": "Hi"}]})
    assert res["valid"] is True
    assert res["normalized"]["blocks"][0]["text"] == "Hi"


def test_validate_spec_tool_bad(make_config: Callable[..., AppConfig]) -> None:
    mcp, _ = _setup(make_config())
    res = _tool(mcp, "validate_spec")(spec={"blocks": [{"type": "marquee"}]})
    assert res["valid"] is False
    assert "not supported" in res["error"]


def test_compose_docx_writes_file(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    mcp, _ = _setup(make_config())
    res = _tool(mcp, "compose_docx")(
        path="documents:/out.docx",
        spec={
            "blocks": [
                {"type": "heading", "text": "Hello"},
                {"type": "paragraph", "text": "World"},
            ]
        },
    )
    assert (docs_dir / "out.docx").exists()
    assert res["blocks"] == 2


def test_compose_docx_refuses_overwrite_without_flag(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    target = docs_dir / "exists.docx"
    target.write_bytes(b"placeholder")
    mcp, _ = _setup(make_config())
    with pytest.raises(ValidationError, match="overwrite"):
        _tool(mcp, "compose_docx")(
            path="documents:/exists.docx",
            spec={"blocks": [{"type": "heading", "text": "x"}]},
        )


def test_compose_docx_overwrite_flag_works(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    target = docs_dir / "ow.docx"
    target.write_bytes(b"placeholder")
    mcp, _ = _setup(make_config())
    _tool(mcp, "compose_docx")(
        path="documents:/ow.docx",
        spec={"blocks": [{"type": "heading", "text": "new"}]},
        overwrite=True,
    )
    # Real DOCX is much larger and starts with 'PK' (zip).
    assert target.read_bytes()[:2] == b"PK"


def test_compose_docx_extension_check(make_config: Callable[..., AppConfig]) -> None:
    mcp, _ = _setup(make_config())
    with pytest.raises(UnsupportedFormatError, match=".docx"):
        _tool(mcp, "compose_docx")(
            path="documents:/out.txt",
            spec={"blocks": [{"type": "heading", "text": "x"}]},
        )


def test_compose_docx_invalid_spec(make_config: Callable[..., AppConfig]) -> None:
    mcp, _ = _setup(make_config())
    with pytest.raises(ValidationError, match="invalid spec"):
        _tool(mcp, "compose_docx")(path="documents:/x.docx", spec={"blocks": "nope"})


def test_compose_pdf_writes_file(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    mcp, _ = _setup(make_config())
    _tool(mcp, "compose_pdf")(
        path="documents:/out.pdf",
        spec={"blocks": [{"type": "heading", "text": "Hi"}]},
    )
    out = docs_dir / "out.pdf"
    assert out.exists()
    assert out.read_bytes()[:4] == b"%PDF"


def test_compose_pdf_extension_check(make_config: Callable[..., AppConfig]) -> None:
    mcp, _ = _setup(make_config())
    with pytest.raises(UnsupportedFormatError, match=".pdf"):
        _tool(mcp, "compose_pdf")(
            path="documents:/out.docx",
            spec={"blocks": [{"type": "heading", "text": "x"}]},
        )


def test_compose_pdf_invalid_spec(make_config: Callable[..., AppConfig]) -> None:
    mcp, _ = _setup(make_config())
    with pytest.raises(ValidationError, match="invalid spec"):
        _tool(mcp, "compose_pdf")(path="documents:/x.pdf", spec={"blocks": "nope"})


def test_compose_from_markdown_to_docx(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    mcp, _ = _setup(make_config())
    res = _tool(mcp, "compose_from_markdown")(
        path="documents:/md.docx",
        markdown="# Title\n\nSome **bold** text.\n",
    )
    assert res["format"] == "docx"
    assert (docs_dir / "md.docx").exists()


def test_compose_from_markdown_to_pdf(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    mcp, _ = _setup(make_config())
    res = _tool(mcp, "compose_from_markdown")(
        path="documents:/md.pdf",
        markdown="# Title\n\nPara.\n",
    )
    assert res["format"] == "pdf"
    out = docs_dir / "md.pdf"
    assert out.exists()
    assert out.read_bytes()[:4] == b"%PDF"


def test_compose_from_markdown_explicit_format(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    mcp, _ = _setup(make_config())
    # Path has no extension matching format → must use explicit ``format``.
    with pytest.raises(UnsupportedFormatError):
        _tool(mcp, "compose_from_markdown")(
            path="documents:/md.txt",
            markdown="# X\n",
        )


def test_compose_from_markdown_empty_rejected(make_config: Callable[..., AppConfig]) -> None:
    mcp, _ = _setup(make_config())
    with pytest.raises(ValidationError, match="non-empty"):
        _tool(mcp, "compose_from_markdown")(
            path="documents:/x.docx",
            markdown="   \n",
        )


def test_compose_from_markdown_refuses_overwrite(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    (docs_dir / "ex.docx").write_bytes(b"x")
    mcp, _ = _setup(make_config())
    with pytest.raises(ValidationError, match="overwrite"):
        _tool(mcp, "compose_from_markdown")(
            path="documents:/ex.docx",
            markdown="# X\n",
        )


def test_path_resolver_invalid_falls_through(
    make_config: Callable[..., AppConfig], tmp_path: Path
) -> None:
    """The path resolver must not crash on URIs outside the workspace —
    it falls back to a literal Path which then surfaces a clean
    'image not found' from the renderer."""
    mcp, _ = _setup(make_config())
    with pytest.raises(Exception):  # HandlerError or similar
        _tool(mcp, "compose_docx")(
            path="documents:/img.docx",
            spec={"blocks": [{"type": "image", "path": "/no/such/place/x.png"}]},
        )


# ── v1.1.0 B.1: compose_docx with template ──


def test_compose_docx_with_template(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    """When `template` is provided, rendered blocks are appended to the template."""
    from docx import Document

    docs_dir, _ = tmp_roots
    template = docs_dir / "tpl.docx"
    tpl_doc = Document()
    tpl_doc.add_heading("UNIVERSITAS UNTIRTA", level=0)
    tpl_doc.add_paragraph("Cover page from template")
    tpl_doc.save(str(template))

    target = docs_dir / "out.docx"
    cfg = make_config()
    ctx = build_context(cfg)
    mcp = FastMCP(name="t-tpl")
    authoring.register(mcp, ctx)
    fn = mcp._tool_manager._tools["compose_docx"].fn
    spec = {
        "blocks": [
            {"type": "heading", "text": "BAB I", "level": 1},
            {"type": "paragraph", "runs": [{"text": "Latar Belakang."}]},
        ],
    }
    result = fn(path=str(target), spec=spec, template=str(template))
    assert result["template"] == str(template)
    out = Document(str(target))
    texts = [p.text for p in out.paragraphs]
    # Template's content preserved.
    assert any("UNTIRTA" in t for t in texts)
    # New content appended.
    assert any("BAB I" in t for t in texts)
    assert any("Latar Belakang" in t for t in texts)


def test_compose_docx_template_not_found(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    target = docs_dir / "notfound.docx"
    cfg = make_config()
    ctx = build_context(cfg)
    mcp = FastMCP(name="t-nf")
    authoring.register(mcp, ctx)
    fn = mcp._tool_manager._tools["compose_docx"].fn
    spec = {"blocks": [{"type": "paragraph", "runs": [{"text": "x"}]}]}
    with pytest.raises(ValidationError):
        fn(path=str(target), spec=spec, template=str(docs_dir / "missing.docx"))


def test_compose_docx_template_wrong_extension(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    bad_template = docs_dir / "wrong.txt"
    bad_template.write_text("not a docx", encoding="utf-8")
    target = docs_dir / "wt_out.docx"
    cfg = make_config()
    ctx = build_context(cfg)
    mcp = FastMCP(name="t-wt")
    authoring.register(mcp, ctx)
    fn = mcp._tool_manager._tools["compose_docx"].fn
    spec = {"blocks": [{"type": "paragraph", "runs": [{"text": "x"}]}]}
    with pytest.raises(UnsupportedFormatError):
        fn(path=str(target), spec=spec, template=str(bad_template))


def test_compose_docx_template_corrupted_raises_handler_error(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    """Template that python-docx can't parse surfaces as HandlerError."""
    docs_dir, _ = tmp_roots
    broken = docs_dir / "broken.docx"
    broken.write_bytes(b"PK\x03\x04not-real-docx")
    target = docs_dir / "br_out.docx"
    cfg = make_config()
    ctx = build_context(cfg)
    mcp = FastMCP(name="t-br")
    authoring.register(mcp, ctx)
    fn = mcp._tool_manager._tools["compose_docx"].fn
    spec = {"blocks": [{"type": "paragraph", "runs": [{"text": "x"}]}]}
    with pytest.raises(HandlerError, match="failed to load template"):
        fn(path=str(target), spec=spec, template=str(broken))


# ── v1.1.0 B.2: compose_to_markdown ──


def test_compose_to_markdown_basic(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    """Round-trip: compose a DOCX, convert to MD, verify content survives."""
    from docx import Document

    docs_dir, _ = tmp_roots
    src = docs_dir / "src.docx"
    doc = Document()
    doc.add_heading("Title", level=1)
    doc.add_paragraph("First paragraph.")
    doc.add_heading("Section", level=2)
    doc.add_paragraph("Second paragraph with some text.")
    doc.save(str(src))

    dst = docs_dir / "out.md"
    cfg = make_config()
    ctx = build_context(cfg)
    mcp = FastMCP(name="t-md")
    authoring.register(mcp, ctx)
    fn = mcp._tool_manager._tools["compose_to_markdown"].fn
    result = fn(src=str(src), dst=str(dst))
    md = dst.read_text(encoding="utf-8")
    # Heading 1 → ATX h1
    assert "# Title" in md
    assert "## Section" in md
    assert "First paragraph" in md
    assert result["size"] > 0


def test_compose_to_markdown_preserves_tables(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    """Mammoth → html2text path keeps table content."""
    from docx import Document

    docs_dir, _ = tmp_roots
    src = docs_dir / "tbl.docx"
    doc = Document()
    t = doc.add_table(rows=2, cols=2)
    t.cell(0, 0).text = "Header A"
    t.cell(0, 1).text = "Header B"
    t.cell(1, 0).text = "Cell A1"
    t.cell(1, 1).text = "Cell B1"
    doc.save(str(src))

    dst = docs_dir / "tbl.md"
    cfg = make_config()
    ctx = build_context(cfg)
    mcp = FastMCP(name="t-tbl")
    authoring.register(mcp, ctx)
    fn = mcp._tool_manager._tools["compose_to_markdown"].fn
    fn(src=str(src), dst=str(dst))
    md = dst.read_text(encoding="utf-8")
    assert "Header A" in md
    assert "Cell B1" in md


def test_compose_to_markdown_extracts_images(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    """When extract_images=True the images are written to dst.parent/images/."""
    import io as _io

    from docx import Document
    from docx.shared import Inches as _Inches
    from PIL import Image as _Image

    docs_dir, _ = tmp_roots
    src = docs_dir / "img.docx"
    doc = Document()
    doc.add_paragraph("Has image:")
    img_buf = _io.BytesIO()
    _Image.new("RGB", (16, 16), "blue").save(img_buf, "PNG")
    img_buf.seek(0)
    doc.add_picture(img_buf, width=_Inches(0.5))
    doc.save(str(src))

    dst = docs_dir / "img.md"
    cfg = make_config()
    ctx = build_context(cfg)
    mcp = FastMCP(name="t-img-md")
    authoring.register(mcp, ctx)
    fn = mcp._tool_manager._tools["compose_to_markdown"].fn
    fn(src=str(src), dst=str(dst), extract_images=True)

    images_dir = dst.parent / "images"
    assert images_dir.exists()
    assert any(images_dir.iterdir())  # at least one image extracted
    md = dst.read_text(encoding="utf-8")
    # Markdown references the relative images path.
    assert "images/" in md


def test_compose_to_markdown_inline_base64(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    """extract_images=False inlines images as base64 - no images dir created."""
    import io as _io

    from docx import Document
    from docx.shared import Inches as _Inches
    from PIL import Image as _Image

    docs_dir, _ = tmp_roots
    src = docs_dir / "inline.docx"
    doc = Document()
    img_buf = _io.BytesIO()
    _Image.new("RGB", (8, 8), "red").save(img_buf, "PNG")
    img_buf.seek(0)
    doc.add_picture(img_buf, width=_Inches(0.4))
    doc.save(str(src))

    dst = docs_dir / "inline.md"
    cfg = make_config()
    ctx = build_context(cfg)
    mcp = FastMCP(name="t-inline")
    authoring.register(mcp, ctx)
    fn = mcp._tool_manager._tools["compose_to_markdown"].fn
    fn(src=str(src), dst=str(dst), extract_images=False)
    assert not (dst.parent / "images").exists()
    md = dst.read_text(encoding="utf-8")
    assert "data:" in md  # base64 data URI


def test_compose_to_markdown_wrong_src_extension(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    src = docs_dir / "not.txt"
    src.write_text("hi", encoding="utf-8")
    dst = docs_dir / "out.md"
    cfg = make_config()
    ctx = build_context(cfg)
    mcp = FastMCP(name="t-wrong-src")
    authoring.register(mcp, ctx)
    fn = mcp._tool_manager._tools["compose_to_markdown"].fn
    with pytest.raises(UnsupportedFormatError, match="must be a .docx"):
        fn(src=str(src), dst=str(dst))


def test_compose_to_markdown_wrong_dst_extension(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    from docx import Document

    docs_dir, _ = tmp_roots
    src = docs_dir / "ok.docx"
    Document().save(str(src))
    Document(src).add_paragraph("x")
    dst = docs_dir / "wrong.txt"
    cfg = make_config()
    ctx = build_context(cfg)
    mcp = FastMCP(name="t-wrong-dst")
    authoring.register(mcp, ctx)
    fn = mcp._tool_manager._tools["compose_to_markdown"].fn
    with pytest.raises(UnsupportedFormatError, match="must end in .md"):
        fn(src=str(src), dst=str(dst))


def test_compose_to_markdown_refuses_overwrite(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    from docx import Document

    docs_dir, _ = tmp_roots
    src = docs_dir / "doc.docx"
    d = Document()
    d.add_paragraph("hello")
    d.save(str(src))
    dst = docs_dir / "exists.md"
    dst.write_text("# old", encoding="utf-8")
    cfg = make_config()
    ctx = build_context(cfg)
    mcp = FastMCP(name="t-ow")
    authoring.register(mcp, ctx)
    fn = mcp._tool_manager._tools["compose_to_markdown"].fn
    with pytest.raises(ValidationError, match="overwrite"):
        fn(src=str(src), dst=str(dst))
    fn(src=str(src), dst=str(dst), overwrite=True)
    assert "hello" in dst.read_text(encoding="utf-8")


def test_render_docx_to_markdown_failed_open(tmp_path: Path) -> None:
    """Mammoth raising OSError surfaces as HandlerError."""
    from dokumen_pintar.authoring.render_markdown import render_docx_to_markdown

    src = tmp_path / "missing.docx"
    dst = tmp_path / "out.md"
    with pytest.raises(HandlerError, match="failed to convert"):
        render_docx_to_markdown(src, dst)


def test_render_docx_to_markdown_empty_doc_raises(tmp_path: Path) -> None:
    """A DOCX that yields no markdown body raises HandlerError."""
    from docx import Document

    from dokumen_pintar.authoring.render_markdown import render_docx_to_markdown

    src = tmp_path / "empty.docx"
    Document().save(str(src))
    dst = tmp_path / "empty.md"
    with pytest.raises(HandlerError, match="no markdown content"):
        render_docx_to_markdown(src, dst)
