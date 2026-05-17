"""Tests for :mod:`dokumen_pintar.tools.search`."""

from __future__ import annotations

from pathlib import Path
from typing import Callable

import pytest
from mcp.server.fastmcp import FastMCP

from dokumen_pintar.config import AppConfig
from dokumen_pintar.context import build_context
from dokumen_pintar.errors import DokumenPintarError
from dokumen_pintar.tools import search


def _setup(cfg: AppConfig) -> tuple[FastMCP, ...]:
    ctx = build_context(cfg)
    mcp = FastMCP(name="test")
    search.register(mcp, ctx)
    return mcp, ctx


def _tool(mcp: FastMCP, name: str):
    return mcp._tool_manager._tools[name].fn


def test_search_filename(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    (docs_dir / "hello.txt").write_text("hi", encoding="utf-8")
    (docs_dir / "world.md").write_text("md", encoding="utf-8")
    mcp, _ = _setup(make_config())
    result = _tool(mcp, "search_filename")(glob_pattern="*.txt")
    assert result["count"] >= 1
    assert any("hello.txt" in m["uri"] for m in result["matches"])


def test_search_filename_with_root(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, ref_dir = tmp_roots
    (docs_dir / "a.txt").write_text("a", encoding="utf-8")
    (ref_dir / "b.txt").write_text("b", encoding="utf-8")
    mcp, _ = _setup(make_config())
    result = _tool(mcp, "search_filename")(glob_pattern="*.txt", root="documents")
    uris = [m["uri"] for m in result["matches"]]
    assert all("documents:" in u for u in uris)


def test_search_content(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    (docs_dir / "needle.txt").write_text("The quick brown fox jumps", encoding="utf-8")
    (docs_dir / "other.txt").write_text("nothing here", encoding="utf-8")
    mcp, _ = _setup(make_config())
    result = _tool(mcp, "search_content")(query="quick brown")
    assert result["truncated"] is False
    assert len(result["matches"]) >= 1
    assert any("needle.txt" in m["uri"] for m in result["matches"])


def test_search_content_regex(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    (docs_dir / "regex.txt").write_text("item123 item456", encoding="utf-8")
    mcp, _ = _setup(make_config())
    result = _tool(mcp, "search_content")(query=r"item\d+", regex=True)
    assert len(result["matches"]) >= 1


def test_search_content_no_match(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    (docs_dir / "miss.txt").write_text("nothing special", encoding="utf-8")
    mcp, _ = _setup(make_config())
    result = _tool(mcp, "search_content")(query="ZZZZNOTFOUND")
    assert result["matches"] == []


def test_search_in_format(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    (docs_dir / "search.txt").write_text("findme here", encoding="utf-8")
    (docs_dir / "data.json").write_text('{"key": "findme"}', encoding="utf-8")
    mcp, _ = _setup(make_config())
    result = _tool(mcp, "search_in_format")(query="findme", format="text")
    assert len(result["matches"]) >= 1
    # Should only find .txt, not .json
    for m in result["matches"]:
        assert m["format"] == "text"


def test_search_in_format_unknown(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    mcp, _ = _setup(make_config())
    with pytest.raises(DokumenPintarError, match="Unknown format"):
        _tool(mcp, "search_in_format")(query="x", format="nonexistent")


# ── Additional search coverage ──


def test_search_filename_limit(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    for i in range(10):
        (docs_dir / f"file{i}.txt").write_text(f"content {i}", encoding="utf-8")
    mcp, _ = _setup(make_config())
    result = _tool(mcp, "search_filename")(glob_pattern="*.txt", limit=3)
    assert result["count"] == 3


def test_search_content_case_sensitive(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    (docs_dir / "case.txt").write_text("FindMe here", encoding="utf-8")
    mcp, _ = _setup(make_config())
    result = _tool(mcp, "search_content")(query="findme", case_sensitive=True)
    assert len(result["matches"]) == 0
    result2 = _tool(mcp, "search_content")(query="FindMe", case_sensitive=True)
    assert len(result2["matches"]) >= 1


def test_search_content_max_files(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    for i in range(5):
        (docs_dir / f"m{i}.txt").write_text(f"needle {i}", encoding="utf-8")
    mcp, _ = _setup(make_config())
    result = _tool(mcp, "search_content")(query="needle", max_files=2)
    # Should have scanned at most 2 files
    assert result["truncated"] is False


def test_search_content_max_results_truncation(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    (docs_dir / "many.txt").write_text("hit\nhit\nhit\nhit\nhit", encoding="utf-8")
    mcp, _ = _setup(make_config())
    result = _tool(mcp, "search_content")(query="hit", max_results=2)
    assert result["truncated"] is True
    assert len(result["matches"]) == 2


def test_search_content_with_glob(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    (docs_dir / "yes.txt").write_text("target", encoding="utf-8")
    (docs_dir / "no.md").write_text("target", encoding="utf-8")
    mcp, _ = _setup(make_config())
    result = _tool(mcp, "search_content")(query="target", glob="*.txt")
    for m in result["matches"]:
        assert ".txt" in m["uri"]


def test_search_in_format_regex(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    (docs_dir / "regex.txt").write_text("item123 item456", encoding="utf-8")
    mcp, _ = _setup(make_config())
    result = _tool(mcp, "search_in_format")(
        query=r"item\d+", format="text", regex=True
    )
    assert len(result["matches"]) >= 1


def test_search_in_format_max_results(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    (docs_dir / "trunc.txt").write_text("x\nx\nx\nx\nx", encoding="utf-8")
    mcp, _ = _setup(make_config())
    result = _tool(mcp, "search_in_format")(
        query="x", format="text", max_results=2
    )
    assert result["truncated"] is True


def test_search_filename_with_root_filter(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    (docs_dir / "rootf.txt").write_text("x", encoding="utf-8")
    mcp, _ = _setup(make_config())
    result = _tool(mcp, "search_filename")(glob_pattern="*.txt", root="documents")
    assert result["count"] >= 1


def test_search_filename_wrong_root(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    (docs_dir / "rootf2.txt").write_text("x", encoding="utf-8")
    mcp, _ = _setup(make_config())
    result = _tool(mcp, "search_filename")(glob_pattern="rootf2.txt", root="nonexistent_root")
    assert result["count"] == 0


def test_search_content_with_root_filter(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    (docs_dir / "rf.txt").write_text("findmeroot", encoding="utf-8")
    mcp, _ = _setup(make_config())
    result = _tool(mcp, "search_content")(query="findmeroot", root="documents")
    assert len(result["matches"]) >= 1


def test_search_content_binary_file_skipped(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    (docs_dir / "binary.bin").write_bytes(b"\x00\x01\x02\x03")
    mcp, _ = _setup(make_config())
    result = _tool(mcp, "search_content")(query="anything")
    # binary.bin should not cause errors, just be skipped (no handler)
    for m in result["matches"]:
        assert "binary.bin" not in m["uri"]


def test_search_content_excludes_hidden(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    (docs_dir / "visible_search.txt").write_text("searchterm", encoding="utf-8")
    sub = docs_dir / ".hidden_dir"
    sub.mkdir()
    (sub / "hidden.txt").write_text("searchterm", encoding="utf-8")
    cfg = make_config()
    cfg.exclude_patterns = [".*"]
    mcp, _ = _setup(cfg)
    result = _tool(mcp, "search_content")(query="searchterm")
    for m in result["matches"]:
        assert ".hidden_dir" not in m["uri"]


def test_search_in_format_case_sensitive(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    (docs_dir / "cs.txt").write_text("HelloWorld", encoding="utf-8")
    mcp, _ = _setup(make_config())
    result = _tool(mcp, "search_in_format")(
        query="helloworld", format="text", case_sensitive=True
    )
    assert len(result["matches"]) == 0
    result2 = _tool(mcp, "search_in_format")(
        query="HelloWorld", format="text", case_sensitive=True
    )
    assert len(result2["matches"]) >= 1


def test_search_content_with_nonexistent_root(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    import shutil
    docs_dir, _ = tmp_roots
    (docs_dir / "find_me.txt").write_text("findable", encoding="utf-8")
    cfg = make_config()
    mcp, _ = _setup(cfg)
    # Search with root filter that doesn't exist
    result = _tool(mcp, "search_content")(query="findable", root="nonexistent_root")
    assert result["matches"] == []


def test_search_content_extract_error_continues(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    from unittest.mock import patch
    docs_dir, _ = tmp_roots
    (docs_dir / "err_search.txt").write_text("findme", encoding="utf-8")
    cfg = make_config()
    mcp, ctx = _setup(cfg)
    # Patch extract_for_search on the text handler to raise
    from dokumen_pintar.handlers.text_handler import TextHandler
    with patch.object(TextHandler, "extract_for_search", side_effect=DokumenPintarError("fail")):
        result = _tool(mcp, "search_content")(query="findme")
    assert result["matches"] == []


def test_search_content_generic_exception_continues(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    from unittest.mock import patch
    docs_dir, _ = tmp_roots
    (docs_dir / "exc_search.txt").write_text("findme2", encoding="utf-8")
    cfg = make_config()
    mcp, ctx = _setup(cfg)
    from dokumen_pintar.handlers.text_handler import TextHandler
    with patch.object(TextHandler, "extract_for_search", side_effect=RuntimeError("oops")):
        result = _tool(mcp, "search_content")(query="findme2")
    assert result["matches"] == []


def test_search_in_format_extract_error(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    from unittest.mock import patch
    docs_dir, _ = tmp_roots
    (docs_dir / "fmt_err.txt").write_text("x", encoding="utf-8")
    cfg = make_config()
    mcp, ctx = _setup(cfg)
    from dokumen_pintar.handlers.text_handler import TextHandler
    with patch.object(TextHandler, "extract_for_search", side_effect=DokumenPintarError("fail")):
        result = _tool(mcp, "search_in_format")(query="x", format="text")
    assert result["matches"] == []


def test_walk_files_root_not_exists(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    import shutil
    docs_dir, _ = tmp_roots
    (docs_dir / "file.txt").write_text("hello", encoding="utf-8")
    cfg = make_config()
    mcp, ctx = _setup(cfg)
    # Remove root dir so it doesn't exist
    shutil.rmtree(docs_dir)
    result = _tool(mcp, "search_content")(query="hello")
    assert result["matches"] == []


def test_walk_files_relative_to_valueerror(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    from unittest.mock import patch
    docs_dir, _ = tmp_roots
    (docs_dir / "ok.txt").write_text("test", encoding="utf-8")
    cfg = make_config()
    mcp, ctx = _setup(cfg)
    _original_rglob = Path.rglob
    def _rglob_with_outside(self, pat):
        yield from _original_rglob(self, pat)
        fake = Path("Z:/outside/fake.txt")
        yield fake
    _original_is_file = Path.is_file
    def _is_file_override(self):
        if str(self).startswith("Z:"):
            return True
        return _original_is_file(self)
    with patch.object(Path, "rglob", _rglob_with_outside):
        with patch.object(Path, "is_file", _is_file_override):
            result = _tool(mcp, "search_content")(query="test")
    assert isinstance(result["matches"], list)



# ── v1.1.0 2.4: search_content heading context ──


def test_search_content_include_context_docx(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    """``include_context=True`` enriches DOCX hits with heading_path + paragraph_index."""
    from docx import Document

    docs_dir, _ = tmp_roots
    p = docs_dir / "report.docx"
    doc = Document()
    doc.add_paragraph("BAB I", style="Heading 1")
    doc.add_paragraph("1.1 Latar Belakang", style="Heading 2")
    doc.add_paragraph("Integrasi data SAP ke sistem internal.")
    doc.add_paragraph("BAB II", style="Heading 1")
    doc.add_paragraph("Pembahasan integrasi lanjutan.")
    doc.save(str(p))

    mcp, _ctx = _setup(make_config())
    result = _tool(mcp, "search_content")(query="integrasi", include_context=True)
    matches = result["matches"]
    assert len(matches) >= 2
    # Hit inside Latar Belakang section.
    bab1_hit = next(m for m in matches if "SAP" in m["snippet"])
    assert bab1_hit["context"]["format"] == "docx"
    assert bab1_hit["context"]["heading_path"] == "BAB I > 1.1 Latar Belakang"
    assert isinstance(bab1_hit["context"]["paragraph_index"], int)
    # Hit inside BAB II.
    bab2_hit = next(m for m in matches if "lanjutan" in m["snippet"])
    assert bab2_hit["context"]["heading_path"] == "BAB II"


def test_search_content_include_context_with_title(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    """A `Title` style resets the breadcrumb stack to a single entry."""
    from docx import Document

    docs_dir, _ = tmp_roots
    p = docs_dir / "title.docx"
    doc = Document()
    doc.add_paragraph("Laporan KP", style="Title")
    doc.add_paragraph("Pendahuluan integrasi sistem.")
    doc.save(str(p))

    mcp, _ctx = _setup(make_config())
    result = _tool(mcp, "search_content")(query="integrasi", include_context=True)
    assert result["matches"][0]["context"]["heading_path"] == "Laporan KP"


def test_search_content_include_context_no_match_returns_no_context(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    """If the snippet doesn't line up with any paragraph (table cell, e.g.),
    paragraph_index falls back to None and heading_path is empty."""
    from docx import Document

    docs_dir, _ = tmp_roots
    p = docs_dir / "tbl.docx"
    doc = Document()
    doc.add_paragraph("Heading", style="Heading 1")
    t = doc.add_table(rows=1, cols=1)
    t.cell(0, 0).text = "table-only-text"
    doc.save(str(p))

    mcp, _ctx = _setup(make_config())
    result = _tool(mcp, "search_content")(
        query="table-only-text", include_context=True
    )
    if result["matches"]:
        ctx_info = result["matches"][0]["context"]
        # The DOCX heading walker only inspects body paragraphs,
        # so a hit that lives in a table cell yields paragraph_index=None.
        assert ctx_info["paragraph_index"] is None
        assert ctx_info["heading_path"] == ""


def test_search_content_include_context_off_keeps_legacy_shape(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    """`include_context=False` (default) does not add a `context` field."""
    from docx import Document

    docs_dir, _ = tmp_roots
    p = docs_dir / "noctx.docx"
    doc = Document()
    doc.add_paragraph("alpha test")
    doc.save(str(p))

    mcp, _ctx = _setup(make_config())
    result = _tool(mcp, "search_content")(query="alpha")
    assert result["matches"]
    assert "context" not in result["matches"][0]


def test_docx_heading_context_corrupted_file_returns_none(tmp_path: Path) -> None:
    from dokumen_pintar.tools.search import _docx_heading_context

    # Truncated zip header - python-docx will refuse to load.
    p = tmp_path / "broken.docx"
    p.write_bytes(b"PK\x03\x04not-a-real-docx")
    assert _docx_heading_context(p) is None


def test_docx_heading_context_non_docx_returns_none(tmp_path: Path) -> None:
    from dokumen_pintar.tools.search import _docx_heading_context

    p = tmp_path / "plain.txt"
    p.write_text("hi", encoding="utf-8")
    assert _docx_heading_context(p) is None



# ── v1.1.0 E.1: Sastrawi Indonesian stemming ──


def test_search_content_stemmed_query_matches_morphological_variants(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    """`mengatakan` query stems to `kata` and matches `berkata` body."""
    docs_dir, _ = tmp_roots
    p = docs_dir / "stem.txt"
    p.write_text("Dia berkata bahwa hujan turun deras.", encoding="utf-8")
    mcp, _ctx = _setup(make_config())
    result = _tool(mcp, "search_content")(
        query="mengatakan", language="id", stem=True
    )
    assert result["matches"], "stemmed search should find berkata"


def test_search_content_stem_requires_indonesian_language(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    """Setting `stem=True` without language='id' raises a clear error."""
    docs_dir, _ = tmp_roots
    p = docs_dir / "x.txt"
    p.write_text("hello", encoding="utf-8")
    mcp, _ctx = _setup(make_config())
    with pytest.raises(DokumenPintarError, match="language='id'"):
        _tool(mcp, "search_content")(query="hi", stem=True)


def test_search_content_no_stem_keeps_legacy_behaviour(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    """`stem=False` (default) does not stem the query."""
    docs_dir, _ = tmp_roots
    p = docs_dir / "ns.txt"
    p.write_text("Dia berkata bahwa hujan", encoding="utf-8")
    mcp, _ctx = _setup(make_config())
    result = _tool(mcp, "search_content")(query="mengatakan")
    # Without stemming, "mengatakan" doesn't appear in the document.
    assert result["matches"] == []


def test_stem_word_handles_empty(tmp_path: Path) -> None:
    from dokumen_pintar.utils.stemming_id import stem_word

    assert stem_word("") == ""
    assert stem_word("   ") == "   "


def test_stem_text_preserves_acronyms() -> None:
    """Short uppercase tokens like 'SAP' are preserved as-is (treated as acronyms)."""
    from dokumen_pintar.utils.stemming_id import stem_text

    out = stem_text("Mengintegrasi data SAP ke sistem")
    assert "SAP" in out


def test_stem_text_handles_empty() -> None:
    from dokumen_pintar.utils.stemming_id import stem_text

    assert stem_text("") == ""


def test_stem_text_lowercases_long_uppercase() -> None:
    """Tokens longer than 5 chars get lowercased + stemmed even if all-caps."""
    from dokumen_pintar.utils.stemming_id import stem_text

    out = stem_text("MENGINTEGRASI")
    assert out == out.lower()