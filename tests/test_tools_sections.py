"""Tests for :mod:`dokumen_pintar.tools.sections`."""

from __future__ import annotations

from pathlib import Path
from typing import Callable

import pytest
from docx import Document
from mcp.server.fastmcp import FastMCP

from dokumen_pintar.config import AppConfig
from dokumen_pintar.context import build_context
from dokumen_pintar.errors import HandlerError, UnsupportedFormatError, ValidationError
from dokumen_pintar.tools import sections


def _setup(cfg: AppConfig) -> tuple[FastMCP, ...]:
    ctx = build_context(cfg)
    mcp = FastMCP(name="t-sections")
    sections.register(mcp, ctx)
    return mcp, ctx


def _tool(mcp: FastMCP, name: str):
    return mcp._tool_manager._tools[name].fn


def _make_chaptered_docx(target: Path) -> None:
    """Build a multi-chapter DOCX for extraction tests."""
    doc = Document()
    doc.add_heading("BAB I", level=1)
    doc.add_paragraph("Pendahuluan paragraph 1")
    doc.add_heading("1.1 Latar Belakang", level=2)
    doc.add_paragraph("Latar belakang body")
    doc.add_heading("BAB II", level=1)
    doc.add_paragraph("Bab dua paragraph")
    doc.add_heading("BAB III", level=1)
    doc.add_paragraph("Bab tiga paragraph")
    doc.save(str(target))


# ── section_extract ──


def test_section_extract_by_heading_pattern(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    src = docs_dir / "src.docx"
    _make_chaptered_docx(src)
    dst = docs_dir / "extracted_bab2.docx"
    mcp, _ctx = _setup(make_config())
    result = _tool(mcp, "section_extract")(src=str(src), dst=str(dst), heading_pattern=r"BAB\s+II$")
    assert result["elements_copied"] >= 2
    out = Document(str(dst))
    paragraphs = [p.text for p in out.paragraphs]
    # BAB II content present.
    assert any("BAB II" in p for p in paragraphs)
    assert any("Bab dua" in p for p in paragraphs)
    # BAB III content NOT present (extraction stopped at next heading).
    assert not any("BAB III" in p for p in paragraphs)


def test_section_extract_by_paragraph_range(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    src = docs_dir / "rng.docx"
    _make_chaptered_docx(src)
    dst = docs_dir / "rng_out.docx"
    mcp, _ctx = _setup(make_config())
    # First two body elements: heading "BAB I" + paragraph after.
    result = _tool(mcp, "section_extract")(src=str(src), dst=str(dst), paragraph_range=[0, 1])
    assert result["elements_copied"] == 2
    out = Document(str(dst))
    texts = [p.text for p in out.paragraphs]
    assert "BAB I" in texts
    assert "Pendahuluan paragraph 1" in texts


def test_section_extract_neither_pattern_nor_range(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    src = docs_dir / "n.docx"
    _make_chaptered_docx(src)
    dst = docs_dir / "n_out.docx"
    mcp, _ctx = _setup(make_config())
    with pytest.raises(ValidationError, match="exactly one"):
        _tool(mcp, "section_extract")(src=str(src), dst=str(dst))


def test_section_extract_both_pattern_and_range(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    src = docs_dir / "both.docx"
    _make_chaptered_docx(src)
    dst = docs_dir / "both_out.docx"
    mcp, _ctx = _setup(make_config())
    with pytest.raises(ValidationError, match="exactly one"):
        _tool(mcp, "section_extract")(
            src=str(src),
            dst=str(dst),
            heading_pattern=r"BAB",
            paragraph_range=[0, 1],
        )


def test_section_extract_no_matching_heading(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    src = docs_dir / "nomatch.docx"
    _make_chaptered_docx(src)
    dst = docs_dir / "nomatch_out.docx"
    mcp, _ctx = _setup(make_config())
    with pytest.raises(ValidationError, match="no content matched"):
        _tool(mcp, "section_extract")(src=str(src), dst=str(dst), heading_pattern=r"BAB\s+ZZ")


def test_section_extract_invalid_paragraph_range(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    src = docs_dir / "inv.docx"
    _make_chaptered_docx(src)
    dst = docs_dir / "inv_out.docx"
    mcp, _ctx = _setup(make_config())
    # negative start
    with pytest.raises(ValidationError, match="invalid paragraph_range"):
        _tool(mcp, "section_extract")(src=str(src), dst=str(dst), paragraph_range=[-1, 5])
    # end < start
    with pytest.raises(ValidationError, match="invalid paragraph_range"):
        _tool(mcp, "section_extract")(src=str(src), dst=str(dst), paragraph_range=[5, 2])
    # wrong length
    with pytest.raises(ValidationError, match="\\[start, end\\] pair"):
        _tool(mcp, "section_extract")(src=str(src), dst=str(dst), paragraph_range=[1, 2, 3])
    # out of range start
    with pytest.raises(ValidationError, match="out of range"):
        _tool(mcp, "section_extract")(src=str(src), dst=str(dst), paragraph_range=[999, 1000])


def test_section_extract_wrong_extensions(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    bad_src = docs_dir / "src.txt"
    bad_src.write_text("not docx", encoding="utf-8")
    src = docs_dir / "ok.docx"
    _make_chaptered_docx(src)
    dst_txt = docs_dir / "out.txt"
    dst_docx = docs_dir / "out2.docx"
    mcp, _ctx = _setup(make_config())
    with pytest.raises(UnsupportedFormatError, match="source must be .docx"):
        _tool(mcp, "section_extract")(src=str(bad_src), dst=str(dst_docx), heading_pattern="x")
    with pytest.raises(UnsupportedFormatError, match="must end in .docx"):
        _tool(mcp, "section_extract")(src=str(src), dst=str(dst_txt), heading_pattern="x")


def test_section_extract_refuses_overwrite(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    src = docs_dir / "ow.docx"
    _make_chaptered_docx(src)
    dst = docs_dir / "exists.docx"
    Document().save(str(dst))
    mcp, _ctx = _setup(make_config())
    with pytest.raises(ValidationError, match="overwrite"):
        _tool(mcp, "section_extract")(src=str(src), dst=str(dst), heading_pattern="BAB I")
    _tool(mcp, "section_extract")(
        src=str(src), dst=str(dst), heading_pattern="BAB I", overwrite=True
    )


def test_section_extract_corrupted_source(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    src = docs_dir / "broken.docx"
    src.write_bytes(b"PK\x03\x04not-real")
    dst = docs_dir / "out.docx"
    mcp, _ctx = _setup(make_config())
    with pytest.raises(HandlerError):
        _tool(mcp, "section_extract")(src=str(src), dst=str(dst), heading_pattern="x")


def test_section_extract_with_title_style(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    """Title style (level 0) is treated as a heading boundary."""
    docs_dir, _ = tmp_roots
    src = docs_dir / "title.docx"
    doc = Document()
    doc.add_paragraph("Cover Title", style="Title")
    doc.add_paragraph("Body para after title")
    doc.add_heading("Section X", level=1)
    doc.add_paragraph("Body of section")
    doc.save(str(src))
    dst = docs_dir / "title_out.docx"
    mcp, _ctx = _setup(make_config())
    result = _tool(mcp, "section_extract")(
        src=str(src), dst=str(dst), heading_pattern="Cover Title"
    )
    # Title (level 0) extracts everything up to the next equal-or-higher,
    # which is nothing - extraction goes to end of document.
    assert result["elements_copied"] >= 1


def test_section_extract_overwrites_existing_takes_pre_snapshot(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    """When overwrite=True with an existing dst, a pre snapshot is taken."""
    docs_dir, _ = tmp_roots
    src = docs_dir / "snap_src.docx"
    _make_chaptered_docx(src)
    dst = docs_dir / "snap_dst.docx"
    Document().save(str(dst))
    cfg = make_config()
    ctx = build_context(cfg)
    mcp = FastMCP(name="t-snap-ex")
    sections.register(mcp, ctx)
    _tool(mcp, "section_extract")(
        src=str(src), dst=str(dst), heading_pattern="BAB I", overwrite=True
    )
    versions = ctx.versions.list_versions(root_name="documents", rel_path="snap_dst.docx")
    actions = {v["action"] for v in versions}
    assert "section_extract_pre" in actions
    assert "section_extract_post" in actions


# ── section_merge ──


def test_section_merge_two_docs(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    a = docs_dir / "a.docx"
    da = Document()
    da.add_heading("Doc A heading", level=1)
    da.add_paragraph("Doc A body")
    da.save(str(a))
    b = docs_dir / "b.docx"
    db = Document()
    db.add_heading("Doc B heading", level=1)
    db.add_paragraph("Doc B body")
    db.save(str(b))
    dst = docs_dir / "merged.docx"
    mcp, _ctx = _setup(make_config())
    result = _tool(mcp, "section_merge")(sources=[str(a), str(b)], dst=str(dst))
    out = Document(str(dst))
    texts = [p.text for p in out.paragraphs]
    assert any("Doc A" in t for t in texts)
    assert any("Doc B" in t for t in texts)
    assert result["preserve_styles"] is False


def test_section_merge_with_page_break(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    """Default page_break_between=True inserts a page break paragraph between sources."""
    docs_dir, _ = tmp_roots
    a = docs_dir / "pa.docx"
    Document().add_paragraph("a body")
    da = Document()
    da.add_paragraph("a body")
    da.save(str(a))
    b = docs_dir / "pb.docx"
    db = Document()
    db.add_paragraph("b body")
    db.save(str(b))
    dst = docs_dir / "pmerged.docx"
    mcp, _ctx = _setup(make_config())
    _tool(mcp, "section_merge")(sources=[str(a), str(b)], dst=str(dst))
    # Verify a w:br type=page exists somewhere in the body.
    out = Document(str(dst))
    from docx.oxml.ns import qn

    has_page_break = any(br.get(qn("w:type")) == "page" for br in out.element.body.iter(qn("w:br")))
    assert has_page_break


def test_section_merge_no_page_break(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    """``page_break_between=False`` skips the explicit break insertion."""
    docs_dir, _ = tmp_roots
    a = docs_dir / "qa.docx"
    da = Document()
    da.add_paragraph("a")
    da.save(str(a))
    b = docs_dir / "qb.docx"
    db = Document()
    db.add_paragraph("b")
    db.save(str(b))
    dst = docs_dir / "qmerged.docx"
    mcp, _ctx = _setup(make_config())
    result = _tool(mcp, "section_merge")(
        sources=[str(a), str(b)], dst=str(dst), page_break_between=False
    )
    assert result["page_break_between"] is False


def test_section_merge_preserve_styles_flag(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    """`preserve_styles=True` is forwarded to docxcompose."""
    docs_dir, _ = tmp_roots
    a = docs_dir / "psa.docx"
    da = Document()
    da.add_paragraph("a")
    da.save(str(a))
    b = docs_dir / "psb.docx"
    db = Document()
    db.add_paragraph("b")
    db.save(str(b))
    dst = docs_dir / "ps_merged.docx"
    mcp, _ctx = _setup(make_config())
    result = _tool(mcp, "section_merge")(
        sources=[str(a), str(b)], dst=str(dst), preserve_styles=True
    )
    assert result["preserve_styles"] is True


def test_section_merge_empty_sources(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    dst = docs_dir / "empty.docx"
    mcp, _ctx = _setup(make_config())
    with pytest.raises(ValidationError, match="must not be empty"):
        _tool(mcp, "section_merge")(sources=[], dst=str(dst))


def test_section_merge_single_source_rejected(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    a = docs_dir / "lone.docx"
    Document().save(str(a))
    dst = docs_dir / "lone_out.docx"
    mcp, _ctx = _setup(make_config())
    with pytest.raises(ValidationError, match="at least 2"):
        _tool(mcp, "section_merge")(sources=[str(a)], dst=str(dst))


def test_section_merge_dst_wrong_extension(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    a = docs_dir / "wa.docx"
    b = docs_dir / "wb.docx"
    Document().save(str(a))
    Document().save(str(b))
    dst = docs_dir / "wrong.txt"
    mcp, _ctx = _setup(make_config())
    with pytest.raises(UnsupportedFormatError, match="must end in .docx"):
        _tool(mcp, "section_merge")(sources=[str(a), str(b)], dst=str(dst))


def test_section_merge_source_wrong_extension(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    a = docs_dir / "src.txt"
    a.write_text("text", encoding="utf-8")
    b = docs_dir / "ok.docx"
    Document().save(str(b))
    dst = docs_dir / "out.docx"
    mcp, _ctx = _setup(make_config())
    with pytest.raises(UnsupportedFormatError, match="must be .docx"):
        _tool(mcp, "section_merge")(sources=[str(a), str(b)], dst=str(dst))


def test_section_merge_missing_source(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    a = docs_dir / "exists.docx"
    Document().save(str(a))
    missing = docs_dir / "missing.docx"
    dst = docs_dir / "out.docx"
    mcp, _ctx = _setup(make_config())
    with pytest.raises(ValidationError, match="source not found"):
        _tool(mcp, "section_merge")(sources=[str(a), str(missing)], dst=str(dst))


def test_section_merge_refuses_overwrite(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    a = docs_dir / "oa.docx"
    b = docs_dir / "ob.docx"
    Document().save(str(a))
    Document().save(str(b))
    dst = docs_dir / "exists_merged.docx"
    Document().save(str(dst))
    mcp, _ctx = _setup(make_config())
    with pytest.raises(ValidationError, match="overwrite"):
        _tool(mcp, "section_merge")(sources=[str(a), str(b)], dst=str(dst))
    _tool(mcp, "section_merge")(sources=[str(a), str(b)], dst=str(dst), overwrite=True)


def test_section_merge_overwrite_takes_pre_snapshot(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    a = docs_dir / "snap_a.docx"
    b = docs_dir / "snap_b.docx"
    Document().save(str(a))
    Document().save(str(b))
    dst = docs_dir / "snap_merge.docx"
    Document().save(str(dst))
    cfg = make_config()
    ctx = build_context(cfg)
    mcp = FastMCP(name="t-snap-merge")
    sections.register(mcp, ctx)
    _tool(mcp, "section_merge")(sources=[str(a), str(b)], dst=str(dst), overwrite=True)
    versions = ctx.versions.list_versions(root_name="documents", rel_path="snap_merge.docx")
    actions = {v["action"] for v in versions}
    assert "section_merge_pre" in actions
    assert "section_merge_post" in actions


def test_section_merge_corrupted_source_raises(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    a = docs_dir / "ok.docx"
    Document().save(str(a))
    bad = docs_dir / "broken.docx"
    bad.write_bytes(b"PK\x03\x04not-a-docx")
    dst = docs_dir / "merge_bad.docx"
    mcp, _ctx = _setup(make_config())
    with pytest.raises(HandlerError):
        _tool(mcp, "section_merge")(sources=[str(a), str(bad)], dst=str(dst))


def test_section_merge_composer_append_failure(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    """Force docxcompose Composer.append to raise; tool surfaces HandlerError."""
    from unittest.mock import patch

    docs_dir, _ = tmp_roots
    a = docs_dir / "ca.docx"
    b = docs_dir / "cb.docx"
    Document().save(str(a))
    Document().save(str(b))
    dst = docs_dir / "ca_merge.docx"
    mcp, _ctx = _setup(make_config())

    def _boom(self, *a, **kw):  # type: ignore[no-untyped-def]
        raise RuntimeError("composer crashed")

    from docxcompose.composer import Composer

    with patch.object(Composer, "append", _boom):
        with pytest.raises(HandlerError, match="failed appending"):
            _tool(mcp, "section_merge")(sources=[str(a), str(b)], dst=str(dst))


def test_section_merge_save_failure(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    """Force docxcompose Composer.save to raise; tool surfaces HandlerError."""
    from unittest.mock import patch

    docs_dir, _ = tmp_roots
    a = docs_dir / "sa.docx"
    b = docs_dir / "sb.docx"
    Document().save(str(a))
    Document().save(str(b))
    dst = docs_dir / "save_merge.docx"
    mcp, _ctx = _setup(make_config())

    from docxcompose.composer import Composer

    with patch.object(Composer, "save", side_effect=OSError("disk full")):
        with pytest.raises(HandlerError, match="failed saving"):
            _tool(mcp, "section_merge")(sources=[str(a), str(b)], dst=str(dst))


# ── helper coverage ──


def test_safe_load_corrupted_returns_handler_error(tmp_path: Path) -> None:
    from dokumen_pintar.tools.sections import _safe_load

    p = tmp_path / "broken.docx"
    p.write_bytes(b"random bytes not docx")
    with pytest.raises(HandlerError):
        _safe_load(p)


def test_extract_by_heading_no_match_returns_empty(tmp_path: Path) -> None:
    from dokumen_pintar.tools.sections import _extract_by_heading

    p = tmp_path / "x.docx"
    d = Document()
    d.add_paragraph("plain")
    d.save(str(p))
    src_doc = Document(str(p))
    assert _extract_by_heading(src_doc, r"NOPE") == []


def test_build_extracted_save_failure(tmp_path: Path) -> None:
    """``_build_extracted_doc`` wraps save errors as HandlerError."""
    from unittest.mock import patch

    from docx import Document as _Doc

    from dokumen_pintar.tools.sections import _build_extracted_doc

    src_path = tmp_path / "s.docx"
    d = _Doc()
    d.add_paragraph("p")
    d.save(str(src_path))
    src = _Doc(str(src_path))
    body = list(src.element.body)[0]
    # Save into a non-existent directory to force failure.
    bogus = tmp_path / "no_such_dir" / "out.docx"

    def explode(self, *a, **kw):  # type: ignore[no-untyped-def]
        raise OSError("disk full")

    with patch("docx.document.Document.save", explode):
        with pytest.raises(HandlerError, match="failed to save"):
            _build_extracted_doc(src, [body], bogus)


def test_relink_images_no_op_when_rid_missing(tmp_path: Path) -> None:
    """Re-link helper silently skips when rId is not present in src rels."""
    from unittest.mock import MagicMock

    from dokumen_pintar.tools.sections import _relink_images

    src_doc = MagicMock()
    src_doc.part.rels = {}
    dst_doc = MagicMock()

    # docxcompose xpath returns a list; provide a minimal stub element.
    fake_blip = MagicMock()
    fake_blip.get.return_value = "rId-missing"
    with __import__("unittest").mock.patch("docxcompose.utils.xpath", return_value=[fake_blip]):
        _relink_images(src_doc, dst_doc, [MagicMock()])
    # No relate_to should have been called.
    dst_doc.part.relate_to.assert_not_called()


def test_section_extract_includes_tables_inside_section(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    """Tables embedded inside a heading section are carried into the extract."""
    docs_dir, _ = tmp_roots
    src = docs_dir / "tbl_in_sec.docx"
    doc = Document()
    doc.add_heading("BAB I", level=1)
    doc.add_paragraph("Some intro paragraph in BAB I")
    t = doc.add_table(rows=1, cols=2)
    t.cell(0, 0).text = "Cell-A"
    t.cell(0, 1).text = "Cell-B"
    doc.add_paragraph("Para after table inside BAB I")
    doc.add_heading("BAB II", level=1)
    doc.add_paragraph("BAB II body")
    doc.save(str(src))

    dst = docs_dir / "tbl_extract.docx"
    mcp, _ctx = _setup(make_config())
    _tool(mcp, "section_extract")(src=str(src), dst=str(dst), heading_pattern=r"^BAB I$")
    out = Document(str(dst))
    para_texts = [p.text for p in out.paragraphs]
    assert any("intro paragraph" in t for t in para_texts)
    assert out.tables  # table survived the extract
    assert out.tables[0].cell(0, 0).text == "Cell-A"


def test_section_extract_chases_style_basedon_chain(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    """`_copy_styles_to` walks basedOn/link/next refs when copying styles."""
    from copy import deepcopy

    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    docs_dir, _ = tmp_roots
    src = docs_dir / "chain.docx"
    doc = Document()
    # Build a custom paragraph style "MyHeading" with basedOn pointing at
    # a second custom style "MyBase" - neither is in python-docx defaults,
    # so the BFS is forced to walk the basedOn chain.
    styles = doc.styles
    base_style = styles.add_style("MyBase", 1)  # 1 == WD_STYLE_TYPE.PARAGRAPH
    base_style.element.append(OxmlElement("w:rPr"))
    custom_style = styles.add_style("MyHeading", 1)
    based_on = OxmlElement("w:basedOn")
    based_on.set(qn("w:val"), "MyBase")
    custom_style.element.insert(0, based_on)
    para = doc.add_paragraph("Custom heading text")
    para.style = custom_style
    doc.save(str(src))

    dst = docs_dir / "chain_out.docx"
    mcp, _ctx = _setup(make_config())
    _tool(mcp, "section_extract")(src=str(src), dst=str(dst), paragraph_range=[0, 0])
    out = Document(str(dst))
    style_ids = {s.style_id for s in out.styles}
    assert "MyHeading" in style_ids
    assert "MyBase" in style_ids  # basedOn chain followed


def test_collect_used_style_ids_skips_blank_val() -> None:
    """``_collect_used_style_ids`` ignores style nodes whose w:val is empty."""
    from unittest.mock import MagicMock

    from dokumen_pintar.tools.sections import _collect_used_style_ids

    el = MagicMock()
    blank_node = MagicMock()
    blank_node.get.return_value = ""  # falsy w:val -> skipped
    populated_node = MagicMock()
    populated_node.get.return_value = "Heading1"

    def fake_iter(tag):
        if "pStyle" in tag:
            return [blank_node]
        if "rStyle" in tag:
            return [populated_node]
        return []

    el.iter = fake_iter
    out = _collect_used_style_ids([el])
    assert out == {"Heading1"}


def test_copy_styles_to_handles_circular_basedon() -> None:
    """`_copy_styles_to` short-circuits when a style chain has been visited."""
    from unittest.mock import MagicMock

    from docx.oxml.ns import qn

    from dokumen_pintar.tools.sections import _copy_styles_to

    style_a = MagicMock()
    based_on_a = MagicMock()
    based_on_a.get.return_value = "B"
    style_a.find.side_effect = lambda tag: based_on_a if tag == qn("w:basedOn") else None

    style_b = MagicMock()
    based_on_b = MagicMock()
    based_on_b.get.return_value = "A"
    style_b.find.side_effect = lambda tag: based_on_b if tag == qn("w:basedOn") else None

    src_doc = MagicMock()
    src_doc.styles.element.get_by_id.side_effect = lambda sid: {
        "A": style_a,
        "B": style_b,
    }.get(sid)

    # Use a custom container so iteration over .styles works without
    # python-docx's real Style objects.
    class _StylesContainer:
        def __init__(self):
            self.element = MagicMock()

        def __iter__(self):
            return iter([])

    dst_doc = MagicMock()
    dst_doc.styles = _StylesContainer()

    _copy_styles_to(src_doc, dst_doc, {"A"})
    # Both styles should have been appended once via deepcopy on the dst element.
    # No infinite loop. Two appends total expected (one per style).
    assert dst_doc.styles.element.append.call_count == 2
