"""Tests for :mod:`dokumen_pintar.tools.templates`."""

from __future__ import annotations

import io
import json
from pathlib import Path
from typing import Callable

import pytest
from docx import Document
from mcp.server.fastmcp import FastMCP
from PIL import Image as PilImage

from dokumen_pintar.config import AppConfig
from dokumen_pintar.context import build_context
from dokumen_pintar.errors import HandlerError, UnsupportedFormatError, ValidationError
from dokumen_pintar.tools import templates


def _setup(cfg: AppConfig) -> tuple[FastMCP, ...]:
    ctx = build_context(cfg)
    mcp = FastMCP(name="t-templates")
    templates.register(mcp, ctx)
    return mcp, ctx


def _tool(mcp: FastMCP, name: str):
    return mcp._tool_manager._tools[name].fn


def _make_template(target: Path, body_lines: list[str]) -> None:
    """Build a tiny Jinja-style template file for tests."""
    doc = Document()
    for line in body_lines:
        doc.add_paragraph(line)
    doc.save(str(target))


# ── template_render ──


def test_template_render_simple_vars(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    template = docs_dir / "tpl.docx"
    _make_template(template, ["Halo, {{ nama }}!", "NIM: {{ nim }}"])
    dst = docs_dir / "out.docx"
    mcp, _ctx = _setup(make_config())
    _tool(mcp, "template_render")(
        template=str(template),
        dst=str(dst),
        vars={"nama": "Firdaus", "nim": "3337230039"},
    )
    out = Document(str(dst))
    text = "\n".join(p.text for p in out.paragraphs)
    assert "Halo, Firdaus!" in text
    assert "NIM: 3337230039" in text


def test_template_render_with_loop(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    template = docs_dir / "loop.docx"
    _make_template(
        template,
        [
            "Daftar kegiatan:",
            "{% for item in entries %}- {{ item.tanggal }}: {{ item.kegiatan }}{% endfor %}",
        ],
    )
    dst = docs_dir / "loop_out.docx"
    mcp, _ctx = _setup(make_config())
    _tool(mcp, "template_render")(
        template=str(template),
        dst=str(dst),
        loops={
            "entries": [
                {"tanggal": "1 Mei", "kegiatan": "Setup"},
                {"tanggal": "2 Mei", "kegiatan": "API"},
            ]
        },
    )
    out = Document(str(dst))
    text = "\n".join(p.text for p in out.paragraphs)
    assert "1 Mei: Setup" in text
    assert "2 Mei: API" in text


def test_template_render_with_conditional(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    template = docs_dir / "cond.docx"
    _make_template(
        template,
        ["{% if show_lampiran %}Lampiran A{% else %}No lampiran{% endif %}"],
    )
    dst = docs_dir / "cond_out.docx"
    mcp, _ctx = _setup(make_config())
    _tool(mcp, "template_render")(
        template=str(template),
        dst=str(dst),
        conditionals={"show_lampiran": True},
    )
    text = "\n".join(p.text for p in Document(str(dst)).paragraphs)
    assert "Lampiran A" in text


def test_template_render_invalid_template_extension(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    bad = docs_dir / "bad.txt"
    bad.write_text("nope", encoding="utf-8")
    dst = docs_dir / "out.docx"
    mcp, _ctx = _setup(make_config())
    with pytest.raises(UnsupportedFormatError, match="template must be .docx"):
        _tool(mcp, "template_render")(template=str(bad), dst=str(dst), vars={})


def test_template_render_invalid_dst_extension(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    template = docs_dir / "x.docx"
    _make_template(template, ["x"])
    dst = docs_dir / "wrong.txt"
    mcp, _ctx = _setup(make_config())
    with pytest.raises(UnsupportedFormatError, match="dst must be .docx"):
        _tool(mcp, "template_render")(template=str(template), dst=str(dst), vars={})


def test_template_render_template_not_found(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    missing = docs_dir / "missing.docx"
    dst = docs_dir / "o.docx"
    mcp, _ctx = _setup(make_config())
    with pytest.raises(ValidationError, match="template not found"):
        _tool(mcp, "template_render")(template=str(missing), dst=str(dst), vars={})


def test_template_render_refuses_overwrite(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    template = docs_dir / "t.docx"
    _make_template(template, ["Hello"])
    dst = docs_dir / "exists.docx"
    Document().save(str(dst))
    mcp, _ctx = _setup(make_config())
    with pytest.raises(ValidationError, match="overwrite"):
        _tool(mcp, "template_render")(template=str(template), dst=str(dst), vars={})
    _tool(mcp, "template_render")(template=str(template), dst=str(dst), vars={}, overwrite=True)


def test_template_render_corrupted_template(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    bad = docs_dir / "broken.docx"
    bad.write_bytes(b"PK\x03\x04not-a-docx")
    dst = docs_dir / "o.docx"
    mcp, _ctx = _setup(make_config())
    # docxtpl loads lazily during render(), so the error surfaces from
    # the render path, not load. Match either prefix.
    with pytest.raises(HandlerError, match=r"failed to (load|render) template"):
        _tool(mcp, "template_render")(template=str(bad), dst=str(dst), vars={})


def test_template_render_invalid_jinja_syntax(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    template = docs_dir / "bad_jinja.docx"
    _make_template(template, ["{% for x in xs %}no endfor"])  # syntax error
    dst = docs_dir / "o.docx"
    mcp, _ctx = _setup(make_config())
    with pytest.raises(HandlerError, match="failed to render"):
        _tool(mcp, "template_render")(
            template=str(template), dst=str(dst), vars={}, loops={"xs": [1]}
        )


def test_template_render_loops_must_be_list(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    template = docs_dir / "lp.docx"
    _make_template(template, ["{% for x in xs %}{{ x }}{% endfor %}"])
    dst = docs_dir / "o.docx"
    mcp, _ctx = _setup(make_config())
    with pytest.raises(ValidationError, match="must be a list"):
        _tool(mcp, "template_render")(
            template=str(template), dst=str(dst), loops={"xs": "not-a-list"}
        )


def test_template_render_inline_image(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    template = docs_dir / "img.docx"
    _make_template(template, ["Logo: {{ logo }}"])
    img_path = docs_dir / "logo.png"
    PilImage.new("RGB", (16, 16), "red").save(img_path)
    dst = docs_dir / "img_out.docx"
    mcp, _ctx = _setup(make_config())
    _tool(mcp, "template_render")(
        template=str(template),
        dst=str(dst),
        inline_images={"logo": str(img_path)},
    )
    # Verify the output is a valid DOCX with an image relationship.
    out = Document(str(dst))
    rels = list(out.part.rels.values())
    assert any("image" in r.reltype for r in rels)


def test_template_render_inline_image_with_width(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    template = docs_dir / "imgw.docx"
    _make_template(template, ["{{ logo }}"])
    img_path = docs_dir / "logo.png"
    PilImage.new("RGB", (16, 16), "blue").save(img_path)
    dst = docs_dir / "imgw_out.docx"
    mcp, _ctx = _setup(make_config())
    _tool(mcp, "template_render")(
        template=str(template),
        dst=str(dst),
        inline_images={"logo": {"path": str(img_path), "width_mm": 30}},
    )
    assert dst.exists()


def test_template_render_inline_image_dict_missing_path(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    template = docs_dir / "ip.docx"
    _make_template(template, ["{{ logo }}"])
    dst = docs_dir / "o.docx"
    mcp, _ctx = _setup(make_config())
    with pytest.raises(ValidationError, match="must contain 'path'"):
        _tool(mcp, "template_render")(
            template=str(template),
            dst=str(dst),
            inline_images={"logo": {"width_mm": 30}},
        )


def test_template_render_inline_image_invalid_type(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    template = docs_dir / "iv.docx"
    _make_template(template, ["{{ logo }}"])
    dst = docs_dir / "o.docx"
    mcp, _ctx = _setup(make_config())
    with pytest.raises(ValidationError, match="must be a path string or dict"):
        _tool(mcp, "template_render")(
            template=str(template),
            dst=str(dst),
            inline_images={"logo": 123},
        )


def test_template_render_inline_image_not_found(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    template = docs_dir / "nf.docx"
    _make_template(template, ["{{ logo }}"])
    dst = docs_dir / "o.docx"
    mcp, _ctx = _setup(make_config())
    with pytest.raises(ValidationError, match="inline image not found"):
        _tool(mcp, "template_render")(
            template=str(template),
            dst=str(dst),
            inline_images={"logo": str(docs_dir / "missing.png")},
        )


# ── template_list / template_install / template_render_named ──


def test_template_list_returns_registry(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    """``template_list`` returns whatever the registry contains (may be empty)."""
    mcp, _ctx = _setup(make_config())
    result = _tool(mcp, "template_list")()
    assert "count" in result
    assert "templates" in result
    assert isinstance(result["templates"], list)


def test_template_list_filter_by_category(
    make_config: Callable[..., AppConfig],
    tmp_roots: tuple[Path, Path],
    monkeypatch,
    tmp_path: Path,
) -> None:
    """The category filter limits results to one folder."""
    # Stand up an isolated registry root with two categories.
    fake_registry = tmp_path / "templates"
    cat_a = fake_registry / "academic_id"
    cat_b = fake_registry / "business_id"
    for cat, name in [(cat_a, "kp_test"), (cat_b, "proposal")]:
        tpl_dir = cat / name
        tpl_dir.mkdir(parents=True)
        d = Document()
        d.add_paragraph(f"{name} template")
        d.save(str(tpl_dir / "template.docx"))
        (tpl_dir / "manifest.json").write_text(json.dumps({"name": name}), encoding="utf-8")
    monkeypatch.setattr(templates, "_REGISTRY_ROOT", fake_registry)
    mcp, _ctx = _setup(make_config())
    result = _tool(mcp, "template_list")(category="academic_id")
    assert result["count"] == 1
    assert result["templates"][0]["category"] == "academic_id"


def test_template_list_handles_missing_registry(
    make_config: Callable[..., AppConfig],
    tmp_roots: tuple[Path, Path],
    monkeypatch,
    tmp_path: Path,
) -> None:
    monkeypatch.setattr(templates, "_REGISTRY_ROOT", tmp_path / "nope")
    mcp, _ctx = _setup(make_config())
    result = _tool(mcp, "template_list")()
    assert result == {"count": 0, "templates": []}


def test_template_list_skips_non_directories(
    make_config: Callable[..., AppConfig],
    tmp_roots: tuple[Path, Path],
    monkeypatch,
    tmp_path: Path,
) -> None:
    """Random files at category or template level are ignored."""
    fake_registry = tmp_path / "templates"
    fake_registry.mkdir()
    (fake_registry / "stray.txt").write_text("random", encoding="utf-8")
    cat = fake_registry / "academic_id"
    cat.mkdir()
    (cat / "stray2.txt").write_text("random", encoding="utf-8")
    tpl_dir = cat / "kp_test"
    tpl_dir.mkdir()
    Document().save(str(tpl_dir / "template.docx"))
    monkeypatch.setattr(templates, "_REGISTRY_ROOT", fake_registry)
    mcp, _ctx = _setup(make_config())
    result = _tool(mcp, "template_list")()
    assert result["count"] == 1


def test_template_list_skips_template_without_docx(
    make_config: Callable[..., AppConfig],
    tmp_roots: tuple[Path, Path],
    monkeypatch,
    tmp_path: Path,
) -> None:
    fake_registry = tmp_path / "templates"
    tpl_dir = fake_registry / "academic_id" / "no_docx"
    tpl_dir.mkdir(parents=True)
    (tpl_dir / "manifest.json").write_text("{}", encoding="utf-8")
    monkeypatch.setattr(templates, "_REGISTRY_ROOT", fake_registry)
    mcp, _ctx = _setup(make_config())
    result = _tool(mcp, "template_list")()
    assert result["count"] == 0


def test_template_list_handles_invalid_manifest_json(
    make_config: Callable[..., AppConfig],
    tmp_roots: tuple[Path, Path],
    monkeypatch,
    tmp_path: Path,
) -> None:
    fake_registry = tmp_path / "templates"
    tpl_dir = fake_registry / "academic_id" / "bad_manifest"
    tpl_dir.mkdir(parents=True)
    Document().save(str(tpl_dir / "template.docx"))
    (tpl_dir / "manifest.json").write_text("{not json}", encoding="utf-8")
    monkeypatch.setattr(templates, "_REGISTRY_ROOT", fake_registry)
    mcp, _ctx = _setup(make_config())
    result = _tool(mcp, "template_list")()
    assert result["count"] == 1
    assert result["templates"][0]["manifest"] is None


def test_template_install_copies_into_workspace(
    make_config: Callable[..., AppConfig],
    tmp_roots: tuple[Path, Path],
    monkeypatch,
    tmp_path: Path,
) -> None:
    fake_registry = tmp_path / "templates"
    tpl_dir = fake_registry / "academic_id" / "kp_test"
    tpl_dir.mkdir(parents=True)
    d = Document()
    d.add_paragraph("packaged template content")
    d.save(str(tpl_dir / "template.docx"))
    monkeypatch.setattr(templates, "_REGISTRY_ROOT", fake_registry)

    docs_dir, _ = tmp_roots
    dst = docs_dir / "installed.docx"
    mcp, _ctx = _setup(make_config())
    result = _tool(mcp, "template_install")(template_id="academic_id/kp_test", dst=str(dst))
    assert dst.exists()
    text = "\n".join(p.text for p in Document(str(dst)).paragraphs)
    assert "packaged template content" in text
    assert result["template_id"] == "academic_id/kp_test"


def test_template_install_invalid_id(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    dst = docs_dir / "x.docx"
    mcp, _ctx = _setup(make_config())
    with pytest.raises(ValidationError, match="must be 'category/name'"):
        _tool(mcp, "template_install")(template_id="no_slash", dst=str(dst))


def test_template_install_unknown_id(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    dst = docs_dir / "x.docx"
    mcp, _ctx = _setup(make_config())
    with pytest.raises(ValidationError, match="not found in registry"):
        _tool(mcp, "template_install")(template_id="academic_id/does_not_exist", dst=str(dst))


def test_template_install_invalid_dst_extension(
    make_config: Callable[..., AppConfig],
    tmp_roots: tuple[Path, Path],
    monkeypatch,
    tmp_path: Path,
) -> None:
    fake_registry = tmp_path / "templates"
    tpl_dir = fake_registry / "academic_id" / "kp_test"
    tpl_dir.mkdir(parents=True)
    Document().save(str(tpl_dir / "template.docx"))
    monkeypatch.setattr(templates, "_REGISTRY_ROOT", fake_registry)

    docs_dir, _ = tmp_roots
    dst = docs_dir / "wrong.txt"
    mcp, _ctx = _setup(make_config())
    with pytest.raises(UnsupportedFormatError, match="dst must be .docx"):
        _tool(mcp, "template_install")(template_id="academic_id/kp_test", dst=str(dst))


def test_template_render_named(
    make_config: Callable[..., AppConfig],
    tmp_roots: tuple[Path, Path],
    monkeypatch,
    tmp_path: Path,
) -> None:
    fake_registry = tmp_path / "templates"
    tpl_dir = fake_registry / "academic_id" / "kp_test"
    tpl_dir.mkdir(parents=True)
    d = Document()
    d.add_paragraph("Halo {{ nama }}!")
    d.save(str(tpl_dir / "template.docx"))
    monkeypatch.setattr(templates, "_REGISTRY_ROOT", fake_registry)

    docs_dir, _ = tmp_roots
    dst = docs_dir / "named.docx"
    mcp, _ctx = _setup(make_config())
    _tool(mcp, "template_render_named")(
        template_id="academic_id/kp_test",
        dst=str(dst),
        vars={"nama": "Firdaus"},
    )
    text = "\n".join(p.text for p in Document(str(dst)).paragraphs)
    assert "Halo Firdaus!" in text


def test_template_render_named_unknown_id(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    dst = docs_dir / "x.docx"
    mcp, _ctx = _setup(make_config())
    with pytest.raises(ValidationError, match="not found in registry"):
        _tool(mcp, "template_render_named")(
            template_id="academic_id/does_not_exist",
            dst=str(dst),
            vars={},
        )


def test_template_render_named_invalid_dst(
    make_config: Callable[..., AppConfig],
    tmp_roots: tuple[Path, Path],
    monkeypatch,
    tmp_path: Path,
) -> None:
    fake_registry = tmp_path / "templates"
    tpl_dir = fake_registry / "academic_id" / "kp_test"
    tpl_dir.mkdir(parents=True)
    Document().save(str(tpl_dir / "template.docx"))
    monkeypatch.setattr(templates, "_REGISTRY_ROOT", fake_registry)
    docs_dir, _ = tmp_roots
    dst = docs_dir / "bad.txt"
    mcp, _ctx = _setup(make_config())
    with pytest.raises(UnsupportedFormatError, match="dst must be .docx"):
        _tool(mcp, "template_render_named")(
            template_id="academic_id/kp_test", dst=str(dst), vars={}
        )


def test_template_render_named_refuses_overwrite(
    make_config: Callable[..., AppConfig],
    tmp_roots: tuple[Path, Path],
    monkeypatch,
    tmp_path: Path,
) -> None:
    fake_registry = tmp_path / "templates"
    tpl_dir = fake_registry / "academic_id" / "kp_test"
    tpl_dir.mkdir(parents=True)
    Document().save(str(tpl_dir / "template.docx"))
    monkeypatch.setattr(templates, "_REGISTRY_ROOT", fake_registry)

    docs_dir, _ = tmp_roots
    dst = docs_dir / "exists.docx"
    Document().save(str(dst))
    mcp, _ctx = _setup(make_config())
    with pytest.raises(ValidationError, match="overwrite"):
        _tool(mcp, "template_render_named")(
            template_id="academic_id/kp_test", dst=str(dst), vars={}
        )
    _tool(mcp, "template_render_named")(
        template_id="academic_id/kp_test",
        dst=str(dst),
        vars={},
        overwrite=True,
    )


def test_registry_path_rejects_dotted_segments() -> None:
    from dokumen_pintar.tools.templates import _registry_path

    with pytest.raises(ValidationError):
        _registry_path("academic_id/.hidden")
    with pytest.raises(ValidationError):
        _registry_path("academic_id//empty")


def test_registry_path_rejects_three_segments() -> None:
    from dokumen_pintar.tools.templates import _registry_path

    with pytest.raises(ValidationError):
        _registry_path("a/b/c")



# ── v1.1.0 E.2: bundled kp_basic template integration ──


def test_template_list_includes_kp_basic(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    """The shipped kp_basic template appears in the live registry listing."""
    mcp, _ctx = _setup(make_config())
    result = _tool(mcp, "template_list")()
    ids = {entry["id"] for entry in result["templates"]}
    assert "academic_id/kp_basic" in ids
    kp = next(e for e in result["templates"] if e["id"] == "academic_id/kp_basic")
    assert kp["category"] == "academic_id"
    assert kp["manifest"] is not None
    assert kp["manifest"]["language"] == "id"


def test_template_render_named_kp_basic_end_to_end(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    """Rendering the bundled template through the real registry produces a usable DOCX."""
    docs_dir, _ = tmp_roots
    dst = docs_dir / "kp_rendered.docx"
    mcp, _ctx = _setup(make_config())
    _tool(mcp, "template_render_named")(
        template_id="academic_id/kp_basic",
        dst=str(dst),
        vars={
            "judul": "Sistem Monitoring SAP",
            "nama": "Firdaus Satrio Utomo",
            "nim": "3337230039",
            "jurusan": "Teknik Informatika",
            "universitas": "Universitas Sultan Ageng Tirtayasa",
            "tahun": "2026",
            "tanggal_pengesahan": "1 Juni 2026",
            "latar_belakang": "Test latar.",
            "tujuan": "Test tujuan.",
            "pelaksanaan": "Test pelaksanaan.",
            "pustaka": "Test pustaka.",
        },
    )
    out = Document(str(dst))
    text = "\n".join(p.text for p in out.paragraphs)
    assert "Sistem Monitoring SAP" in text
    assert "Firdaus Satrio Utomo" in text
    assert "3337230039" in text
    # Required headings for the academic_id_kp lint preset.
    assert "LEMBAR PENGESAHAN" in text
    assert "BAB I PENDAHULUAN" in text
    assert "LOG BOOK" in text
    assert "DAFTAR PUSTAKA" in text
