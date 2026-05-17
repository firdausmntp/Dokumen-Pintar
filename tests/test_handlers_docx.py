"""Tests for :class:`dokumen_pintar.handlers.docx_handler.DocxHandler`."""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from dokumen_pintar.errors import HandlerError
from dokumen_pintar.handlers.docx_handler import DocxHandler


@pytest.fixture
def handler() -> DocxHandler:
    return DocxHandler()


def _create_docx(path: Path, paragraphs: list[str] | None = None) -> None:
    doc = Document()
    for text in (paragraphs or ["Hello World", "Second paragraph", "Third paragraph"]):
        doc.add_paragraph(text)
    doc.save(str(path))


def test_detect(handler: DocxHandler) -> None:
    assert handler.detect(Path("doc.docx")) is True
    assert handler.detect(Path("doc.txt")) is False


def test_read_text(handler: DocxHandler, tmp_path: Path) -> None:
    target = tmp_path / "test.docx"
    _create_docx(target, ["alpha", "beta"])
    text = handler.read_text(target)
    assert "alpha" in text
    assert "beta" in text


def test_read_meta(handler: DocxHandler, tmp_path: Path) -> None:
    target = tmp_path / "test.docx"
    _create_docx(target)
    meta = handler.read_meta(target)
    assert meta["format"] == "docx"
    assert meta["paragraph_count"] >= 3
    assert "core_props" in meta
    assert "headings" in meta


def test_write_text_roundtrip(handler: DocxHandler, tmp_path: Path) -> None:
    target = tmp_path / "test.docx"
    handler.write_text(target, "line1\nline2\nline3")
    text = handler.read_text(target)
    assert "line1" in text
    assert "line2" in text
    assert "line3" in text


def test_extract_for_search(handler: DocxHandler, tmp_path: Path) -> None:
    target = tmp_path / "test.docx"
    _create_docx(target, ["searchable content here"])
    text = handler.extract_for_search(target)
    assert "searchable" in text


def test_extract_for_search_invalid_file(handler: DocxHandler, tmp_path: Path) -> None:
    target = tmp_path / "bad.docx"
    target.write_bytes(b"not a real docx")
    text = handler.extract_for_search(target)
    assert text == ""


def test_structured_get_paragraphs(handler: DocxHandler, tmp_path: Path) -> None:
    target = tmp_path / "test.docx"
    _create_docx(target, ["p0", "p1"])
    result = handler.structured_get(target, "paragraphs")
    assert isinstance(result, list)
    assert result[0]["text"] == "p0"


def test_structured_get_paragraph_by_index(handler: DocxHandler, tmp_path: Path) -> None:
    target = tmp_path / "test.docx"
    _create_docx(target, ["first", "second"])
    result = handler.structured_get(target, "paragraph:0")
    assert result["text"] == "first"


def test_structured_get_paragraph_out_of_range(handler: DocxHandler, tmp_path: Path) -> None:
    target = tmp_path / "test.docx"
    _create_docx(target, ["only one"])
    with pytest.raises(HandlerError, match="out of range"):
        handler.structured_get(target, "paragraph:999")


def test_structured_get_headings(handler: DocxHandler, tmp_path: Path) -> None:
    target = tmp_path / "test.docx"
    doc = Document()
    doc.add_heading("Title Heading", level=1)
    doc.add_paragraph("Normal text")
    doc.save(str(target))
    headings = handler.structured_get(target, "headings")
    assert isinstance(headings, list)
    assert len(headings) >= 1


def test_structured_get_core_props(handler: DocxHandler, tmp_path: Path) -> None:
    target = tmp_path / "test.docx"
    _create_docx(target)
    props = handler.structured_get(target, "core_props")
    assert isinstance(props, dict)


def test_structured_get_tables(handler: DocxHandler, tmp_path: Path) -> None:
    target = tmp_path / "test.docx"
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "A1"
    table.cell(0, 1).text = "B1"
    table.cell(1, 0).text = "A2"
    table.cell(1, 1).text = "B2"
    doc.save(str(target))
    result = handler.structured_get(target, "tables")
    assert isinstance(result, list)
    assert result[0][0] == ["A1", "B1"]


def test_structured_get_unsupported(handler: DocxHandler, tmp_path: Path) -> None:
    target = tmp_path / "test.docx"
    _create_docx(target)
    with pytest.raises(HandlerError, match="unsupported"):
        handler.structured_get(target, "invalid_expr")


def test_structured_set_paragraph(handler: DocxHandler, tmp_path: Path) -> None:
    target = tmp_path / "test.docx"
    _create_docx(target, ["original", "keep"])
    handler.structured_set(target, "paragraph:0", {"text": "modified"})
    result = handler.structured_get(target, "paragraph:0")
    assert result["text"] == "modified"


def test_structured_set_paragraph_non_dict_raises(handler: DocxHandler, tmp_path: Path) -> None:
    target = tmp_path / "test.docx"
    _create_docx(target)
    with pytest.raises(HandlerError, match="dict"):
        handler.structured_set(target, "paragraph:0", "not a dict")


def test_structured_set_paragraph_missing_text_raises(handler: DocxHandler, tmp_path: Path) -> None:
    target = tmp_path / "test.docx"
    _create_docx(target)
    with pytest.raises(HandlerError, match="text"):
        handler.structured_set(target, "paragraph:0", {"style": "Normal"})


def test_structured_set_unsupported(handler: DocxHandler, tmp_path: Path) -> None:
    target = tmp_path / "test.docx"
    _create_docx(target)
    with pytest.raises(HandlerError, match="unsupported"):
        handler.structured_set(target, "invalid_expr", "value")


def test_structured_delete_paragraph(handler: DocxHandler, tmp_path: Path) -> None:
    target = tmp_path / "test.docx"
    _create_docx(target, ["delete_me", "keep_me"])
    handler.structured_delete(target, "paragraph:0")
    text = handler.read_text(target)
    assert "delete_me" not in text
    assert "keep_me" in text


def test_structured_delete_unsupported(handler: DocxHandler, tmp_path: Path) -> None:
    target = tmp_path / "test.docx"
    _create_docx(target)
    with pytest.raises(HandlerError, match="unsupported"):
        handler.structured_delete(target, "invalid_expr")


# ── Additional docx coverage ──


def test_structured_get_table_by_index(handler: DocxHandler, tmp_path: Path) -> None:
    target = tmp_path / "tbl.docx"
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "X"
    table.cell(0, 1).text = "Y"
    doc.save(str(target))
    result = handler.structured_get(target, "table:0")
    assert result[0][0] == "X"


def test_structured_get_table_out_of_range(handler: DocxHandler, tmp_path: Path) -> None:
    target = tmp_path / "tbl2.docx"
    _create_docx(target)
    with pytest.raises(HandlerError, match="out of range"):
        handler.structured_get(target, "table:999")


def test_structured_set_core_props(handler: DocxHandler, tmp_path: Path) -> None:
    target = tmp_path / "cp.docx"
    _create_docx(target)
    handler.structured_set(target, "core_props", {"title": "Test Doc"})
    props = handler.structured_get(target, "core_props")
    assert props["title"] == "Test Doc"


def test_structured_set_core_props_non_dict(handler: DocxHandler, tmp_path: Path) -> None:
    target = tmp_path / "cp2.docx"
    _create_docx(target)
    with pytest.raises(HandlerError, match="dict"):
        handler.structured_set(target, "core_props", "not dict")


def test_structured_delete_table(handler: DocxHandler, tmp_path: Path) -> None:
    target = tmp_path / "dtbl.docx"
    doc = Document()
    doc.add_paragraph("text")
    doc.add_table(rows=2, cols=2)
    doc.save(str(target))
    handler.structured_delete(target, "table:0")
    meta = handler.read_meta(target)
    assert meta["table_count"] == 0


def test_structured_delete_table_out_of_range(handler: DocxHandler, tmp_path: Path) -> None:
    target = tmp_path / "dtbl2.docx"
    _create_docx(target)
    with pytest.raises(HandlerError, match="out of range"):
        handler.structured_delete(target, "table:999")


def test_structured_set_paragraph_out_of_range(handler: DocxHandler, tmp_path: Path) -> None:
    target = tmp_path / "por.docx"
    _create_docx(target, ["only one"])
    with pytest.raises(HandlerError, match="out of range"):
        handler.structured_set(target, "paragraph:999", {"text": "x"})


def test_structured_delete_paragraph_out_of_range(handler: DocxHandler, tmp_path: Path) -> None:
    target = tmp_path / "dpor.docx"
    _create_docx(target, ["only one"])
    with pytest.raises(HandlerError, match="out of range"):
        handler.structured_delete(target, "paragraph:999")


def test_extract_for_search_with_table(handler: DocxHandler, tmp_path: Path) -> None:
    target = tmp_path / "srtbl.docx"
    doc = Document()
    doc.add_paragraph("para text")
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "cell text"
    doc.save(str(target))
    text = handler.extract_for_search(target)
    assert "para text" in text
    assert "cell text" in text


def test_open_invalid_docx(handler: DocxHandler, tmp_path: Path) -> None:
    target = tmp_path / "bad.docx"
    target.write_bytes(b"not a docx file")
    with pytest.raises(HandlerError, match="valid docx"):
        handler.read_text(target)


def test_structured_get_invalid_file(handler: DocxHandler, tmp_path: Path) -> None:
    target = tmp_path / "bad.docx"
    target.write_bytes(b"not a valid docx")
    with pytest.raises(HandlerError):
        handler.structured_get(target, "paragraphs")


# ── More DOCX coverage ──

from dokumen_pintar.handlers.docx_handler import _heading_level, _parse_index_expr


def test_heading_level_none() -> None:
    assert _heading_level(None) is None
    assert _heading_level("") is None


def test_heading_level_title() -> None:
    assert _heading_level("Title") == 0


def test_heading_level_heading_1() -> None:
    assert _heading_level("Heading 1") == 1
    assert _heading_level("Heading 3") == 3


def test_heading_level_normal() -> None:
    assert _heading_level("Normal") is None


def test_parse_index_expr_invalid() -> None:
    with pytest.raises(HandlerError, match="invalid index"):
        _parse_index_expr("paragraph:abc", "paragraph:")


def test_parse_index_expr_empty() -> None:
    with pytest.raises(HandlerError, match="invalid index"):
        _parse_index_expr("paragraph:", "paragraph:")


def test_structured_set_paragraph_with_style(
    handler: DocxHandler, tmp_path: Path
) -> None:
    target = tmp_path / "pstyle.docx"
    doc = Document()
    doc.add_paragraph("Test text")
    doc.save(str(target))
    handler.structured_set(target, "paragraph:0", {"text": "New", "style": "Normal"})
    result = handler.structured_get(target, "paragraph:0")
    assert result["text"] == "New"


def test_structured_set_paragraph_unknown_style(
    handler: DocxHandler, tmp_path: Path
) -> None:
    target = tmp_path / "unkn.docx"
    doc = Document()
    doc.add_paragraph("Test")
    doc.save(str(target))
    with pytest.raises(HandlerError, match="unknown style"):
        handler.structured_set(target, "paragraph:0", {"text": "X", "style": "NonexistentStyle"})


def test_structured_set_paragraph_invalid_style_type(
    handler: DocxHandler, tmp_path: Path
) -> None:
    target = tmp_path / "ist.docx"
    doc = Document()
    doc.add_paragraph("Test")
    doc.save(str(target))
    with pytest.raises(HandlerError, match="style.*must be a string"):
        handler.structured_set(target, "paragraph:0", {"text": "X", "style": 42})


def test_structured_set_paragraph_missing_text(
    handler: DocxHandler, tmp_path: Path
) -> None:
    target = tmp_path / "mt.docx"
    doc = Document()
    doc.add_paragraph("Test")
    doc.save(str(target))
    with pytest.raises(HandlerError, match="missing required"):
        handler.structured_set(target, "paragraph:0", {"no_text_key": "val"})


def test_structured_set_paragraph_not_dict(
    handler: DocxHandler, tmp_path: Path
) -> None:
    target = tmp_path / "nd.docx"
    doc = Document()
    doc.add_paragraph("Test")
    doc.save(str(target))
    with pytest.raises(HandlerError, match="dict"):
        handler.structured_set(target, "paragraph:0", "not a dict")


def test_structured_set_core_props_unknown(
    handler: DocxHandler, tmp_path: Path
) -> None:
    target = tmp_path / "ucprop.docx"
    doc = Document()
    doc.add_paragraph("Test")
    doc.save(str(target))
    with pytest.raises(HandlerError, match="unknown core property"):
        handler.structured_set(target, "core_props", {"nonexistent_prop": "val"})


def test_read_text_invalid_file(handler: DocxHandler, tmp_path: Path) -> None:
    target = tmp_path / "badr.docx"
    target.write_bytes(b"not valid docx data")
    with pytest.raises(HandlerError):
        handler.read_text(target)


def test_extract_for_search_invalid(handler: DocxHandler, tmp_path: Path) -> None:
    target = tmp_path / "bads.docx"
    target.write_bytes(b"corrupt")
    text = handler.extract_for_search(target)
    assert text == ""


def test_write_text_creates_docx(handler: DocxHandler, tmp_path: Path) -> None:
    target = tmp_path / "wr.docx"
    handler.write_text(target, "Hello\nWorld")
    text = handler.read_text(target)
    assert "Hello" in text
    assert "World" in text


def test_write_text_reraises_handler_error_unwrapped(
    handler: DocxHandler, tmp_path: Path
) -> None:
    """A HandlerError raised inside the try-block must propagate as-is, not
    get wrapped a second time as 'failed to write docx'."""
    from unittest.mock import patch

    target = tmp_path / "reraise.docx"
    sentinel = HandlerError("explicit failure from inside Document()")
    with patch(
        "dokumen_pintar.handlers.docx_handler.Document", side_effect=sentinel
    ):
        with pytest.raises(HandlerError, match="explicit failure from inside") as ei:
            handler.write_text(target, "x")
        # Must be the same exception, not a wrapped one with "failed to write".
        assert ei.value is sentinel


def test_structured_delete_table(handler: DocxHandler, tmp_path: Path) -> None:
    target = tmp_path / "dt.docx"
    doc = Document()
    doc.add_paragraph("Above table")
    t = doc.add_table(rows=2, cols=2)
    t.cell(0, 0).text = "A"
    doc.add_paragraph("Below table")
    doc.save(str(target))
    
    tables_before = handler.structured_get(target, "tables")
    assert len(tables_before) == 1
    handler.structured_delete(target, "table:0")
    tables_after = handler.structured_get(target, "tables")
    assert len(tables_after) == 0


def test_open_generic_exception(handler: DocxHandler, tmp_path: Path) -> None:
    target = tmp_path / "gen.docx"
    target.write_bytes(b"\x00" * 100)
    with pytest.raises(HandlerError, match="(valid docx|failed to open)"):
        handler.read_text(target)


def test_structured_set_core_props_value_error(
    handler: DocxHandler, tmp_path: Path
) -> None:
    target = tmp_path / "cpval.docx"
    doc = Document()
    doc.add_paragraph("Test")
    doc.save(str(target))
    with pytest.raises(HandlerError, match="failed to set core property"):
        handler.structured_set(target, "core_props", {"created": "not-a-date"})


def test_structured_set_unsupported_expr(
    handler: DocxHandler, tmp_path: Path
) -> None:
    target = tmp_path / "unse.docx"
    doc = Document()
    doc.add_paragraph("Test")
    doc.save(str(target))
    with pytest.raises(HandlerError, match="unsupported"):
        handler.structured_set(target, "bad_expr", "val")


def test_structured_delete_unsupported_expr(
    handler: DocxHandler, tmp_path: Path
) -> None:
    target = tmp_path / "delu.docx"
    doc = Document()
    doc.add_paragraph("Test")
    doc.save(str(target))
    with pytest.raises(HandlerError, match="unsupported"):
        handler.structured_delete(target, "bad_expr")


# ── Additional DOCX coverage tests ──

def test_open_generic_exception(handler: DocxHandler, tmp_path: Path) -> None:
    from unittest.mock import patch
    target = tmp_path / "badopen.docx"
    doc = Document()
    doc.add_paragraph("x")
    doc.save(str(target))
    with patch("dokumen_pintar.handlers.docx_handler.Document", side_effect=RuntimeError("corrupt")):
        with pytest.raises(HandlerError, match="failed to open docx"):
            handler.read_text(target)


def test_parse_index_valueerror(handler: DocxHandler, tmp_path: Path) -> None:
    from dokumen_pintar.handlers.docx_handler import _parse_index_expr
    with pytest.raises(HandlerError, match="invalid index"):
        _parse_index_expr("paragraph:", "paragraph:")


def test_heading_level_title() -> None:
    from dokumen_pintar.handlers.docx_handler import _heading_level
    assert _heading_level("Title") == 0
    assert _heading_level("Heading 2") == 2
    assert _heading_level(None) is None
    assert _heading_level("Normal") is None


def test_write_text_save_exception(handler: DocxHandler, tmp_path: Path) -> None:
    from unittest.mock import patch
    target = tmp_path / "saveerr.docx"
    with patch("dokumen_pintar.handlers.docx_handler.Document") as MockDoc:
        mock_doc = MockDoc.return_value
        mock_doc.save.side_effect = RuntimeError("save fail")
        with pytest.raises(HandlerError, match="failed to write docx"):
            handler.write_text(target, "hello")


def test_structured_set_save_exception(handler: DocxHandler, tmp_path: Path) -> None:
    from unittest.mock import patch
    target = tmp_path / "setsave.docx"
    doc = Document()
    doc.add_paragraph("Test")
    doc.save(str(target))
    with patch("docx.document.Document.save", side_effect=RuntimeError("save fail")):
        with pytest.raises(HandlerError, match="failed to save docx"):
            handler.structured_set(target, "paragraph:0", {"text": "new"})


def test_structured_delete_save_exception(handler: DocxHandler, tmp_path: Path) -> None:
    from unittest.mock import patch
    target = tmp_path / "delsave.docx"
    doc = Document()
    doc.add_paragraph("Test")
    doc.save(str(target))
    with patch("docx.document.Document.save", side_effect=RuntimeError("save fail")):
        with pytest.raises(HandlerError, match="failed to save docx"):
            handler.structured_delete(target, "paragraph:0")


def test_structured_delete_paragraph_no_parent(handler: DocxHandler, tmp_path: Path) -> None:
    from unittest.mock import patch, MagicMock
    target = tmp_path / "noparen.docx"
    doc = Document()
    doc.add_paragraph("Test")
    doc.save(str(target))
    # Mock the Document returned by _open so we control _element.getparent()
    mock_doc = MagicMock()
    mock_p = MagicMock()
    mock_p._element.getparent.return_value = None
    mock_doc.paragraphs = [mock_p]
    with patch("dokumen_pintar.handlers.docx_handler._open", return_value=mock_doc):
        with pytest.raises(HandlerError, match="no parent"):
            handler.structured_delete(target, "paragraph:0")


def test_structured_delete_table_no_parent(handler: DocxHandler, tmp_path: Path) -> None:
    from unittest.mock import patch, MagicMock
    target = tmp_path / "notblpar.docx"
    doc = Document()
    doc.add_table(rows=1, cols=1)
    doc.save(str(target))
    mock_doc = MagicMock()
    mock_t = MagicMock()
    mock_t._element.getparent.return_value = None
    mock_doc.tables = [mock_t]
    with patch("dokumen_pintar.handlers.docx_handler._open", return_value=mock_doc):
        with pytest.raises(HandlerError, match="no parent"):
            handler.structured_delete(target, "table:0")


def test_heading_level_non_digit_tail(handler: DocxHandler) -> None:
    from dokumen_pintar.handlers.docx_handler import _heading_level
    # "Heading X" where X is not a digit → 35->37 branch False → falls through to title check
    assert _heading_level("Heading Alpha") is None
    assert _heading_level("Title") == 0


def test_extract_for_search_empty_table_cell(handler: DocxHandler, tmp_path: Path) -> None:
    target = tmp_path / "empcell.docx"
    doc = Document()
    t = doc.add_table(rows=2, cols=2)
    t.cell(0, 0).text = "A"
    t.cell(0, 1).text = ""  # empty cell
    t.cell(1, 0).text = ""  # empty cell
    t.cell(1, 1).text = "D"
    doc.save(str(target))
    result = handler.extract_for_search(target)
    assert "A" in result
    assert "D" in result



# ── v1.1.0 Quick Win 2.2: struct_get table sub-expressions ──


def _make_table_docx(target: Path) -> None:
    """Build a 3x3 docx table for sub-expression tests."""
    doc = Document()
    t = doc.add_table(rows=3, cols=3)
    cells = [
        ["Endpoint", "Method", "Description"],
        ["/api/sap", "POST", "Trigger sync"],
        ["/api/log", "GET", "Get logs"],
    ]
    for r, row in enumerate(cells):
        for c, text in enumerate(row):
            t.cell(r, c).text = text
    doc.save(str(target))


def test_structured_get_table_cell_a1(handler: DocxHandler, tmp_path: Path) -> None:
    target = tmp_path / "tbl_a1.docx"
    _make_table_docx(target)
    assert handler.structured_get(target, "table:0!A1") == "Endpoint"
    assert handler.structured_get(target, "table:0!B2") == "POST"
    assert handler.structured_get(target, "table:0!C3") == "Get logs"


def test_structured_get_table_row(handler: DocxHandler, tmp_path: Path) -> None:
    target = tmp_path / "tbl_row.docx"
    _make_table_docx(target)
    row = handler.structured_get(target, "table:0!row:1")
    assert row == ["/api/sap", "POST", "Trigger sync"]


def test_structured_get_table_col(handler: DocxHandler, tmp_path: Path) -> None:
    target = tmp_path / "tbl_col.docx"
    _make_table_docx(target)
    col = handler.structured_get(target, "table:0!col:0")
    assert col == ["Endpoint", "/api/sap", "/api/log"]


def test_structured_get_table_subexpr_invalid(handler: DocxHandler, tmp_path: Path) -> None:
    target = tmp_path / "tbl_inv.docx"
    _make_table_docx(target)
    # Empty sub-expr after !
    with pytest.raises(HandlerError, match="empty table sub-expression"):
        handler.structured_get(target, "table:0!")
    # Nonsense sub-expr
    with pytest.raises(HandlerError, match="unrecognised table sub-expression"):
        handler.structured_get(target, "table:0!nonsense")


def test_structured_get_table_subexpr_row_out_of_range(
    handler: DocxHandler, tmp_path: Path
) -> None:
    target = tmp_path / "tbl_oor.docx"
    _make_table_docx(target)
    with pytest.raises(HandlerError, match="row index"):
        handler.structured_get(target, "table:0!row:99")
    with pytest.raises(HandlerError, match="col index"):
        handler.structured_get(target, "table:0!col:99")


def test_structured_get_table_a1_out_of_range(handler: DocxHandler, tmp_path: Path) -> None:
    target = tmp_path / "tbl_a1oor.docx"
    _make_table_docx(target)
    # Z column doesn't exist in 3-col table
    with pytest.raises(HandlerError, match="cell column"):
        handler.structured_get(target, "table:0!Z1")
    # row 99 doesn't exist
    with pytest.raises(HandlerError, match="cell row"):
        handler.structured_get(target, "table:0!A99")


def test_structured_get_table_a1_double_letter_columns(
    handler: DocxHandler, tmp_path: Path
) -> None:
    """Verify A1 notation parses double-letter columns (AA = col 26, AB = 27, ...)."""
    from dokumen_pintar.handlers.docx_handler import _parse_table_subexpr

    # 30-column table semantics, just exercise the parser directly.
    kind, r, c = _parse_table_subexpr("AA1", ncols=30, nrows=2)
    assert kind == "cell"
    assert r == 0
    assert c == 26
    kind, r, c = _parse_table_subexpr("AB2", ncols=30, nrows=5)
    assert kind == "cell"
    assert r == 1
    assert c == 27



# ── v1.1.0 B.5: paragraph_runs structured access ──


def test_paragraph_runs_get(handler: DocxHandler, tmp_path: Path) -> None:
    """`paragraph_runs:N` returns each run's text + bold/italic/underline flags."""
    target = tmp_path / "runs.docx"
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("Plain ")
    bold = p.add_run("bold ")
    bold.bold = True
    italic = p.add_run("italic")
    italic.italic = True
    doc.save(str(target))
    result = handler.structured_get(target, "paragraph_runs:0")
    assert result["index"] == 0
    assert result["text"] == "Plain bold italic"
    runs = result["runs"]
    assert runs[0]["text"] == "Plain "
    assert runs[0]["bold"] is False
    assert runs[1]["bold"] is True
    assert runs[2]["italic"] is True


def test_paragraph_runs_set_replaces_runs(handler: DocxHandler, tmp_path: Path) -> None:
    target = tmp_path / "runs_set.docx"
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("old plain")
    doc.save(str(target))
    handler.structured_set(
        target,
        "paragraph_runs:0",
        [
            {"text": "Hello "},
            {"text": "world", "bold": True, "italic": True, "underline": True},
        ],
    )
    out = Document(str(target))
    runs = list(out.paragraphs[0].runs)
    assert [r.text for r in runs] == ["Hello ", "world"]
    assert runs[1].bold is True
    assert runs[1].italic is True
    assert runs[1].underline is True


def test_paragraph_runs_set_value_must_be_list(
    handler: DocxHandler, tmp_path: Path
) -> None:
    target = tmp_path / "bad_value.docx"
    doc = Document()
    doc.add_paragraph("x")
    doc.save(str(target))
    with pytest.raises(HandlerError, match="must be a list"):
        handler.structured_set(target, "paragraph_runs:0", {"text": "wrong"})


def test_paragraph_runs_set_run_spec_must_be_dict(
    handler: DocxHandler, tmp_path: Path
) -> None:
    target = tmp_path / "bad_run.docx"
    doc = Document()
    doc.add_paragraph("x")
    doc.save(str(target))
    with pytest.raises(HandlerError, match="must be a dict"):
        handler.structured_set(target, "paragraph_runs:0", ["not a dict"])


def test_paragraph_runs_set_text_must_be_str(
    handler: DocxHandler, tmp_path: Path
) -> None:
    target = tmp_path / "bad_text.docx"
    doc = Document()
    doc.add_paragraph("x")
    doc.save(str(target))
    with pytest.raises(HandlerError, match="must be a string"):
        handler.structured_set(target, "paragraph_runs:0", [{"text": 123}])


def test_paragraph_runs_get_out_of_range(handler: DocxHandler, tmp_path: Path) -> None:
    target = tmp_path / "oor.docx"
    doc = Document()
    doc.add_paragraph("only one")
    doc.save(str(target))
    with pytest.raises(HandlerError, match="out of range"):
        handler.structured_get(target, "paragraph_runs:99")


def test_paragraph_runs_set_out_of_range(handler: DocxHandler, tmp_path: Path) -> None:
    target = tmp_path / "oor_set.docx"
    doc = Document()
    doc.add_paragraph("only one")
    doc.save(str(target))
    with pytest.raises(HandlerError, match="out of range"):
        handler.structured_set(target, "paragraph_runs:99", [{"text": "x"}])