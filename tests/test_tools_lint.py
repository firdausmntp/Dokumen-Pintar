"""Tests for :mod:`dokumen_pintar.tools.lint` and the lint subsystem."""

from __future__ import annotations

from pathlib import Path
from typing import Callable

import pytest
from docx import Document
from mcp.server.fastmcp import FastMCP

from dokumen_pintar.config import AppConfig
from dokumen_pintar.context import build_context
from dokumen_pintar.errors import HandlerError, UnsupportedFormatError, ValidationError
from dokumen_pintar.lint import default_registry
from dokumen_pintar.tools import lint as lint_tool


def _setup(cfg: AppConfig) -> tuple[FastMCP, ...]:
    ctx = build_context(cfg)
    mcp = FastMCP(name="t-lint")
    lint_tool.register(mcp, ctx)
    return mcp, ctx


def _tool(mcp: FastMCP, name: str):
    return mcp._tool_manager._tools[name].fn


# ── happy paths ──


def test_document_lint_default_preset(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    target = docs_dir / "doc.docx"
    doc = Document()
    doc.add_paragraph("trailing space   ")
    doc.add_heading("", level=1)  # empty heading
    doc.save(str(target))
    mcp, _ctx = _setup(make_config())
    result = _tool(mcp, "document_lint")(path=str(target))
    rule_ids = {i["rule"] for i in result["issues"]}
    assert "trailing_whitespace" in rule_ids
    assert "empty_heading" in rule_ids
    assert result["summary"]["warnings"] >= 2


def test_document_lint_severity_filter(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    target = docs_dir / "sev.docx"
    doc = Document()
    doc.add_heading("only one heading", level=1)
    doc.add_paragraph("trailing  ")  # warn
    doc.save(str(target))
    mcp, _ctx = _setup(make_config())
    result = _tool(mcp, "document_lint")(path=str(target), severity_filter="warn")
    severities = {i["severity"] for i in result["issues"]}
    assert severities <= {"warn"}


def test_document_lint_invalid_severity(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    target = docs_dir / "x.docx"
    Document().save(str(target))
    mcp, _ctx = _setup(make_config())
    with pytest.raises(ValidationError, match="severity_filter"):
        _tool(mcp, "document_lint")(path=str(target), severity_filter="critical")


def test_document_lint_wrong_extension(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    target = docs_dir / "x.txt"
    target.write_text("hi", encoding="utf-8")
    mcp, _ctx = _setup(make_config())
    with pytest.raises(UnsupportedFormatError, match="must be .docx"):
        _tool(mcp, "document_lint")(path=str(target))


def test_document_lint_corrupted_file(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    target = docs_dir / "broken.docx"
    target.write_bytes(b"PK\x03\x04bad")
    mcp, _ctx = _setup(make_config())
    with pytest.raises(HandlerError):
        _tool(mcp, "document_lint")(path=str(target))


def test_document_lint_unknown_rule_or_preset(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    target = docs_dir / "u.docx"
    Document().save(str(target))
    mcp, _ctx = _setup(make_config())
    with pytest.raises(ValidationError, match="unknown rule or preset"):
        _tool(mcp, "document_lint")(path=str(target), rules=["does_not_exist"])


def test_document_lint_rules_invalid_type(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    target = docs_dir / "u.docx"
    Document().save(str(target))
    mcp, _ctx = _setup(make_config())
    with pytest.raises(ValidationError, match="must be a preset name"):
        _tool(mcp, "document_lint")(path=str(target), rules=42)
    with pytest.raises(ValidationError, match="must contain strings"):
        _tool(mcp, "document_lint")(path=str(target), rules=[42])


def test_document_lint_rules_combined_preset_and_rule(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    """The list form accepts both preset names AND individual rule IDs."""
    docs_dir, _ = tmp_roots
    target = docs_dir / "comb.docx"
    doc = Document()
    doc.add_paragraph("trailing  ")
    doc.save(str(target))
    mcp, _ctx = _setup(make_config())
    result = _tool(mcp, "document_lint")(path=str(target), rules=["default", "trailing_whitespace"])
    # default already includes trailing_whitespace; combined list
    # de-duplicates so the rule fires only once per paragraph.
    assert any(i["rule"] == "trailing_whitespace" for i in result["issues"])


def test_document_lint_duplicate_heading(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    target = docs_dir / "dup.docx"
    doc = Document()
    doc.add_heading("BAB I", level=1)
    doc.add_paragraph("body 1")
    doc.add_heading("BAB I", level=1)  # duplicate
    doc.save(str(target))
    mcp, _ctx = _setup(make_config())
    result = _tool(mcp, "document_lint")(path=str(target))
    dup = [i for i in result["issues"] if i["rule"] == "duplicate_heading"]
    assert dup


def test_document_lint_heading_hierarchy_skip(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    target = docs_dir / "hsk.docx"
    doc = Document()
    doc.add_heading("Chapter", level=1)
    doc.add_heading("Subsection", level=3)  # skips level 2
    doc.save(str(target))
    mcp, _ctx = _setup(make_config())
    result = _tool(mcp, "document_lint")(path=str(target))
    skip = [i for i in result["issues"] if i["rule"] == "heading_hierarchy_skip"]
    assert skip


def test_document_lint_heading_hierarchy_skip_after_title(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    """Title (level 0) resets the hierarchy floor; H1 after Title is fine."""
    docs_dir, _ = tmp_roots
    target = docs_dir / "title_h1.docx"
    doc = Document()
    doc.add_paragraph("Cover", style="Title")
    doc.add_heading("Chapter", level=1)
    doc.save(str(target))
    mcp, _ctx = _setup(make_config())
    result = _tool(mcp, "document_lint")(path=str(target))
    assert not any(i["rule"] == "heading_hierarchy_skip" for i in result["issues"])


def test_document_lint_title_case_id(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    target = docs_dir / "tc.docx"
    doc = Document()
    doc.add_heading("BAB i pendahuluan", level=1)  # not title-cased
    doc.save(str(target))
    mcp, _ctx = _setup(make_config())
    result = _tool(mcp, "document_lint")(path=str(target), rules="academic_id")
    tc = [i for i in result["issues"] if i["rule"] == "title_case_id"]
    assert tc


def test_document_lint_required_section_missing(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    target = docs_dir / "rs.docx"
    doc = Document()
    doc.add_heading("Pendahuluan", level=1)
    doc.save(str(target))
    mcp, _ctx = _setup(make_config())
    result = _tool(mcp, "document_lint")(
        path=str(target), rules=["required_section_lembar_pengesahan"]
    )
    assert any(i["rule"] == "required_section_lembar_pengesahan" for i in result["issues"])


def test_document_lint_required_section_present(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    target = docs_dir / "ps.docx"
    doc = Document()
    doc.add_heading("LEMBAR PENGESAHAN", level=1)
    doc.save(str(target))
    mcp, _ctx = _setup(make_config())
    result = _tool(mcp, "document_lint")(
        path=str(target), rules=["required_section_lembar_pengesahan"]
    )
    assert not any(i["rule"] == "required_section_lembar_pengesahan" for i in result["issues"])


def test_document_lint_kp_preset_includes_logbook(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    target = docs_dir / "kp.docx"
    doc = Document()
    doc.add_heading("Cover", level=1)
    doc.save(str(target))
    mcp, _ctx = _setup(make_config())
    result = _tool(mcp, "document_lint")(path=str(target), rules="academic_id_kp")
    rule_ids = {i["rule"] for i in result["issues"]}
    assert "required_section_log_book" in rule_ids
    # The KP preset extends academic_id which extends default.
    assert "required_section_lembar_pengesahan" in rule_ids


def test_document_lint_skripsi_preset(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    target = docs_dir / "sk.docx"
    Document().save(str(target))
    mcp, _ctx = _setup(make_config())
    result = _tool(mcp, "document_lint")(path=str(target), rules="academic_id_skripsi")
    rule_ids = {i["rule"] for i in result["issues"]}
    assert "required_section_abstrak" in rule_ids
    assert "required_section_metodologi" in rule_ids


# ── document_lint_fix ──


def test_document_lint_fix_dry_run_lists_planned(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    target = docs_dir / "fx.docx"
    doc = Document()
    doc.add_paragraph("trailing  ")
    doc.save(str(target))
    original_size = target.stat().st_size
    mcp, _ctx = _setup(make_config())
    result = _tool(mcp, "document_lint_fix")(path=str(target))
    assert result["dry_run"] is True
    assert result["fix_count"] >= 1
    # File untouched.
    assert target.stat().st_size == original_size


def test_document_lint_fix_apply_trailing_whitespace(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    target = docs_dir / "wsfix.docx"
    doc = Document()
    doc.add_paragraph("clean")
    doc.add_paragraph("trailing   ")
    doc.save(str(target))
    mcp, _ctx = _setup(make_config())
    result = _tool(mcp, "document_lint_fix")(path=str(target), dry_run=False)
    assert result["applied_count"] >= 1
    out = Document(str(target))
    # All paragraphs should be rstripped now.
    for p in out.paragraphs:
        assert p.text == p.text.rstrip()


def test_document_lint_fix_apply_empty_heading_removal(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    target = docs_dir / "ehfix.docx"
    doc = Document()
    doc.add_heading("Real heading", level=1)
    doc.add_heading("", level=1)  # empty heading
    doc.add_paragraph("body")
    doc.save(str(target))
    mcp, _ctx = _setup(make_config())
    _tool(mcp, "document_lint_fix")(path=str(target), dry_run=False)
    out = Document(str(target))
    headings = [p for p in out.paragraphs if (p.style.name or "").startswith("Heading")]
    # The empty heading should have been removed.
    assert all(h.text.strip() for h in headings)


def test_document_lint_fix_only_severities(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    target = docs_dir / "sf.docx"
    doc = Document()
    doc.add_paragraph("trailing  ")
    doc.add_heading("", level=1)
    doc.save(str(target))
    mcp, _ctx = _setup(make_config())
    # Restrict to 'error' only - should match nothing in the default preset.
    result = _tool(mcp, "document_lint_fix")(path=str(target), only_severities=["error"])
    assert result["fix_count"] == 0


def test_document_lint_fix_only_severities_invalid(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    target = docs_dir / "sfi.docx"
    Document().save(str(target))
    mcp, _ctx = _setup(make_config())
    with pytest.raises(ValidationError, match="invalid values"):
        _tool(mcp, "document_lint_fix")(path=str(target), only_severities=["bogus"])


def test_document_lint_fix_wrong_extension(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    target = docs_dir / "fx.txt"
    target.write_text("hi", encoding="utf-8")
    mcp, _ctx = _setup(make_config())
    with pytest.raises(UnsupportedFormatError):
        _tool(mcp, "document_lint_fix")(path=str(target))


def test_document_lint_fix_file_not_found(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    missing = docs_dir / "no.docx"
    mcp, _ctx = _setup(make_config())
    with pytest.raises(ValidationError, match="not found"):
        _tool(mcp, "document_lint_fix")(path=str(missing))


def test_document_lint_fix_save_failure(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    from unittest.mock import patch

    docs_dir, _ = tmp_roots
    target = docs_dir / "fs.docx"
    doc = Document()
    doc.add_paragraph("trailing   ")
    doc.save(str(target))
    mcp, _ctx = _setup(make_config())
    with patch("docx.document.Document.save", side_effect=OSError("disk full")):
        with pytest.raises(HandlerError, match="failed to save"):
            _tool(mcp, "document_lint_fix")(path=str(target), dry_run=False)


def test_document_lint_fix_apply_fix_returns_false(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    """A rule whose apply_fix returns False lands in the `skipped` bucket."""
    from unittest.mock import patch

    docs_dir, _ = tmp_roots
    target = docs_dir / "ret.docx"
    doc = Document()
    doc.add_paragraph("trailing   ")
    doc.save(str(target))
    mcp, _ctx = _setup(make_config())
    with patch.object(
        default_registry.rule("trailing_whitespace"),
        "apply_fix",
        return_value=False,
    ):
        result = _tool(mcp, "document_lint_fix")(path=str(target), dry_run=False)
    assert result["applied_count"] == 0
    assert len(result["skipped"]) >= 1


# ── lint subsystem internals ──


def test_register_rule_rejects_empty_id() -> None:
    from dokumen_pintar.lint.base import LintRule, _Registry

    class _NoId(LintRule):
        id = ""

    reg = _Registry()
    with pytest.raises(ValueError, match="empty id"):
        reg.register(_NoId)


def test_resolve_preset_unknown() -> None:
    from dokumen_pintar.lint.base import _Registry

    reg = _Registry()
    with pytest.raises(KeyError):
        reg.resolve_preset("nope")


def test_resolve_preset_cycle_detection() -> None:
    from dokumen_pintar.lint.base import _Registry

    reg = _Registry()
    reg.add_preset("a", rules=[], extends="b")
    reg.add_preset("b", rules=[], extends="a")
    with pytest.raises(ValueError, match="cycle"):
        reg.resolve_preset("a")


def test_resolve_preset_unknown_in_extends_chain() -> None:
    from dokumen_pintar.lint.base import _Registry

    reg = _Registry()
    reg.add_preset("a", rules=[], extends="missing")
    with pytest.raises(KeyError, match="extends chain"):
        reg.resolve_preset("a")


def test_lint_rule_default_apply_fix_returns_false() -> None:
    from unittest.mock import MagicMock

    from dokumen_pintar.lint.base import Issue, LintRule

    rule = LintRule()
    issue = Issue(rule="x", severity="warn", location={})
    assert rule.apply_fix(MagicMock(), issue) is False


def test_lint_rule_check_default_raises() -> None:
    from unittest.mock import MagicMock

    from dokumen_pintar.lint.base import LintRule

    rule = LintRule()
    with pytest.raises(NotImplementedError):
        list(rule.check(MagicMock()))


def test_unknown_rule_lookup() -> None:
    from dokumen_pintar.lint.base import _Registry

    reg = _Registry()
    with pytest.raises(KeyError, match="unknown rule"):
        reg.rule("nope")


def test_required_section_rule_no_pattern_yields_nothing() -> None:
    """Subclass without pattern returns no issues."""
    from unittest.mock import MagicMock

    from dokumen_pintar.lint.rules import RequiredSectionRule

    rule = RequiredSectionRule()
    assert list(rule.check(MagicMock())) == []


def test_trailing_whitespace_apply_fix_invalid_index() -> None:
    """apply_fix with out-of-range paragraph index returns False."""
    from unittest.mock import MagicMock

    from dokumen_pintar.lint.base import Issue
    from dokumen_pintar.lint.rules import TrailingWhitespaceRule

    rule = TrailingWhitespaceRule()
    fake_doc = MagicMock()
    fake_doc.paragraphs = []
    issue = Issue(
        rule="trailing_whitespace",
        severity="warn",
        location={"paragraph": 99},
    )
    assert rule.apply_fix(fake_doc, issue) is False
    issue_no_idx = Issue(rule="trailing_whitespace", severity="warn", location={})
    assert rule.apply_fix(fake_doc, issue_no_idx) is False


def test_empty_heading_apply_fix_invalid_index() -> None:
    from unittest.mock import MagicMock

    from dokumen_pintar.lint.base import Issue
    from dokumen_pintar.lint.rules import EmptyHeadingRule

    rule = EmptyHeadingRule()
    fake_doc = MagicMock()
    fake_doc.paragraphs = []
    issue = Issue(rule="empty_heading", severity="warn", location={"paragraph": 5})
    assert rule.apply_fix(fake_doc, issue) is False
    issue_no_idx = Issue(rule="empty_heading", severity="warn", location={})
    assert rule.apply_fix(fake_doc, issue_no_idx) is False



def test_title_case_id_skips_non_heading(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    """Plain paragraphs (level=None) skip the title-case check entirely."""
    docs_dir, _ = tmp_roots
    target = docs_dir / "tcskip.docx"
    doc = Document()
    doc.add_paragraph("plain body without title-case rules")
    doc.save(str(target))
    mcp, _ctx = _setup(make_config())
    result = _tool(mcp, "document_lint")(
        path=str(target), rules=["title_case_id"]
    )
    assert result["issues"] == []


def test_title_case_id_lowercases_conjunctions(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    """`dan`, `atau`, etc. lowercased mid-heading is fine."""
    docs_dir, _ = tmp_roots
    target = docs_dir / "conj.docx"
    doc = Document()
    # 'Penelitian dan Pembahasan' - conjunction `dan` mid-heading is acceptable.
    doc.add_heading("Penelitian dan Pembahasan", level=1)
    doc.save(str(target))
    mcp, _ctx = _setup(make_config())
    result = _tool(mcp, "document_lint")(
        path=str(target), rules=["title_case_id"]
    )
    assert result["issues"] == []


def test_required_section_skips_non_heading_paragraphs(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    """Required-section pattern only matches heading-styled paragraphs.

    Plain body paragraph with the same text shouldn't satisfy the rule.
    """
    docs_dir, _ = tmp_roots
    target = docs_dir / "rsbody.docx"
    doc = Document()
    doc.add_paragraph("LEMBAR PENGESAHAN")  # body text, not a heading
    doc.save(str(target))
    mcp, _ctx = _setup(make_config())
    result = _tool(mcp, "document_lint")(
        path=str(target), rules=["required_section_lembar_pengesahan"]
    )
    # Section not found -> issue raised.
    assert any(
        i["rule"] == "required_section_lembar_pengesahan"
        for i in result["issues"]
    )


def test_lint_fix_skips_non_autofixable(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    """`document_lint_fix` filters non-auto_fixable issues out of the plan
    (covers the `if not issue.auto_fixable: continue` branch in the loop)."""
    docs_dir, _ = tmp_roots
    target = docs_dir / "nofix.docx"
    doc = Document()
    # duplicate_heading is NOT auto_fixable; trailing_whitespace IS.
    doc.add_heading("BAB I", level=1)
    doc.add_heading("BAB I", level=1)  # duplicate
    doc.add_paragraph("trailing  ")
    doc.save(str(target))
    mcp, _ctx = _setup(make_config())
    result = _tool(mcp, "document_lint_fix")(path=str(target))
    # The duplicate heading is in the lint result but skipped from the fix plan.
    fix_rules = {f["rule"] for f in result["fixes"]}
    assert "duplicate_heading" not in fix_rules
    assert "trailing_whitespace" in fix_rules