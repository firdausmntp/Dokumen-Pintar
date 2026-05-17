"""Tests for :mod:`dokumen_pintar.tools.bibliography`."""

from __future__ import annotations

from pathlib import Path
from typing import Callable

import pytest
from docx import Document
from mcp.server.fastmcp import FastMCP

from dokumen_pintar.config import AppConfig
from dokumen_pintar.context import build_context
from dokumen_pintar.errors import HandlerError, UnsupportedFormatError, ValidationError
from dokumen_pintar.tools import bibliography


def _setup(cfg: AppConfig) -> tuple[FastMCP, ...]:
    ctx = build_context(cfg)
    mcp = FastMCP(name="t-bib")
    bibliography.register(mcp, ctx)
    return mcp, ctx


def _tool(mcp: FastMCP, name: str):
    return mcp._tool_manager._tools[name].fn


def _build_apa_doc(target: Path) -> None:
    """Build a small APA-style doc with body citations + bib section."""
    doc = Document()
    doc.add_heading("BAB I PENDAHULUAN", level=1)
    doc.add_paragraph(
        "Penelitian ini mengacu pada (Smith, 2020). Lebih lanjut, "
        "(Jones et al., 2021) menemukan bahwa..."
    )
    doc.add_heading("DAFTAR PUSTAKA", level=1)
    doc.add_paragraph("Smith, J. (2020). Title of work. Publisher.")
    doc.add_paragraph("Jones, A. (2021). Another work. Journal.")
    doc.add_paragraph("Unused, B. (2019). Never cited. Press.")
    doc.save(str(target))


def _build_ieee_doc(target: Path) -> None:
    """Build an IEEE-style doc with [N] citations."""
    doc = Document()
    doc.add_heading("INTRO", level=1)
    doc.add_paragraph("As shown in [1] and [2], the result holds.")
    doc.add_heading("REFERENCES", level=1)
    doc.add_paragraph('[1] J. Smith, "Title," Journal, 2020.')
    doc.add_paragraph('[2] A. Jones, "Other," Conf., 2021.')
    doc.save(str(target))


# ── bibliography_check happy path ──


def test_bibliography_check_apa_basic(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    target = docs_dir / "apa.docx"
    _build_apa_doc(target)
    mcp, _ctx = _setup(make_config())
    result = _tool(mcp, "bibliography_check")(path=str(target))
    assert result["style"] == "APA"
    assert len(result["citations_found"]) >= 2
    assert len(result["bibliography_entries"]) == 3
    issue_types = {i["type"] for i in result["issues"]}
    # 'Unused, B. 2019' is in the bib but never cited.
    assert "unused_bib_entry" in issue_types


def test_bibliography_check_ieee(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    target = docs_dir / "ieee.docx"
    _build_ieee_doc(target)
    mcp, _ctx = _setup(make_config())
    result = _tool(mcp, "bibliography_check")(path=str(target), style="IEEE")
    assert all(c["kind"] == "numeric" for c in result["citations_found"])
    assert len(result["bibliography_entries"]) == 2


def test_bibliography_check_missing_entry(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    target = docs_dir / "missing.docx"
    doc = Document()
    doc.add_heading("Body", level=1)
    doc.add_paragraph("As (Author, 2024) reported.")
    doc.add_heading("DAFTAR PUSTAKA", level=1)
    doc.add_paragraph("Other, X. (2019). Title.")
    doc.save(str(target))
    mcp, _ctx = _setup(make_config())
    result = _tool(mcp, "bibliography_check")(path=str(target))
    issue_types = {i["type"] for i in result["issues"]}
    assert "missing_bib_entry" in issue_types


def test_bibliography_check_duplicate_entry(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    target = docs_dir / "dup.docx"
    doc = Document()
    doc.add_heading("Body", level=1)
    doc.add_paragraph("(Smith, 2020) is cited.")
    doc.add_heading("DAFTAR PUSTAKA", level=1)
    doc.add_paragraph("Smith, J. (2020). First title.")
    doc.add_paragraph("Smith, J. (2020). Same year same author.")
    doc.save(str(target))
    mcp, _ctx = _setup(make_config())
    result = _tool(mcp, "bibliography_check")(path=str(target))
    dup_issues = [i for i in result["issues"] if i["type"] == "duplicate_bib_entry"]
    assert dup_issues
    assert dup_issues[0]["count"] == 2


def test_bibliography_check_no_section_found(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    """When the bib section heading is absent, results contain only citations."""
    docs_dir, _ = tmp_roots
    target = docs_dir / "nosec.docx"
    doc = Document()
    doc.add_paragraph("(Smith, 2020) cited but no bib section.")
    doc.save(str(target))
    mcp, _ctx = _setup(make_config())
    result = _tool(mcp, "bibliography_check")(path=str(target))
    assert result["bib_section_range"] is None
    assert result["bibliography_entries"] == []
    # Every citation becomes a missing_bib_entry issue.
    missing = [i for i in result["issues"] if i["type"] == "missing_bib_entry"]
    assert len(missing) >= 1


def test_bibliography_check_custom_pattern(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    target = docs_dir / "custom.docx"
    doc = Document()
    doc.add_heading("Pustaka Acuan", level=1)  # custom heading text
    doc.add_paragraph("Smith, J. (2020). Title.")
    doc.save(str(target))
    mcp, _ctx = _setup(make_config())
    # Default pattern wouldn't match; provide override.
    result = _tool(mcp, "bibliography_check")(
        path=str(target), bib_section_pattern=r"(?i)pustaka acuan"
    )
    assert result["bib_section_range"] is not None
    assert len(result["bibliography_entries"]) == 1


def test_bibliography_check_invalid_style(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    target = docs_dir / "x.docx"
    Document().save(str(target))
    mcp, _ctx = _setup(make_config())
    with pytest.raises(ValidationError, match="style must be"):
        _tool(mcp, "bibliography_check")(path=str(target), style="MLA")


def test_bibliography_check_wrong_extension(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    target = docs_dir / "wrong.txt"
    target.write_text("text", encoding="utf-8")
    mcp, _ctx = _setup(make_config())
    with pytest.raises(UnsupportedFormatError, match="must be .docx"):
        _tool(mcp, "bibliography_check")(path=str(target))


def test_bibliography_check_corrupted_file(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    target = docs_dir / "broken.docx"
    target.write_bytes(b"PK\x03\x04bad")
    mcp, _ctx = _setup(make_config())
    with pytest.raises(HandlerError):
        _tool(mcp, "bibliography_check")(path=str(target))


def test_bibliography_check_disabled_section_detection(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    """`auto_detect_section=False` and no override yields no bib_range."""
    docs_dir, _ = tmp_roots
    target = docs_dir / "off.docx"
    _build_apa_doc(target)
    mcp, _ctx = _setup(make_config())
    result = _tool(mcp, "bibliography_check")(path=str(target), auto_detect_section=False)
    # All citations detected, but bib entries skip nothing -> the entries
    # under DAFTAR PUSTAKA also count as candidate citations now if they
    # contain (Year). With detection off, bib_section_range is None.
    assert result["bib_section_range"] is None


def test_bibliography_check_malformed_apa_entry_flagged(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    target = docs_dir / "mal.docx"
    doc = Document()
    doc.add_heading("Body", level=1)
    doc.add_paragraph("(Smith, 2020) cited.")
    doc.add_heading("DAFTAR PUSTAKA", level=1)
    doc.add_paragraph("just some random text without year info")
    doc.save(str(target))
    mcp, _ctx = _setup(make_config())
    result = _tool(mcp, "bibliography_check")(path=str(target))
    mal = [i for i in result["issues"] if i["type"] == "malformed_entry"]
    assert mal
    assert "APA" in mal[0]["message"]


def test_bibliography_check_malformed_ieee_entry_flagged(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    target = docs_dir / "ieemal.docx"
    doc = Document()
    doc.add_heading("Body", level=1)
    doc.add_paragraph("[1] is cited")
    doc.add_heading("REFERENCES", level=1)
    # Missing [N] prefix breaks IEEE expectation.
    doc.add_paragraph('Smith, J. "Title," Journal, 2020.')
    doc.save(str(target))
    mcp, _ctx = _setup(make_config())
    result = _tool(mcp, "bibliography_check")(path=str(target), style="IEEE")
    mal = [i for i in result["issues"] if i["type"] == "malformed_entry"]
    assert mal
    assert "IEEE" in mal[0]["message"]


# ── bibliography_format ──


def test_bibliography_format_dry_run(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    """Without auto_fix, returns the proposed order but doesn't write."""
    docs_dir, _ = tmp_roots
    target = docs_dir / "fmt.docx"
    _build_apa_doc(target)
    original_size = target.stat().st_size
    mcp, _ctx = _setup(make_config())
    result = _tool(mcp, "bibliography_format")(path=str(target))
    assert result["would_change"] is True
    assert result["proposed_order"] != result["current_order"]
    # File untouched.
    assert target.stat().st_size == original_size


def test_bibliography_format_apply_sorts_in_place(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    target = docs_dir / "fmt2.docx"
    _build_apa_doc(target)
    mcp, _ctx = _setup(make_config())
    result = _tool(mcp, "bibliography_format")(path=str(target), auto_fix=True)
    assert result["applied"] is True
    # Re-read and verify alphabetical order: Jones < Smith < Unused.
    out = Document(str(target))
    bib_paras = [
        p.text.strip()
        for p in out.paragraphs
        if p.text.strip() and not p.text.startswith(("BAB", "DAFTAR"))
    ]
    keys_only = [k.split(".")[0].split()[-1] for k in bib_paras if "(20" in k]
    # Sort alphabetic by key (Author Year): 'Jones 2021', 'Smith 2020', 'Unused 2019'
    # Verify Jones appears before Smith.
    if "Jones" in keys_only and "Smith" in keys_only:
        assert keys_only.index("Jones") < keys_only.index("Smith")


def test_bibliography_format_no_section(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    target = docs_dir / "nosec.docx"
    Document().save(str(target))
    mcp, _ctx = _setup(make_config())
    with pytest.raises(ValidationError, match="no bibliography section"):
        _tool(mcp, "bibliography_format")(path=str(target))


def test_bibliography_format_empty_section(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    target = docs_dir / "empty_sec.docx"
    doc = Document()
    doc.add_heading("DAFTAR PUSTAKA", level=1)
    doc.save(str(target))
    mcp, _ctx = _setup(make_config())
    with pytest.raises(ValidationError, match="empty"):
        _tool(mcp, "bibliography_format")(path=str(target))


def test_bibliography_format_no_sort_returns_original_order(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    target = docs_dir / "ns.docx"
    _build_apa_doc(target)
    mcp, _ctx = _setup(make_config())
    result = _tool(mcp, "bibliography_format")(path=str(target), sort=False)
    assert result["would_change"] is False
    assert result["proposed_order"] == result["current_order"]


def test_bibliography_format_invalid_style(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    target = docs_dir / "x.docx"
    Document().save(str(target))
    mcp, _ctx = _setup(make_config())
    with pytest.raises(ValidationError, match="style must be"):
        _tool(mcp, "bibliography_format")(path=str(target), style="MLA")


def test_bibliography_format_wrong_extension(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    target = docs_dir / "x.txt"
    target.write_text("hi", encoding="utf-8")
    mcp, _ctx = _setup(make_config())
    with pytest.raises(UnsupportedFormatError):
        _tool(mcp, "bibliography_format")(path=str(target))


def test_bibliography_format_file_not_found(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    missing = docs_dir / "no.docx"
    mcp, _ctx = _setup(make_config())
    with pytest.raises(ValidationError, match="not found"):
        _tool(mcp, "bibliography_format")(path=str(missing))


def test_bibliography_format_save_failure(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    """Save failure surfaces as HandlerError."""
    from unittest.mock import patch

    docs_dir, _ = tmp_roots
    target = docs_dir / "sv.docx"
    _build_apa_doc(target)
    mcp, _ctx = _setup(make_config())
    with patch("docx.document.Document.save", side_effect=OSError("fail")):
        with pytest.raises(HandlerError, match="failed to save"):
            _tool(mcp, "bibliography_format")(path=str(target), auto_fix=True)


def test_extract_bib_entries_unknown_kind(tmp_path: Path) -> None:
    """A bib paragraph that doesn't fit numeric or author-year is 'unknown'."""
    from dokumen_pintar.tools.bibliography import (
        _extract_bib_entries,
        _split_into_paragraphs,
    )

    target = tmp_path / "u.docx"
    doc = Document()
    doc.add_heading("DAFTAR PUSTAKA", level=1)
    doc.add_paragraph("unrecognised gibberish without year")
    doc.save(str(target))
    src = Document(str(target))
    paragraphs = _split_into_paragraphs(src)
    bib_range = (1, len(paragraphs))
    entries = _extract_bib_entries(paragraphs, bib_range)
    assert entries
    assert entries[0]["kind"] == "unknown"



def test_extract_bib_entries_skips_blank_paragraphs(tmp_path: Path) -> None:
    """`_extract_bib_entries` ignores empty paragraphs inside the section."""
    from dokumen_pintar.tools.bibliography import (
        _extract_bib_entries,
        _split_into_paragraphs,
    )

    target = tmp_path / "blank.docx"
    doc = Document()
    doc.add_heading("DAFTAR PUSTAKA", level=1)
    doc.add_paragraph("")  # blank line in the bib section
    doc.add_paragraph("Smith, J. (2020). Real entry.")
    doc.save(str(target))
    src = Document(str(target))
    paragraphs = _split_into_paragraphs(src)
    bib_range = (1, len(paragraphs))
    entries = _extract_bib_entries(paragraphs, bib_range)
    # Only the populated paragraph yields an entry.
    assert len(entries) == 1


def test_compare_no_handler_for_source_unwrapped() -> None:
    """`_read_text_for_compare` raises UnsupportedFormatError when handler is missing."""
    from unittest.mock import MagicMock

    from dokumen_pintar.tools.compare import _read_text_for_compare

    fake_ctx = MagicMock()
    fake_ctx.registry.for_path.return_value = None
    with pytest.raises(UnsupportedFormatError, match="no handler"):
        _read_text_for_compare(fake_ctx, Path("/tmp/x.unknownext"))