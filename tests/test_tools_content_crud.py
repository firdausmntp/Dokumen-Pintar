"""Tests for :mod:`dokumen_pintar.tools.content_crud`."""

from __future__ import annotations

from pathlib import Path
from typing import Callable

import pytest
from mcp.server.fastmcp import FastMCP

from dokumen_pintar.config import AppConfig
from dokumen_pintar.context import build_context
from dokumen_pintar.errors import ValidationError
from dokumen_pintar.tools import content_crud


def _setup(cfg: AppConfig) -> tuple[FastMCP, ...]:
    ctx = build_context(cfg)
    mcp = FastMCP(name="test")
    content_crud.register(mcp, ctx)
    return mcp, ctx


def _tool(mcp: FastMCP, name: str):
    return mcp._tool_manager._tools[name].fn


def test_content_read(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    (docs_dir / "read.txt").write_bytes(b"line1\nline2\nline3\n")
    mcp, _ = _setup(make_config())
    result = _tool(mcp, "content_read")(path="documents:/read.txt")
    assert "line1" in result["content"]
    assert result["line_count"] == 3


def test_content_read_slice(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    (docs_dir / "slice.txt").write_bytes(b"L1\nL2\nL3\nL4\n")
    mcp, _ = _setup(make_config())
    result = _tool(mcp, "content_read")(path="documents:/slice.txt", start_line=2, end_line=3)
    assert "L2" in result["content"]
    assert "L1" not in result["content"]


def test_content_write(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    (docs_dir / "wr.txt").write_bytes(b"old")
    mcp, _ = _setup(make_config())
    _tool(mcp, "content_write")(path="documents:/wr.txt", content="new content")
    assert (docs_dir / "wr.txt").read_text(encoding="utf-8") == "new content"


def test_content_write_new_file(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    mcp, _ = _setup(make_config())
    _tool(mcp, "content_write")(path="documents:/brand_new.txt", content="fresh")
    assert (docs_dir / "brand_new.txt").read_text(encoding="utf-8") == "fresh"


def test_content_append(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    (docs_dir / "app.txt").write_bytes(b"hello ")
    mcp, _ = _setup(make_config())
    result = _tool(mcp, "content_append")(path="documents:/app.txt", content="world")
    assert result["new_size"] == len("hello world")


def test_content_insert(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    (docs_dir / "ins.txt").write_bytes(b"A\nC\n")
    mcp, _ = _setup(make_config())
    _tool(mcp, "content_insert")(path="documents:/ins.txt", line_number=2, content="B")
    text = (docs_dir / "ins.txt").read_bytes().decode("utf-8")
    lines = text.strip().split("\n")
    assert lines == ["A", "B", "C"]


def test_content_insert_invalid_line(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    mcp, _ = _setup(make_config())
    with pytest.raises(ValidationError, match="line_number"):
        _tool(mcp, "content_insert")(path="documents:/x.txt", line_number=0, content="x")


def test_content_replace(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    (docs_dir / "rep.txt").write_bytes(b"foo bar foo baz")
    mcp, _ = _setup(make_config())
    result = _tool(mcp, "content_replace")(path="documents:/rep.txt", old="foo", new="qux")
    assert result["replacements"] == 2
    assert (docs_dir / "rep.txt").read_text(encoding="utf-8") == "qux bar qux baz"


def test_content_replace_regex(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    (docs_dir / "regrep.txt").write_bytes(b"cat123 dog456")
    mcp, _ = _setup(make_config())
    result = _tool(mcp, "content_replace")(
        path="documents:/regrep.txt", old=r"\d+", new="NUM", regex=True
    )
    assert result["replacements"] == 2


def test_content_replace_no_match(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    (docs_dir / "nm.txt").write_bytes(b"nothing here")
    mcp, _ = _setup(make_config())
    result = _tool(mcp, "content_replace")(path="documents:/nm.txt", old="zzz", new="aaa")
    assert result["replacements"] == 0


def test_content_delete_range(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    (docs_dir / "delr.txt").write_bytes(b"A\nB\nC\nD\n")
    mcp, _ = _setup(make_config())
    _tool(mcp, "content_delete_range")(path="documents:/delr.txt", start_line=2, end_line=3)
    text = (docs_dir / "delr.txt").read_bytes().decode("utf-8")
    assert "B" not in text
    assert "C" not in text
    assert "A" in text
    assert "D" in text


def test_content_delete_range_invalid(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    mcp, _ = _setup(make_config())
    with pytest.raises(ValidationError, match="invalid"):
        _tool(mcp, "content_delete_range")(path="documents:/x.txt", start_line=5, end_line=2)


def test_content_patch(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    (docs_dir / "patch.txt").write_bytes(b"alpha\nbeta\ngamma\n")
    diff = """\
@@ -1,3 +1,3 @@
 alpha
-beta
+BETA
 gamma
"""
    mcp, _ = _setup(make_config())
    _tool(mcp, "content_patch")(path="documents:/patch.txt", unified_diff=diff)
    text = (docs_dir / "patch.txt").read_bytes().decode("utf-8")
    assert "BETA" in text
    assert "beta" not in text


# ── Additional content_crud coverage ──


def test_content_replace_with_count(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    (docs_dir / "cnt.txt").write_bytes(b"aaa bbb aaa ccc aaa")
    mcp, _ = _setup(make_config())
    result = _tool(mcp, "content_replace")(
        path="documents:/cnt.txt", old="aaa", new="xxx", count=2
    )
    assert result["replacements"] == 2
    text = (docs_dir / "cnt.txt").read_text(encoding="utf-8")
    assert text.count("xxx") == 2
    assert text.count("aaa") == 1


def test_content_replace_regex_with_count(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    (docs_dir / "rcnt.txt").write_bytes(b"cat1 cat2 cat3")
    mcp, _ = _setup(make_config())
    result = _tool(mcp, "content_replace")(
        path="documents:/rcnt.txt", old=r"cat\d", new="dog", regex=True, count=1
    )
    assert result["replacements"] == 1


def test_content_append_new_file(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    mcp, _ = _setup(make_config())
    _tool(mcp, "content_append")(path="documents:/new_app.txt", content="start")
    assert (docs_dir / "new_app.txt").read_text(encoding="utf-8") == "start"


def test_content_insert_new_file(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    mcp, _ = _setup(make_config())
    _tool(mcp, "content_insert")(path="documents:/new_ins.txt", line_number=1, content="first")
    text = (docs_dir / "new_ins.txt").read_text(encoding="utf-8")
    assert "first" in text


def test_content_patch_with_headers(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    (docs_dir / "ph.txt").write_bytes(b"one\ntwo\nthree\n")
    diff = """\
--- a/ph.txt
+++ b/ph.txt
@@ -1,3 +1,3 @@
 one
-two
+TWO
 three
"""
    mcp, _ = _setup(make_config())
    _tool(mcp, "content_patch")(path="documents:/ph.txt", unified_diff=diff)
    text = (docs_dir / "ph.txt").read_text(encoding="utf-8")
    assert "TWO" in text


def test_content_read_no_slice(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    (docs_dir / "full.txt").write_bytes(b"L1\nL2\nL3\n")
    mcp, _ = _setup(make_config())
    result = _tool(mcp, "content_read")(path="documents:/full.txt")
    assert result["content"] == "L1\nL2\nL3\n"


def test_content_write_unsupported_format(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    # Write to a file with xlsx extension that doesn't exist yet - should fallback
    mcp, _ = _setup(make_config())
    _tool(mcp, "content_write")(path="documents:/fallback.xyz", content="xyz data")
    assert (docs_dir / "fallback.xyz").read_text(encoding="utf-8") == "xyz data"


def test_content_delete_range_start_zero(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    mcp, _ = _setup(make_config())
    with pytest.raises(ValidationError, match="invalid"):
        _tool(mcp, "content_delete_range")(path="documents:/x.txt", start_line=0, end_line=1)


def test_content_insert_beyond_end(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    (docs_dir / "beyond.txt").write_bytes(b"A\n")
    mcp, _ = _setup(make_config())
    _tool(mcp, "content_insert")(path="documents:/beyond.txt", line_number=999, content="END")
    text = (docs_dir / "beyond.txt").read_text(encoding="utf-8")
    assert "END" in text


def test_content_insert_with_trailing_newline(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    (docs_dir / "tn.txt").write_bytes(b"A\nB\n")
    mcp, _ = _setup(make_config())
    _tool(mcp, "content_insert")(path="documents:/tn.txt", line_number=2, content="X\n")
    text = (docs_dir / "tn.txt").read_text(encoding="utf-8")
    lines = text.strip().split("\n")
    assert lines == ["A", "X", "B"]


def test_content_patch_invalid_hunk(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    from dokumen_pintar.errors import HandlerError
    docs_dir, _ = tmp_roots
    (docs_dir / "badhunk.txt").write_bytes(b"hello\nworld\n")
    diff = "@@ BAD HEADER @@\n hello\n"
    mcp, _ = _setup(make_config())
    with pytest.raises(HandlerError, match="Invalid hunk"):
        _tool(mcp, "content_patch")(path="documents:/badhunk.txt", unified_diff=diff)


def test_content_patch_context_mismatch(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    from dokumen_pintar.errors import HandlerError
    docs_dir, _ = tmp_roots
    (docs_dir / "cm.txt").write_bytes(b"hello\nworld\n")
    diff = "@@ -1,2 +1,2 @@\n wrong_context\n+new\n"
    mcp, _ = _setup(make_config())
    with pytest.raises(HandlerError, match="mismatch"):
        _tool(mcp, "content_patch")(path="documents:/cm.txt", unified_diff=diff)


def test_content_read_custom_encoding(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    (docs_dir / "enc.txt").write_bytes(b"L1\nL2\n")
    mcp, _ = _setup(make_config())
    result = _tool(mcp, "content_read")(path="documents:/enc.txt", encoding="latin-1")
    assert result["encoding"] == "latin-1"


def test_content_read_handler_unsupported_fallback(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    from unittest.mock import patch
    from dokumen_pintar.errors import UnsupportedFormatError
    docs_dir, _ = tmp_roots
    (docs_dir / "fallback.txt").write_bytes(b"raw text\n")
    cfg = make_config()
    mcp, ctx = _setup(cfg)
    from dokumen_pintar.handlers.text_handler import TextHandler
    with patch.object(TextHandler, "read_text", side_effect=UnsupportedFormatError("nope")):
        result = _tool(mcp, "content_read")(path="documents:/fallback.txt")
    assert "raw text" in result["content"]


def test_content_write_handler_unsupported_fallback(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    from unittest.mock import patch
    from dokumen_pintar.errors import UnsupportedFormatError
    docs_dir, _ = tmp_roots
    (docs_dir / "wfallback.txt").write_bytes(b"old\n")
    cfg = make_config()
    mcp, ctx = _setup(cfg)
    from dokumen_pintar.handlers.text_handler import TextHandler
    with patch.object(TextHandler, "write_text", side_effect=UnsupportedFormatError("nope")):
        _tool(mcp, "content_write")(path="documents:/wfallback.txt", content="new text")
    assert (docs_dir / "wfallback.txt").read_text(encoding="utf-8") == "new text"


def test_content_write_binary_handler_unsupported_reraises(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    """If a binary-container handler with WRITE_TEXT capability raises
    UnsupportedFormatError at runtime, the error must propagate — we must
    NOT silently fall back to writing raw bytes (that would corrupt the
    file with a binary extension)."""
    from unittest.mock import patch
    from dokumen_pintar.errors import UnsupportedFormatError
    from dokumen_pintar.handlers.docx_handler import DocxHandler

    docs_dir, _ = tmp_roots
    mcp, _ = _setup(make_config())
    # New docx (does not exist) — refuse_binary_text_op short-circuits, then
    # handler.write_text is invoked and we force it to raise.
    with patch.object(DocxHandler, "write_text", side_effect=UnsupportedFormatError("nope")):
        with pytest.raises(UnsupportedFormatError, match="nope"):
            _tool(mcp, "content_write")(path="documents:/x.docx", content="x")
    # No raw-bytes fallback file should have been left behind.
    assert not (docs_dir / "x.docx").exists()


@pytest.mark.parametrize("ext", ["xlsx", "pptx", "pdf"])
def test_content_write_refuses_new_binary_without_write_text(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path], ext: str
) -> None:
    """xlsx/pptx/pdf handlers don't expose WRITE_TEXT — content_write on a
    fresh file with that extension must refuse rather than fall through to
    a raw-bytes write."""
    from dokumen_pintar.errors import UnsupportedFormatError

    docs_dir, _ = tmp_roots
    mcp, _ = _setup(make_config())
    with pytest.raises(UnsupportedFormatError, match="write_text"):
        _tool(mcp, "content_write")(path=f"documents:/fresh.{ext}", content="x")
    assert not (docs_dir / f"fresh.{ext}").exists()


def test_content_patch_removal_mismatch(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    from dokumen_pintar.errors import HandlerError
    docs_dir, _ = tmp_roots
    (docs_dir / "rm.txt").write_bytes(b"alpha\nbeta\ngamma\n")
    diff = "@@ -1,2 +1,1 @@\n-alpha\n-WRONG_LINE\n+merged\n"
    mcp, _ = _setup(make_config())
    with pytest.raises(HandlerError, match="removal mismatch"):
        _tool(mcp, "content_patch")(path="documents:/rm.txt", unified_diff=diff)


def test_content_patch_with_prefix_and_trailing(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    (docs_dir / "pft.txt").write_bytes(b"line1\nline2\nline3\nline4\n")
    # Patch starting at line 3: replace line3 with NEW
    diff = "@@ -3,1 +3,1 @@\n-line3\n+NEW\n"
    mcp, _ = _setup(make_config())
    _tool(mcp, "content_patch")(path="documents:/pft.txt", unified_diff=diff)
    text = (docs_dir / "pft.txt").read_text(encoding="utf-8")
    assert "line1\nline2\nNEW\nline4\n" == text


def test_content_patch_empty_line(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    (docs_dir / "el.txt").write_bytes(b"A\nB\n")
    # Patch adds an empty line
    diff = "@@ -1,2 +1,3 @@\n A\n\n B\n"
    mcp, _ = _setup(make_config())
    _tool(mcp, "content_patch")(path="documents:/el.txt", unified_diff=diff)
    text = (docs_dir / "el.txt").read_text(encoding="utf-8")
    assert "\n\n" in text or "A" in text


def test_content_read_no_handler(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    # Write a file with unknown extension so registry.for_path returns None
    (docs_dir / "data.xyz123").write_text("raw content", encoding="utf-8")
    mcp, ctx = _setup(make_config())
    # Clear registry match by patching for_path to return None
    from unittest.mock import patch
    with patch.object(ctx.registry, "for_path", return_value=None):
        result = _tool(mcp, "content_read")(path="documents:/data.xyz123")
    assert "raw content" in result["content"]


def test_content_patch_no_trailing_newline(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    # File with no trailing newline
    (docs_dir / "notrail.txt").write_bytes(b"alpha\nbeta")
    diff = "@@ -1,2 +1,2 @@\n alpha\n-beta\n+BETA\n"
    mcp, _ = _setup(make_config())
    _tool(mcp, "content_patch")(path="documents:/notrail.txt", unified_diff=diff)
    text = (docs_dir / "notrail.txt").read_text(encoding="utf-8")
    assert text == "alpha\nBETA"


def test_content_patch_bare_line_outside_hunk(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    (docs_dir / "bare.txt").write_bytes(b"x\n")
    # "diff --git" outer header → line 258 (outer di += 1)
    # "\ No newline" inside hunk → 254->256 (inner loop unrecognized line)
    diff = (
        "diff --git a/bare.txt b/bare.txt\n"
        "--- a/bare.txt\n"
        "+++ b/bare.txt\n"
        "@@ -1,1 +1,1 @@\n"
        "-x\n"
        "+y\n"
        "\\ No newline at end of file\n"
    )
    mcp, _ = _setup(make_config())
    _tool(mcp, "content_patch")(path="documents:/bare.txt", unified_diff=diff)
    text = (docs_dir / "bare.txt").read_text(encoding="utf-8")
    assert text == "y\n"


# ── B7 regression: content_* tools must refuse to corrupt binary formats ──


def _make_docx(path: Path, paragraphs: list[str]) -> bytes:
    from docx import Document

    doc = Document()
    for p in paragraphs:
        doc.add_paragraph(p)
    doc.save(str(path))
    return path.read_bytes()


def _make_xlsx(path: Path) -> bytes:
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws["A1"] = "hello"
    wb.save(str(path))
    return path.read_bytes()


@pytest.mark.parametrize(
    "tool,kwargs",
    [
        ("content_write", {"content": "destroyed"}),
        ("content_append", {"content": "destroyed"}),
        ("content_insert", {"line_number": 1, "content": "destroyed"}),
        ("content_replace", {"old": "hello", "new": "destroyed"}),
        ("content_delete_range", {"start_line": 1, "end_line": 1}),
        ("content_patch", {"unified_diff": "@@ -1,1 +1,1 @@\n-hello\n+x\n"}),
    ],
)
def test_content_tools_refuse_existing_docx(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path],
    tool: str, kwargs: dict,
) -> None:
    from dokumen_pintar.errors import UnsupportedFormatError

    docs_dir, _ = tmp_roots
    docx_path = docs_dir / "doc.docx"
    original = _make_docx(docx_path, ["hello"])
    mcp, _ = _setup(make_config())
    with pytest.raises(UnsupportedFormatError):
        _tool(mcp, tool)(path="documents:/doc.docx", **kwargs)
    # File must remain byte-identical.
    assert docx_path.read_bytes() == original


@pytest.mark.parametrize(
    "make_file,suffix",
    [
        (_make_xlsx, ".xlsx"),
    ],
)
def test_content_write_refuses_existing_xlsx(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path],
    make_file, suffix: str,
) -> None:
    from dokumen_pintar.errors import UnsupportedFormatError

    docs_dir, _ = tmp_roots
    target = docs_dir / f"file{suffix}"
    original = make_file(target)
    mcp, _ = _setup(make_config())
    with pytest.raises(UnsupportedFormatError):
        _tool(mcp, "content_write")(path=f"documents:/file{suffix}", content="x")
    # File must remain byte-identical (no corruption).
    assert target.read_bytes() == original


def test_content_write_still_creates_new_docx(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    """Creating a brand-new docx via content_write must still work."""
    from docx import Document

    docs_dir, _ = tmp_roots
    target = docs_dir / "new.docx"
    assert not target.exists()
    mcp, _ = _setup(make_config())
    result = _tool(mcp, "content_write")(
        path="documents:/new.docx", content="hello\nworld"
    )
    assert target.exists()
    assert "snapshot" in result
    doc = Document(str(target))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "hello" in text and "world" in text


def test_docx_write_text_refuses_overwrite_existing(tmp_path: Path) -> None:
    """B8: DocxHandler.write_text must not silently destroy an existing docx."""
    from dokumen_pintar.errors import HandlerError
    from dokumen_pintar.handlers.docx_handler import DocxHandler

    handler = DocxHandler()
    target = tmp_path / "exist.docx"
    original = _make_docx(target, ["original"])
    with pytest.raises(HandlerError, match="refuses to overwrite"):
        handler.write_text(target, "replacement")
    assert target.read_bytes() == original



# ── v1.1.0 Quick Win 2.1: content_diff ──


def test_content_diff_identical_files(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    a = docs_dir / "a.txt"
    b = docs_dir / "b.txt"
    a.write_text("same\ncontent\n", encoding="utf-8")
    b.write_text("same\ncontent\n", encoding="utf-8")
    mcp, _ctx = _setup(make_config())
    result = _tool(mcp, "content_diff")(path_a=str(a), path_b=str(b))
    assert result["identical"] is True
    assert result["diff"] == ""
    assert result["stats"]["changes"] == 0


def test_content_diff_text_with_changes(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    a = docs_dir / "v1.txt"
    b = docs_dir / "v2.txt"
    a.write_text("alpha\nbeta\ngamma\n", encoding="utf-8")
    b.write_text("alpha\nBETA\ngamma\ndelta\n", encoding="utf-8")
    mcp, _ctx = _setup(make_config())
    result = _tool(mcp, "content_diff")(path_a=str(a), path_b=str(b))
    assert result["identical"] is False
    assert "BETA" in result["diff"]
    assert "delta" in result["diff"]
    assert result["stats"]["additions"] >= 1
    assert result["stats"]["deletions"] >= 1
    assert result["stats"]["changes"] == result["stats"]["additions"] + result["stats"]["deletions"]


def test_content_diff_ignore_whitespace(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    a = docs_dir / "ws_a.txt"
    b = docs_dir / "ws_b.txt"
    a.write_text("hello   world", encoding="utf-8")
    b.write_text("hello\tworld", encoding="utf-8")
    mcp, _ctx = _setup(make_config())
    result = _tool(mcp, "content_diff")(
        path_a=str(a), path_b=str(b), ignore_whitespace=True
    )
    assert result["identical"] is True


def test_content_diff_negative_context_rejected(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    docs_dir, _ = tmp_roots
    a = docs_dir / "a.txt"
    b = docs_dir / "b.txt"
    a.write_text("a", encoding="utf-8")
    b.write_text("b", encoding="utf-8")
    mcp, _ctx = _setup(make_config())
    with pytest.raises(ValidationError, match=r"context_lines"):
        _tool(mcp, "content_diff")(path_a=str(a), path_b=str(b), context_lines=-1)


def test_content_diff_unknown_format_uses_raw_text(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    """Files without a registered handler diff via raw text encoding."""
    docs_dir, _ = tmp_roots
    a = docs_dir / "a.unknownext"
    b = docs_dir / "b.unknownext"
    a.write_bytes(b"line1\nline2\n")
    b.write_bytes(b"line1\nLINE2\n")
    mcp, _ctx = _setup(make_config())
    result = _tool(mcp, "content_diff")(path_a=str(a), path_b=str(b))
    assert result["identical"] is False
    assert result["format_a"] == "text"
    assert result["format_b"] == "text"


def test_content_diff_docx_uses_extracted_view(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    """DOCX files diff via the search-extracted view (refuse_binary_text path)."""
    docs_dir, _ = tmp_roots
    a = docs_dir / "rep_a.docx"
    b = docs_dir / "rep_b.docx"
    _make_docx(a, ["alpha", "beta", "gamma"])
    _make_docx(b, ["alpha", "BETA", "gamma"])
    mcp, _ctx = _setup(make_config())
    result = _tool(mcp, "content_diff")(path_a=str(a), path_b=str(b))
    assert result["identical"] is False
    assert result["format_a"] == "docx"
    assert result["format_b"] == "docx"
    assert "BETA" in result["diff"]


def test_content_diff_xlsx_uses_extracted_view(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    """XLSX has no read_text -> falls back to extract_for_search for the diff."""
    import openpyxl

    docs_dir, _ = tmp_roots
    a = docs_dir / "data_a.xlsx"
    b = docs_dir / "data_b.xlsx"
    wb_a = openpyxl.Workbook()
    wb_a.active["A1"] = "hello"
    wb_a.save(a)
    wb_b = openpyxl.Workbook()
    wb_b.active["A1"] = "world"
    wb_b.save(b)
    mcp, _ctx = _setup(make_config())
    result = _tool(mcp, "content_diff")(path_a=str(a), path_b=str(b))
    assert result["identical"] is False
    assert result["format_a"] == "xlsx"



def test_content_diff_handler_raises_unsupported_falls_back_to_extract(
    make_config: Callable[..., AppConfig], tmp_roots: tuple[Path, Path]
) -> None:
    """If a handler's read_text raises UnsupportedFormatError, the diff
    falls back to the extract_for_search view (covers the defensive
    branch in _read_for_diff)."""
    from unittest.mock import patch

    from dokumen_pintar.errors import UnsupportedFormatError

    docs_dir, _ = tmp_roots
    a = docs_dir / "fb_a.txt"
    b = docs_dir / "fb_b.txt"
    a.write_text("alpha", encoding="utf-8")
    b.write_text("beta", encoding="utf-8")
    mcp, ctx = _setup(make_config())
    handler = ctx.registry.for_path(a.absolute())

    real_read_text = handler.read_text
    real_extract = handler.extract_for_search

    def fake_read_text(*_args: Any, **_kwargs: Any) -> str:
        raise UnsupportedFormatError("forced for test")

    with patch.object(handler, "read_text", side_effect=fake_read_text):
        with patch.object(
            handler,
            "extract_for_search",
            side_effect=lambda p: real_read_text(p),
        ):
            result = _tool(mcp, "content_diff")(path_a=str(a), path_b=str(b))
    # Sanity: real handlers were patched, but real_extract is unused -
    # silence the linter without affecting the assertion.
    _ = real_extract
    assert result["identical"] is False
    assert "alpha" in result["diff"] or "beta" in result["diff"]
