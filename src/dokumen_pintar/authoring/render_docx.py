"""Render a :class:`DocumentSpec` to a `.docx` file via python-docx."""

from __future__ import annotations

from pathlib import Path
from typing import Any, Callable

from docx import Document
from docx.shared import Cm, Pt, RGBColor

from dokumen_pintar.errors import HandlerError

from .spec import DocumentSpec


_HEX_RX = "0123456789abcdefABCDEF"


def _parse_hex_color(value: str) -> RGBColor | None:
    v = value.strip().lstrip("#")
    if len(v) != 6 or any(c not in _HEX_RX for c in v):
        return None
    return RGBColor(int(v[0:2], 16), int(v[2:4], 16), int(v[4:6], 16))


def _add_runs(paragraph: Any, runs: list[dict[str, Any]]) -> None:
    for r in runs:
        run = paragraph.add_run(r.get("text", ""))
        if r.get("bold"):
            run.bold = True
        if r.get("italic"):
            run.italic = True
        if r.get("underline"):
            run.underline = True
        if r.get("code"):
            try:
                run.font.name = "Consolas"
            except Exception:  # pragma: no cover — defensive
                pass
        if "font_size" in r:
            run.font.size = Pt(float(r["font_size"]))
        if "color" in r:
            color = _parse_hex_color(str(r["color"]))
            if color is not None:
                run.font.color.rgb = color


def _resolve_image_path(path_str: str, path_resolver: Callable[[str], Path] | None) -> Path:
    if path_resolver is not None:
        return path_resolver(path_str)
    return Path(path_str)


def _render_block(
    doc: Any,
    block: dict[str, Any],
    *,
    path_resolver: Callable[[str], Path] | None,
) -> None:
    btype = block["type"]
    if btype == "heading":
        doc.add_heading(block["text"], level=int(block["level"]))
        return
    if btype == "paragraph":
        para = doc.add_paragraph()
        _add_runs(para, block.get("runs", []))
        return
    if btype == "list":
        style = "List Number" if block.get("ordered") else "List Bullet"
        for item in block["items"]:
            doc.add_paragraph(item, style=style)
        return
    if btype == "table":
        rows = block["rows"]
        header = block.get("header")
        n_cols = max(
            (len(header) if header else 0),
            *(len(r) for r in rows) if rows else (0,),
        )
        if n_cols == 0:
            return
        n_rows = (1 if header else 0) + len(rows)
        if n_rows == 0:  # pragma: no cover — guarded by the n_cols check
            # above (an empty header + empty rows hits n_cols==0 first).
            return
        table = doc.add_table(rows=n_rows, cols=n_cols)
        table.style = "Table Grid"
        row_offset = 0
        if header:
            hdr_cells = table.rows[0].cells
            for ci, cell_text in enumerate(header):
                hdr_cells[ci].text = cell_text
            row_offset = 1
        for ri, row in enumerate(rows):
            cells = table.rows[ri + row_offset].cells
            for ci, val in enumerate(row):
                cells[ci].text = val
        return
    if btype == "image":
        img_path = _resolve_image_path(block["path"], path_resolver)
        if not img_path.exists():
            raise HandlerError(f"image not found: {img_path}")
        kwargs: dict[str, Any] = {}
        if "width_cm" in block:
            kwargs["width"] = Cm(float(block["width_cm"]))
        doc.add_picture(str(img_path), **kwargs)
        if "caption" in block:
            cap = doc.add_paragraph(block["caption"])
            try:
                cap.style = doc.styles["Caption"]
            except KeyError:  # pragma: no cover — fallback
                pass
        return
    if btype == "page_break":
        doc.add_page_break()
        return
    if btype == "code":
        para = doc.add_paragraph()
        run = para.add_run(block["text"])
        try:
            run.font.name = "Consolas"
        except Exception:  # pragma: no cover
            pass
        run.font.size = Pt(10)
        return
    if btype == "math":
        # Plain rendering — DOCX OMML is non-trivial; ship as monospaced
        # inline placeholder. Future: integrate latex2omml.
        para = doc.add_paragraph()
        run = para.add_run(f"[math] {block['latex']}")
        run.italic = True
        return
    if btype == "hr":
        # Approximate horizontal rule with an empty paragraph + bottom border
        # would require deep OXML; use a centered glyph row as a pragmatic
        # fallback that survives in every viewer.
        doc.add_paragraph("─" * 40)
        return
    if btype == "blockquote":
        para = doc.add_paragraph(block["text"])
        try:
            para.style = doc.styles["Quote"]
        except KeyError:  # pragma: no cover
            pass
        return
    raise HandlerError(f"unsupported block type for docx: {btype!r}")


def _apply_meta(doc: Any, meta: dict[str, Any]) -> None:
    if not meta:
        return
    cp = doc.core_properties
    for key in ("title", "author", "subject", "keywords"):
        if key in meta and meta[key] is not None:
            try:
                setattr(cp, key, meta[key])
            except (AttributeError, TypeError, ValueError):  # pragma: no cover
                pass


def render_docx(
    spec: DocumentSpec,
    out_path: Path,
    *,
    path_resolver: Callable[[str], Path] | None = None,
    template: Path | None = None,
) -> None:
    """Render ``spec`` to ``out_path``.

    Caller is responsible for snapshots/locks.

    ``path_resolver`` lets the caller resolve image URIs (e.g.
    ``kp:/img.png``) through PathGuard before they are fed to python-docx.

    ``template`` (optional) is an existing DOCX whose styles, page setup,
    headers, and footers should be inherited. The template's body is
    preserved as-is; the rendered blocks are appended after it. Useful
    for university or corporate templates with cover pages and standard
    section properties.
    """
    if template is not None:
        try:
            doc = Document(str(template))
        except Exception as exc:  # noqa: BLE001 - python-docx surfaces several types
            raise HandlerError(f"failed to load template: {template} ({exc})") from exc
    else:
        doc = Document()
    _apply_meta(doc, spec.meta)
    for block in spec.blocks:
        _render_block(doc, block, path_resolver=path_resolver)
    try:
        doc.save(str(out_path))
    except Exception as exc:  # noqa: BLE001
        raise HandlerError(f"failed to write docx: {out_path} ({exc})") from exc
