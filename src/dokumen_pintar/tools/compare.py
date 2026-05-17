"""Document comparison with visual diff export.

Generates a comparison DOCX from two source documents in one of three
styles:

- ``track_changes``: Word-style insertions/deletions inline. Rendered as
  paragraphs with green ``[+ inserted +]`` / red ``[- deleted -]``
  markers. (Real Word ``w:ins`` / ``w:del`` track-changes XML is not
  emitted because round-tripping it through python-docx is fragile;
  the visual marker form survives every viewer.)
- ``side_by_side``: a two-column table, A on the left, B on the right.
- ``diff_doc``: colored unified diff (red removals, green additions,
  black context).

The compared text comes from each handler's ``read_text`` view, so DOCX,
PDF, and other formats are all valid inputs. The output is always a DOCX.
"""

from __future__ import annotations

import difflib
from pathlib import Path
from typing import Any

from mcp.server.fastmcp import FastMCP

from ..context import AppContext
from ..errors import HandlerError, UnsupportedFormatError, ValidationError
from ..utils.locks import file_lock
from ._common import resolve_for_read, resolve_for_write, summarize_resolved


_VALID_STYLES = frozenset({"track_changes", "side_by_side", "diff_doc"})


def _read_text_for_compare(ctx: AppContext, path: Path) -> str:
    """Use the registered handler's read_text for comparison."""
    handler = ctx.registry.for_path(path)
    if handler is None:
        raise UnsupportedFormatError(f"no handler for {path.suffix!r}; cannot read for comparison")
    try:
        return handler.read_text(path)
    except (
        UnsupportedFormatError
    ):  # pragma: no cover - no v1.1.0 handler raises here, but kept as a safety net
        # Fall back to extracted view (e.g. xlsx/pptx).
        return handler.extract_for_search(path)


def _make_track_changes_doc(text_a: str, text_b: str, dst: Path, label_a: str, label_b: str) -> int:
    """Build a track-changes-style DOCX. Returns paragraph count."""
    from docx import Document
    from docx.shared import RGBColor

    doc = Document()
    doc.add_heading(f"Compare: {label_a} → {label_b}", level=1)
    matcher = difflib.SequenceMatcher(a=text_a.splitlines(), b=text_b.splitlines())
    n_paragraphs = 1  # heading
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == "equal":
            for line in text_a.splitlines()[i1:i2]:
                doc.add_paragraph(line)
                n_paragraphs += 1
        elif tag == "delete":
            for line in text_a.splitlines()[i1:i2]:
                p = doc.add_paragraph()
                run = p.add_run(f"[- {line} -]")
                run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
                run.bold = True
                n_paragraphs += 1
        elif tag == "insert":
            for line in text_b.splitlines()[j1:j2]:
                p = doc.add_paragraph()
                run = p.add_run(f"[+ {line} +]")
                run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
                run.bold = True
                n_paragraphs += 1
        else:  # tag == "replace"
            for line in text_a.splitlines()[i1:i2]:
                p = doc.add_paragraph()
                run = p.add_run(f"[- {line} -]")
                run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
                n_paragraphs += 1
            for line in text_b.splitlines()[j1:j2]:
                p = doc.add_paragraph()
                run = p.add_run(f"[+ {line} +]")
                run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
                n_paragraphs += 1
    doc.save(str(dst))
    return n_paragraphs


def _make_side_by_side_doc(text_a: str, text_b: str, dst: Path, label_a: str, label_b: str) -> int:
    """Build a two-column comparison table."""
    from docx import Document

    doc = Document()
    doc.add_heading(f"Side-by-side: {label_a} ↔ {label_b}", level=1)
    lines_a = text_a.splitlines()
    lines_b = text_b.splitlines()
    rows = max(len(lines_a), len(lines_b)) + 1  # +1 for header row
    table = doc.add_table(rows=rows, cols=2)
    table.style = "Light Grid"
    table.cell(0, 0).text = label_a
    table.cell(0, 1).text = label_b
    for i in range(rows - 1):
        table.cell(i + 1, 0).text = lines_a[i] if i < len(lines_a) else ""
        table.cell(i + 1, 1).text = lines_b[i] if i < len(lines_b) else ""
    doc.save(str(dst))
    return rows


def _make_diff_doc(text_a: str, text_b: str, dst: Path, label_a: str, label_b: str) -> int:
    """Build a unified diff DOCX with colored runs."""
    from docx import Document
    from docx.shared import RGBColor

    doc = Document()
    doc.add_heading(f"Diff: {label_a} → {label_b}", level=1)
    lines_a = text_a.splitlines(keepends=True)
    lines_b = text_b.splitlines(keepends=True)
    diff_lines = list(
        difflib.unified_diff(
            lines_a,
            lines_b,
            fromfile=label_a,
            tofile=label_b,
            n=3,
        )
    )
    n_paragraphs = 1
    for line in diff_lines:
        line_stripped = line.rstrip("\r\n")
        p = doc.add_paragraph()
        run = p.add_run(line_stripped if line_stripped else " ")
        if line.startswith("+++") or line.startswith("---"):
            run.bold = True
        elif line.startswith("@@"):
            run.italic = True
            run.font.color.rgb = RGBColor(0x00, 0x66, 0xCC)
        elif line.startswith("+"):
            run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
        elif line.startswith("-"):
            run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
        n_paragraphs += 1
    doc.save(str(dst))
    return n_paragraphs


def register(mcp: FastMCP, ctx: AppContext) -> None:
    """Register the document_compare tool."""

    @mcp.tool(
        name="document_compare",
        description=(
            "Generate a comparison DOCX from two source documents. "
            "`style` is one of: 'track_changes' (default; insertion/deletion "
            "markers inline), 'side_by_side' (two-column table), 'diff_doc' "
            "(colored unified diff). Source format detected automatically; "
            "outputs always DOCX. Refuses to overwrite unless "
            "`overwrite=True`. Snapshots pre+post."
        ),
    )
    def document_compare(
        src_a: str,
        src_b: str,
        dst: str,
        style: str = "track_changes",
        overwrite: bool = False,
    ) -> dict[str, Any]:
        if style not in _VALID_STYLES:
            raise ValidationError(f"style must be one of {sorted(_VALID_STYLES)} (got {style!r})")
        resolved_a = resolve_for_read(ctx, src_a)
        resolved_b = resolve_for_read(ctx, src_b)
        resolved_dst = resolve_for_write(ctx, dst)
        if resolved_dst.absolute.suffix.lower() != ".docx":
            raise UnsupportedFormatError(
                f"document_compare: dst must be .docx, got {resolved_dst.absolute.suffix!r}"
            )
        if resolved_dst.absolute.exists() and not overwrite:
            raise ValidationError(
                f"destination exists: {resolved_dst.absolute}; pass overwrite=True"
            )

        text_a = _read_text_for_compare(ctx, resolved_a.absolute)
        text_b = _read_text_for_compare(ctx, resolved_b.absolute)
        label_a = resolved_a.absolute.name
        label_b = resolved_b.absolute.name

        rel = resolved_dst.rel_to_root.as_posix()
        root_name = resolved_dst.root.name
        with file_lock(resolved_dst.absolute):
            if resolved_dst.absolute.exists():
                ctx.versions.snapshot(
                    root_name=root_name,
                    rel_path=rel,
                    source=resolved_dst.absolute,
                    action="document_compare_pre",
                )
            try:
                if style == "track_changes":
                    n_paragraphs = _make_track_changes_doc(
                        text_a, text_b, resolved_dst.absolute, label_a, label_b
                    )
                elif style == "side_by_side":
                    n_paragraphs = _make_side_by_side_doc(
                        text_a, text_b, resolved_dst.absolute, label_a, label_b
                    )
                else:  # diff_doc
                    n_paragraphs = _make_diff_doc(
                        text_a, text_b, resolved_dst.absolute, label_a, label_b
                    )
            except Exception as exc:  # noqa: BLE001 - python-docx serialiser
                raise HandlerError(
                    f"failed to write comparison docx: {resolved_dst.absolute} ({exc})"
                ) from exc
            snap = ctx.versions.snapshot(
                root_name=root_name,
                rel_path=rel,
                source=resolved_dst.absolute,
                action="document_compare_post",
            )
        # Crude stat: line counts in each.
        ctx.audit.log(
            "document_compare",
            src_a=str(resolved_a.absolute),
            src_b=str(resolved_b.absolute),
            dst=str(resolved_dst.absolute),
            style=style,
            paragraphs=n_paragraphs,
        )
        return {
            "src_a": summarize_resolved(resolved_a),
            "src_b": summarize_resolved(resolved_b),
            "dst": summarize_resolved(resolved_dst),
            "style": style,
            "paragraphs": n_paragraphs,
            "lines_a": len(text_a.splitlines()),
            "lines_b": len(text_b.splitlines()),
            "snapshot": snap,
        }
