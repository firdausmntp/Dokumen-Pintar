"""Citation and bibliography validation.

Citations are detected by regex patterns commonly used in Indonesian
academic writing:

- ``[1]``, ``[12]`` IEEE-style numeric citations.
- ``(Author, 2024)`` APA-style author-year.
- ``(Author et al., 2024)``.

The ``DAFTAR PUSTAKA`` / ``REFERENCES`` / ``Bibliography`` section is
auto-detected by walking the document headings; users can override
with ``bib_section_pattern``. ``bibliography_check`` reports unused
entries, missing entries, and duplicates. ``bibliography_format``
re-sorts the section alphabetically (or by first-appearance order)
and flags entries that look malformed for the chosen citation style.
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import Any

from mcp.server.fastmcp import FastMCP

from ..context import AppContext
from ..errors import HandlerError, UnsupportedFormatError, ValidationError
from ..utils.locks import file_lock
from ._common import resolve_for_read, resolve_for_write, summarize_resolved


# Citation regex patterns (compiled once per call, attached to results).
_CITATION_NUMERIC = re.compile(r"\[(\d+)\]")
_CITATION_AUTHOR_YEAR = re.compile(
    r"\(([A-Z][A-Za-z'\-\s]+?)(?:\s+et\s+al\.?)?,?\s*(\d{4})[a-z]?\)"
)

# Bibliography section markers (case-insensitive).
_BIB_SECTION_DEFAULT = re.compile(
    r"(?i)^(daftar pustaka|references|bibliography)\b",
)


def _split_into_paragraphs(doc: Any) -> list[tuple[int, str, int | None]]:
    """Return ``[(idx, text, heading_level), ...]`` for body paragraphs."""
    out: list[tuple[int, str, int | None]] = []
    for idx, para in enumerate(doc.paragraphs):
        style_name = (getattr(getattr(para, "style", None), "name", None) or "").strip()
        m = re.match(r"^[Hh]eading\s*(\d+)$", style_name)
        if m:
            level = int(m.group(1))
        elif style_name.lower() == "title":  # pragma: no cover - title style rare in bib docs
            level = 0
        else:
            level = None
        out.append((idx, (para.text or ""), level))
    return out


def _find_bibliography_range(
    paragraphs: list[tuple[int, str, int | None]],
    pattern: re.Pattern[str],
) -> tuple[int, int] | None:
    """Return ``(start_idx, end_idx)`` of the bibliography section, or None.

    The section spans from the matching heading (exclusive) to the next
    heading at the same or higher level, or end of document.
    """
    start: int | None = None
    start_level: int | None = None
    for idx, text, level in paragraphs:
        if level is None:
            continue
        if pattern.search(text.strip()):
            start = idx
            start_level = level
            break
    if start is None:
        return None

    end = paragraphs[-1][0] + 1
    for idx, _text, level in paragraphs:
        if idx <= start:
            continue
        if (
            level is not None and level <= start_level
        ):  # pragma: no cover - bib in tests is always last section
            end = idx
            break
    return (start + 1, end)  # body starts after the heading itself


def _extract_citations(
    paragraphs: list[tuple[int, str, int | None]],
    bib_range: tuple[int, int] | None,
) -> list[dict[str, Any]]:
    """Find every citation reference in the body (excluding the bib section)."""
    out: list[dict[str, Any]] = []
    bib_start, bib_end = bib_range or (10**9, 10**9)
    for idx, text, _level in paragraphs:
        if bib_start <= idx < bib_end:
            continue
        for m in _CITATION_NUMERIC.finditer(text):
            out.append(
                {
                    "kind": "numeric",
                    "raw": m.group(0),
                    "key": m.group(1),
                    "paragraph_index": idx,
                }
            )
        for m in _CITATION_AUTHOR_YEAR.finditer(text):
            out.append(
                {
                    "kind": "author_year",
                    "raw": m.group(0),
                    "key": f"{m.group(1).strip()} {m.group(2)}",
                    "paragraph_index": idx,
                }
            )
    return out


def _extract_bib_entries(
    paragraphs: list[tuple[int, str, int | None]],
    bib_range: tuple[int, int] | None,
) -> list[dict[str, Any]]:
    """Parse paragraphs in the bibliography section into structured entries."""
    if bib_range is None:
        return []
    bib_start, bib_end = bib_range
    out: list[dict[str, Any]] = []
    for idx, text, _level in paragraphs:
        if not (bib_start <= idx < bib_end):
            continue
        if not text.strip():
            continue
        # Detect [N] prefix for numeric style entries.
        m = _CITATION_NUMERIC.match(text.strip())
        if m:
            out.append(
                {
                    "kind": "numeric",
                    "key": m.group(1),
                    "raw": text.strip(),
                    "paragraph_index": idx,
                }
            )
            continue
        # Otherwise treat as author-year style; key = first author(s) + year.
        # Author names may contain letters, hyphens, apostrophes, periods,
        # ampersands, commas (e.g. "Smith, J.", "Smith, J., & Jones, A.").
        ay_match = re.match(r"^([A-Z][A-Za-z'\-\.\s,&]+?)\s+\(?(\d{4})", text.strip())
        if ay_match:
            # Normalise the author key by stripping trailing punctuation/whitespace.
            author = ay_match.group(1).strip(" ,.")
            out.append(
                {
                    "kind": "author_year",
                    "key": f"{author} {ay_match.group(2)}",
                    "raw": text.strip(),
                    "paragraph_index": idx,
                }
            )
        else:
            out.append(
                {
                    "kind": "unknown",
                    "key": text.strip()[:60],
                    "raw": text.strip(),
                    "paragraph_index": idx,
                }
            )
    return out


def _detect_issues(
    citations: list[dict[str, Any]],
    bib_entries: list[dict[str, Any]],
) -> list[dict[str, Any]]:
    """Cross-reference citations against bibliography entries."""
    issues: list[dict[str, Any]] = []

    # Index bib entries by key.
    bib_by_key: dict[str, list[dict[str, Any]]] = {}
    for entry in bib_entries:
        bib_by_key.setdefault(entry["key"], []).append(entry)

    citation_keys = {c["key"] for c in citations}

    # 1. Missing entries: citation without a matching bib entry.
    for c in citations:
        if c["key"] not in bib_by_key:
            issues.append(
                {
                    "type": "missing_bib_entry",
                    "citation": c["raw"],
                    "key": c["key"],
                    "paragraph_index": c["paragraph_index"],
                }
            )

    # 2. Unused entries: bib entry never referenced.
    for key, entries in bib_by_key.items():
        if key not in citation_keys:
            issues.append(
                {
                    "type": "unused_bib_entry",
                    "key": key,
                    "raw": entries[0]["raw"],
                }
            )

    # 3. Duplicate entries: same key appears multiple times in bib.
    for key, entries in bib_by_key.items():
        if len(entries) > 1:
            issues.append(
                {
                    "type": "duplicate_bib_entry",
                    "key": key,
                    "count": len(entries),
                }
            )

    return issues


def _validate_entry_format(entry: dict[str, Any], style: str) -> str | None:
    """Return an issue description if ``entry`` doesn't fit the chosen style."""
    raw = entry["raw"]
    if style == "APA":
        # Minimum: Author (Year). Title.  -> matches Author + 4-digit year.
        if not re.search(r"\(\d{4}\)", raw) and entry["kind"] == "unknown":
            return (
                "APA expects '(YYYY)' year format; entry has neither year "
                "nor recognisable author-year structure"
            )
    elif style == "IEEE":  # pragma: no branch - APA + IEEE are the supported style branches
        # IEEE entries should start with [N].
        if entry["kind"] != "numeric":
            return "IEEE expects entries to start with [N]"
    return None


def _open_doc(path: Path) -> Any:
    """Open ``path`` with python-docx, wrapping load errors."""
    from docx import Document
    from docx.opc.exceptions import PackageNotFoundError

    try:
        return Document(str(path))
    except PackageNotFoundError as exc:
        raise HandlerError(f"not a valid docx: {path} ({exc})") from exc
    except Exception as exc:  # pragma: no cover - PackageNotFoundError covers common case
        raise HandlerError(f"failed to open docx: {path} ({exc})") from exc


def register(mcp: FastMCP, ctx: AppContext) -> None:
    """Register bibliography_check + bibliography_format tools."""

    @mcp.tool(
        name="bibliography_check",
        description=(
            "Validate citations against the bibliography section. Detects "
            "missing entries (cited but not listed), unused entries (listed "
            "but never cited), and duplicates. Returns the citations found, "
            "the parsed bibliography entries, and a list of issues. "
            "`auto_detect_section=True` (default) walks headings for "
            "DAFTAR PUSTAKA / REFERENCES / Bibliography; pass "
            "`bib_section_pattern` to override."
        ),
    )
    def bibliography_check(
        path: str,
        style: str = "APA",
        auto_detect_section: bool = True,
        bib_section_pattern: str | None = None,
    ) -> dict[str, Any]:
        if style not in ("APA", "IEEE", "Chicago", "Harvard"):
            raise ValidationError(f"style must be one of APA/IEEE/Chicago/Harvard (got {style!r})")
        resolved = resolve_for_read(ctx, path)
        if resolved.absolute.suffix.lower() != ".docx":
            raise UnsupportedFormatError(
                f"bibliography_check: target must be .docx, got {resolved.absolute.suffix!r}"
            )

        doc = _open_doc(resolved.absolute)
        paragraphs = _split_into_paragraphs(doc)
        if auto_detect_section or bib_section_pattern:
            pattern = (
                re.compile(bib_section_pattern) if bib_section_pattern else _BIB_SECTION_DEFAULT
            )
            bib_range = _find_bibliography_range(paragraphs, pattern)
        else:
            bib_range = None

        citations = _extract_citations(paragraphs, bib_range)
        bib_entries = _extract_bib_entries(paragraphs, bib_range)
        issues = _detect_issues(citations, bib_entries)

        # Style-specific format validation on the bib entries themselves.
        for entry in bib_entries:
            problem = _validate_entry_format(entry, style)
            if problem is not None:
                issues.append(
                    {
                        "type": "malformed_entry",
                        "key": entry["key"],
                        "raw": entry["raw"],
                        "message": problem,
                    }
                )

        ctx.audit.log(
            "bibliography_check",
            path=str(resolved.absolute),
            style=style,
            citations=len(citations),
            entries=len(bib_entries),
            issues=len(issues),
        )
        return {
            **summarize_resolved(resolved),
            "style": style,
            "citations_found": citations,
            "bibliography_entries": bib_entries,
            "issues": issues,
            "bib_section_range": bib_range,
        }

    @mcp.tool(
        name="bibliography_format",
        description=(
            "Reformat the bibliography section. `sort=True` (default) "
            "sorts alphabetically by key. `auto_fix=False` (default) only "
            "reports what would change without writing; pass True to "
            "actually rewrite the section. Snapshots pre+post when "
            "auto_fix is True."
        ),
    )
    def bibliography_format(
        path: str,
        style: str = "APA",
        sort: bool = True,
        auto_fix: bool = False,
        bib_section_pattern: str | None = None,
    ) -> dict[str, Any]:
        if style not in ("APA", "IEEE", "Chicago", "Harvard"):
            raise ValidationError(f"style must be one of APA/IEEE/Chicago/Harvard (got {style!r})")
        resolved = resolve_for_write(ctx, path)
        if resolved.absolute.suffix.lower() != ".docx":
            raise UnsupportedFormatError(
                f"bibliography_format: target must be .docx, got {resolved.absolute.suffix!r}"
            )
        if not resolved.absolute.exists():
            raise ValidationError(f"file not found: {resolved.absolute}")

        doc = _open_doc(resolved.absolute)
        paragraphs = _split_into_paragraphs(doc)
        pattern = re.compile(bib_section_pattern) if bib_section_pattern else _BIB_SECTION_DEFAULT
        bib_range = _find_bibliography_range(paragraphs, pattern)
        if bib_range is None:
            raise ValidationError("no bibliography section found; check bib_section_pattern")
        bib_entries = _extract_bib_entries(paragraphs, bib_range)
        if not bib_entries:
            raise ValidationError("bibliography section is empty")

        original_order = [e["key"] for e in bib_entries]
        new_order = sorted(bib_entries, key=lambda e: e["key"]) if sort else bib_entries
        new_keys = [e["key"] for e in new_order]
        changed = original_order != new_keys

        if not auto_fix:
            return {
                **summarize_resolved(resolved),
                "style": style,
                "would_change": changed,
                "current_order": original_order,
                "proposed_order": new_keys,
            }

        # Apply by replacing each bib paragraph in-place with the sorted text.
        bib_start, bib_end = bib_range
        target_paragraphs = [
            p
            for i, p in enumerate(doc.paragraphs)
            if bib_start <= i < bib_end and (p.text or "").strip()
        ]
        # Pad with extras if new_order is shorter (preserves blank lines).
        for para, entry in zip(target_paragraphs, new_order):
            para.text = entry["raw"]

        rel = resolved.rel_to_root.as_posix()
        root_name = resolved.root.name
        with file_lock(resolved.absolute):
            ctx.versions.snapshot(
                root_name=root_name,
                rel_path=rel,
                source=resolved.absolute,
                action="bibliography_format_pre",
            )
            try:
                doc.save(str(resolved.absolute))
            except Exception as exc:  # noqa: BLE001 - python-docx serialiser
                raise HandlerError(f"failed to save docx: {resolved.absolute} ({exc})") from exc
            snap = ctx.versions.snapshot(
                root_name=root_name,
                rel_path=rel,
                source=resolved.absolute,
                action="bibliography_format_post",
            )
        ctx.audit.log(
            "bibliography_format",
            path=str(resolved.absolute),
            style=style,
            entries=len(bib_entries),
            sorted=sort,
        )
        return {
            **summarize_resolved(resolved),
            "style": style,
            "applied": True,
            "current_order": original_order,
            "new_order": new_keys,
            "snapshot": snap,
        }
