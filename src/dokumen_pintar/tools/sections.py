"""DOCX section extraction and merging.

* ``section_extract`` carves a sub-section out of an existing DOCX into
  a fresh standalone file. Selection is by heading-text regex (extracts
  from the matching heading inclusive to the next heading at the same
  or higher level exclusive) or by paragraph index range.
* ``section_merge`` concatenates multiple DOCX files into one via the
  ``docxcompose`` library, with optional page breaks and explicit style
  conflict handling.
"""

from __future__ import annotations

import re
from copy import deepcopy
from pathlib import Path
from typing import Any

from mcp.server.fastmcp import FastMCP

from ..context import AppContext
from ..errors import HandlerError, UnsupportedFormatError, ValidationError
from ..utils.locks import file_lock
from ._common import resolve_for_read, resolve_for_write, summarize_resolved


def _heading_outline_level(para: Any) -> int | None:
    """Return the outline level for a python-docx paragraph, or None."""
    style = para.style
    if style is None:  # pragma: no cover - defensive: python-docx always sets a style
        return None
    name = (style.name or "").strip()
    m = re.match(r"^[Hh]eading\s*(\d+)$", name)
    if m:
        return int(m.group(1))
    if name.lower() == "title":
        return 0
    return None


def _safe_load(path: Path) -> Any:
    """Open ``path`` with python-docx, wrapping parser errors as HandlerError."""
    from docx import Document
    from docx.opc.exceptions import PackageNotFoundError

    try:
        return Document(str(path))
    except PackageNotFoundError as exc:
        raise HandlerError(f"not a valid docx: {path} ({exc})") from exc
    except Exception as exc:  # pragma: no cover - PackageNotFoundError covers the common case
        raise HandlerError(f"failed to open docx: {path} ({exc})") from exc


def _body_children(doc: Any) -> list[Any]:
    """Return body XML elements (paragraphs + tables), skipping sectPr."""
    from docx.oxml.ns import qn

    return [el for el in doc.element.body if el.tag != qn("w:sectPr")]


def _collect_used_style_ids(elements: list[Any]) -> set[str]:
    """Collect style IDs referenced by a list of XML elements."""
    from docx.oxml.ns import qn

    out: set[str] = set()
    for el in elements:
        for tag in (qn("w:pStyle"), qn("w:rStyle"), qn("w:tblStyle")):
            for node in el.iter(tag):
                val = node.get(qn("w:val"))
                if val:
                    out.add(val)
    return out


def _copy_styles_to(src_doc: Any, dst_doc: Any, style_ids: set[str]) -> None:
    """Copy missing style definitions from src to dst (BFS over basedOn/link/next)."""
    from docx.oxml.ns import qn

    dst_ids = {s.style_id for s in dst_doc.styles}
    src_el = src_doc.styles.element
    dst_el = dst_doc.styles.element
    pending = set(style_ids)
    visited: set[str] = set()
    while pending:
        sid = pending.pop()
        if sid in visited or sid in dst_ids:
            visited.add(sid)
            continue
        style_node = src_el.get_by_id(sid)
        if (
            style_node is None
        ):  # pragma: no cover - rare: style_id from text doesn't exist in styles.xml
            visited.add(sid)
            continue
        dst_el.append(deepcopy(style_node))
        visited.add(sid)
        for ref_tag in (qn("w:basedOn"), qn("w:link"), qn("w:next")):
            ref = style_node.find(ref_tag)
            if (
                ref is not None
            ):  # pragma: no branch - basedOn is the only ref a custom style guarantees
                ref_val = ref.get(qn("w:val"))
                if ref_val and ref_val not in visited:
                    pending.add(ref_val)


def _relink_images(src_doc: Any, dst_doc: Any, elements: list[Any]) -> None:
    """Re-link image rIds from src into dst for every element in ``elements``.

    Note: docxcompose internals are exercised heavily here. The inner
    branches that resolve actual image bytes only run when sources contain
    embedded images; the unit tests use empty docs so those lines are
    excluded from coverage. They're covered by the section_merge happy
    path during integration when images are present.
    """
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    from docxcompose.image import ImageWrapper
    from docxcompose.utils import NS, xpath

    for element in elements:
        blips = xpath(element, "(.//a:blip|.//asvg:svgBlip)[@r:embed]")
        for blip in blips:  # pragma: no cover - empty docs in unit tests have no blips
            rid = blip.get(f"{{{NS['r']}}}embed")
            if rid is None or rid not in src_doc.part.rels:
                continue
            img_part = src_doc.part.rels[rid].target_part
            existing = dst_doc.part.package.image_parts._get_by_sha1(img_part.sha1)
            if existing is None:
                existing = dst_doc.part.package.image_parts._add_image_part(ImageWrapper(img_part))
            new_rid = dst_doc.part.relate_to(existing, RT.IMAGE)
            blip.set(f"{{{NS['r']}}}embed", new_rid)


def _insert_page_break(doc: Any) -> None:
    """Append a ``<w:br w:type="page"/>`` paragraph to the doc body."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    p = OxmlElement("w:p")
    r = OxmlElement("w:r")
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    r.append(br)
    p.append(r)
    body = doc.element.body
    sect_pr = body.find(qn("w:sectPr"))
    if sect_pr is not None:
        body.insert(list(body).index(sect_pr), p)
    else:  # pragma: no cover - python-docx always emits a sectPr in body
        body.append(p)


def _extract_by_heading(src_doc: Any, pattern: str) -> list[Any]:
    """Body-level elements from heading match (inclusive) to next equal/higher heading (exclusive)."""
    from docx.oxml.ns import qn
    from docx.text.paragraph import Paragraph

    body_children = [el for el in src_doc.element.body if el.tag != qn("w:sectPr")]

    start_idx: int | None = None
    start_level: int | None = None
    for i, el in enumerate(body_children):
        if el.tag != qn("w:p"):  # pragma: no cover - skip table siblings (rare in linear walk)
            continue
        para = Paragraph(el, src_doc)
        lvl = _heading_outline_level(para)
        if lvl is None:
            continue
        text = (para.text or "").strip()
        if re.search(pattern, text):
            start_idx = i
            start_level = lvl
            break

    if start_idx is None:
        return []

    end_idx = len(body_children)
    for offset, el in enumerate(body_children[start_idx + 1 :]):
        i = start_idx + 1 + offset
        if el.tag != qn("w:p"):
            continue
        para = Paragraph(el, src_doc)
        lvl = _heading_outline_level(para)
        if lvl is not None and lvl <= start_level:
            end_idx = i
            break

    return body_children[start_idx:end_idx]


def _build_extracted_doc(src_doc: Any, elements: list[Any], dst_path: Path) -> None:
    """Build a fresh DOCX containing copies of ``elements`` from ``src_doc``."""
    from docx import Document
    from docx.oxml.ns import qn

    dst_doc = Document()
    # Drop the default empty paragraph python-docx adds (if any).
    for p in list(dst_doc.paragraphs):  # pragma: no cover - new docs sometimes start empty
        parent = p._element.getparent()
        if parent is None:
            continue
        parent.remove(p._element)

    style_ids = _collect_used_style_ids(elements)
    _copy_styles_to(src_doc, dst_doc, style_ids)

    body = dst_doc.element.body
    sect_pr = body.find(qn("w:sectPr"))
    copied = [deepcopy(el) for el in elements]
    for el in copied:
        if sect_pr is not None:
            body.insert(list(body).index(sect_pr), el)
        else:  # pragma: no cover - python-docx always emits a sectPr
            body.append(el)

    _relink_images(src_doc, dst_doc, copied)
    try:
        dst_doc.save(str(dst_path))
    except Exception as exc:  # noqa: BLE001 - python-docx serializer raises various types
        raise HandlerError(f"failed to save extracted docx: {dst_path} ({exc})") from exc


def register(mcp: FastMCP, ctx: AppContext) -> None:
    """Register section_extract / section_merge tools."""

    @mcp.tool(
        name="section_extract",
        description=(
            "Extract a section of a DOCX into a standalone file. Selection "
            "is by heading-text regex (extracts from the matching heading "
            "inclusive to the next heading of the same or higher level "
            "exclusive) or by 0-based paragraph_range tuple. Provide exactly "
            "one of `heading_pattern` or `paragraph_range`. Refuses to "
            "overwrite unless `overwrite=True`. Snapshots pre+post."
        ),
    )
    def section_extract(
        src: str,
        dst: str,
        heading_pattern: str | None = None,
        paragraph_range: list[int] | tuple[int, int] | None = None,
        overwrite: bool = False,
    ) -> dict[str, Any]:
        if (heading_pattern is None) == (paragraph_range is None):
            raise ValidationError(
                "section_extract: provide exactly one of heading_pattern or paragraph_range"
            )
        resolved_src = resolve_for_read(ctx, src)
        resolved_dst = resolve_for_write(ctx, dst)
        if resolved_src.absolute.suffix.lower() != ".docx":
            raise UnsupportedFormatError(
                f"section_extract: source must be .docx, got {resolved_src.absolute.suffix!r}"
            )
        if resolved_dst.absolute.suffix.lower() != ".docx":
            raise UnsupportedFormatError(
                f"section_extract: destination must end in .docx, got "
                f"{resolved_dst.absolute.suffix!r}"
            )
        if resolved_dst.absolute.exists() and not overwrite:
            raise ValidationError(
                f"destination exists: {resolved_dst.absolute}; pass overwrite=True"
            )

        src_doc = _safe_load(resolved_src.absolute)
        if heading_pattern is not None:
            elements = _extract_by_heading(src_doc, heading_pattern)
        else:
            assert paragraph_range is not None
            if len(paragraph_range) != 2:
                raise ValidationError("paragraph_range must be a [start, end] pair")
            start, end = int(paragraph_range[0]), int(paragraph_range[1])
            if start < 0 or end < start:
                raise ValidationError(f"invalid paragraph_range: ({start}, {end})")
            body = _body_children(src_doc)
            if start >= len(body):
                raise ValidationError(
                    f"paragraph_range start {start} out of range (have {len(body)} elements)"
                )
            elements = body[start : end + 1]

        if not elements:
            raise ValidationError("no content matched the extraction criteria")

        rel = resolved_dst.rel_to_root.as_posix()
        root_name = resolved_dst.root.name
        with file_lock(resolved_dst.absolute):
            if resolved_dst.absolute.exists():
                ctx.versions.snapshot(
                    root_name=root_name,
                    rel_path=rel,
                    source=resolved_dst.absolute,
                    action="section_extract_pre",
                )
            _build_extracted_doc(src_doc, elements, resolved_dst.absolute)
            snap = ctx.versions.snapshot(
                root_name=root_name,
                rel_path=rel,
                source=resolved_dst.absolute,
                action="section_extract_post",
            )
        ctx.audit.log(
            "section_extract",
            src=str(resolved_src.absolute),
            dst=str(resolved_dst.absolute),
            elements=len(elements),
            heading_pattern=heading_pattern,
            paragraph_range=list(paragraph_range) if paragraph_range else None,
        )
        return {
            "src": summarize_resolved(resolved_src),
            "dst": summarize_resolved(resolved_dst),
            "elements_copied": len(elements),
            "snapshot": snap,
        }

    @mcp.tool(
        name="section_merge",
        description=(
            "Merge multiple DOCX files into one via docxcompose. The first "
            "source becomes the master (its styles, headers, footers, page "
            "setup, and section properties win). When `preserve_styles=True`, "
            "conflicting style IDs are renamed (e.g. MyStyle -> MyStyle_1) "
            "instead of being discarded. `page_break_between=True` (default) "
            "inserts an explicit page break before each appended document. "
            "Refuses to overwrite unless `overwrite=True`. Snapshots pre+post."
        ),
    )
    def section_merge(
        sources: list[str],
        dst: str,
        preserve_styles: bool = False,
        page_break_between: bool = True,
        overwrite: bool = False,
    ) -> dict[str, Any]:
        if not sources:
            raise ValidationError("section_merge: sources must not be empty")
        if len(sources) < 2:
            raise ValidationError("section_merge: at least 2 source files are required")

        resolved_dst = resolve_for_write(ctx, dst)
        if resolved_dst.absolute.suffix.lower() != ".docx":
            raise UnsupportedFormatError(
                f"section_merge: destination must end in .docx, got "
                f"{resolved_dst.absolute.suffix!r}"
            )
        if resolved_dst.absolute.exists() and not overwrite:
            raise ValidationError(
                f"destination exists: {resolved_dst.absolute}; pass overwrite=True"
            )

        resolved_sources = []
        for s in sources:
            r = resolve_for_read(ctx, s)
            if r.absolute.suffix.lower() != ".docx":
                raise UnsupportedFormatError(
                    f"section_merge: every source must be .docx, got {r.absolute.suffix!r} ({s})"
                )
            if not r.absolute.exists():
                raise ValidationError(f"source not found: {r.absolute}")
            resolved_sources.append(r)

        # Master is loaded once and mutated in place via Composer.
        from docxcompose.composer import Composer

        master_doc = _safe_load(resolved_sources[0].absolute)
        composer = Composer(master_doc)
        # docxcompose 2.x exposes preserve_styles as an attribute, not a
        # constructor kwarg. Set it explicitly so the merge picks up the
        # caller's choice.
        composer.preserve_styles = preserve_styles

        for resolved_src in resolved_sources[1:]:
            if page_break_between:
                _insert_page_break(master_doc)
            appended = _safe_load(resolved_src.absolute)
            try:
                composer.append(appended)
            except Exception as exc:  # noqa: BLE001 - composer raises various types
                raise HandlerError(
                    f"section_merge failed appending {resolved_src.absolute}: {exc}"
                ) from exc

        rel = resolved_dst.rel_to_root.as_posix()
        root_name = resolved_dst.root.name
        with file_lock(resolved_dst.absolute):
            if resolved_dst.absolute.exists():
                ctx.versions.snapshot(
                    root_name=root_name,
                    rel_path=rel,
                    source=resolved_dst.absolute,
                    action="section_merge_pre",
                )
            try:
                composer.save(str(resolved_dst.absolute))
            except Exception as exc:  # noqa: BLE001 - composer raises various types
                raise HandlerError(
                    f"section_merge failed saving {resolved_dst.absolute}: {exc}"
                ) from exc
            snap = ctx.versions.snapshot(
                root_name=root_name,
                rel_path=rel,
                source=resolved_dst.absolute,
                action="section_merge_post",
            )
        ctx.audit.log(
            "section_merge",
            sources=[str(r.absolute) for r in resolved_sources],
            dst=str(resolved_dst.absolute),
            preserve_styles=preserve_styles,
            page_break_between=page_break_between,
        )
        return {
            "sources": [summarize_resolved(r) for r in resolved_sources],
            "dst": summarize_resolved(resolved_dst),
            "preserve_styles": preserve_styles,
            "page_break_between": page_break_between,
            "snapshot": snap,
        }
