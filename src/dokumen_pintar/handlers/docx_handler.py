"""DOCX format handler (python-docx backed)."""

from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document
from docx.opc.exceptions import PackageNotFoundError

from dokumen_pintar.errors import HandlerError
from dokumen_pintar.handlers.base import (
    FormatHandler,
    HandlerCapability,
    default_registry,
)


def _open(path: Path) -> Any:
    try:
        return Document(str(path))
    except PackageNotFoundError as exc:
        raise HandlerError(f"not a valid docx: {path} ({exc})") from exc
    except Exception as exc:  # noqa: BLE001
        raise HandlerError(f"failed to open docx: {path} ({exc})") from exc


def _heading_level(style_name: str | None) -> int | None:
    if not style_name:
        return None
    name = style_name.strip()
    # python-docx heading styles are named "Heading 1", "Heading 2", ...
    if name.lower().startswith("heading "):
        tail = name.split(" ", 1)[1].strip()
        if tail.isdigit():
            return int(tail)
    if name.lower() == "title":
        return 0
    return None


def _collect_headings(doc: Any) -> list[dict[str, Any]]:
    headings: list[dict[str, Any]] = []
    for idx, para in enumerate(doc.paragraphs):
        style_name = getattr(getattr(para, "style", None), "name", None)
        level = _heading_level(style_name)
        if level is not None:
            headings.append(
                {
                    "index": idx,
                    "level": level,
                    "text": para.text,
                }
            )
    return headings


def _core_props_dict(doc: Any) -> dict[str, Any]:
    cp = doc.core_properties
    created = getattr(cp, "created", None)
    modified = getattr(cp, "modified", None)
    return {
        "title": getattr(cp, "title", None),
        "author": getattr(cp, "author", None),
        "created": created.isoformat() if created is not None else None,
        "modified": modified.isoformat() if modified is not None else None,
    }


def _parse_index_expr(expr: str, prefix: str) -> int:
    raw = expr[len(prefix) :].strip()
    if not raw or not raw.lstrip("-").isdigit():
        raise HandlerError(f"invalid index in expression: {expr!r}")
    try:
        return int(raw)
    except ValueError as exc:  # pragma: no cover
        raise HandlerError(f"invalid index in expression: {expr!r}") from exc


class DocxHandler:
    """Handler for Microsoft Word `.docx` files via python-docx."""

    name: str = "docx"
    extensions: tuple[str, ...] = (".docx",)
    capabilities: HandlerCapability = (
        HandlerCapability.READ_TEXT
        | HandlerCapability.WRITE_TEXT
        | HandlerCapability.STRUCTURED_GET
        | HandlerCapability.STRUCTURED_SET
        | HandlerCapability.STRUCTURED_DELETE
        | HandlerCapability.SEARCH_EXTRACTED
        | HandlerCapability.WRITE_META
    )

    def detect(self, path: Path) -> bool:
        return path.suffix.lower() in self.extensions

    # ---------- reading ----------

    def read_text(self, path: Path, **_: Any) -> str:
        doc = _open(path)
        return "\n".join(p.text for p in doc.paragraphs)

    def read_meta(self, path: Path) -> dict[str, Any]:
        stat = path.stat()
        doc = _open(path)
        return {
            "format": self.name,
            "size": stat.st_size,
            "mtime": stat.st_mtime,
            "paragraph_count": len(doc.paragraphs),
            "table_count": len(doc.tables),
            "section_count": len(doc.sections),
            "core_props": _core_props_dict(doc),
            "headings": _collect_headings(doc),
        }

    def extract_for_search(self, path: Path) -> str:
        try:
            doc = _open(path)
        except HandlerError:
            return ""
        parts: list[str] = [p.text for p in doc.paragraphs if p.text]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text
                    if text:
                        parts.append(text)
        return "\n".join(parts)

    # ---------- writing ----------

    def write_text(self, path: Path, content: str, **_: Any) -> None:
        # write_text on docx CREATES a fresh document containing only plain
        # paragraphs — refuse to clobber an existing file, which would silently
        # destroy any styles, tables, images, headers/footers, comments, or
        # other rich content. Callers that need to mutate an existing docx
        # should use the structured API (structured_set with `paragraph:N`).
        if path.exists():
            raise HandlerError(
                f"write_text refuses to overwrite existing docx '{path}' "
                "(would discard styles, tables, images, and other rich "
                "content). Delete the file first if you really want to "
                "replace it, or use structured_set to mutate paragraphs."
            )
        try:
            doc = Document()
            for line in content.split("\n"):
                doc.add_paragraph(line)
            doc.save(str(path))
        except HandlerError:
            raise
        except Exception as exc:  # noqa: BLE001
            raise HandlerError(f"failed to write docx: {path} ({exc})") from exc

    # ---------- structured ----------

    def structured_get(self, path: Path, expr: str) -> Any:
        doc = _open(path)
        key = expr.strip()

        if key == "paragraphs":
            return [
                {
                    "index": i,
                    "text": p.text,
                    "style": getattr(getattr(p, "style", None), "name", None),
                }
                for i, p in enumerate(doc.paragraphs)
            ]

        if key.startswith("paragraph:"):
            idx = _parse_index_expr(key, "paragraph:")
            paragraphs = doc.paragraphs
            if idx < 0 or idx >= len(paragraphs):
                raise HandlerError(f"paragraph index out of range: {idx}")
            p = paragraphs[idx]
            return {
                "index": idx,
                "text": p.text,
                "style": getattr(getattr(p, "style", None), "name", None),
            }

        if key == "tables":
            return [
                [[cell.text for cell in row.cells] for row in table.rows] for table in doc.tables
            ]

        if key.startswith("table:"):
            idx = _parse_index_expr(key, "table:")
            tables = doc.tables
            if idx < 0 or idx >= len(tables):
                raise HandlerError(f"table index out of range: {idx}")
            table = tables[idx]
            return [[cell.text for cell in row.cells] for row in table.rows]

        if key == "headings":
            return _collect_headings(doc)

        if key == "core_props":
            return _core_props_dict(doc)

        raise HandlerError(f"unsupported structured expression: {expr!r}")

    def structured_set(self, path: Path, expr: str, value: Any) -> None:
        doc = _open(path)
        key = expr.strip()

        if key.startswith("paragraph:"):
            idx = _parse_index_expr(key, "paragraph:")
            paragraphs = doc.paragraphs
            if idx < 0 or idx >= len(paragraphs):
                raise HandlerError(f"paragraph index out of range: {idx}")
            if not isinstance(value, dict):
                raise HandlerError(
                    "paragraph value must be a dict with 'text' and optional 'style'"
                )
            if "text" not in value or not isinstance(value["text"], str):
                raise HandlerError("paragraph value missing required 'text' string")
            p = paragraphs[idx]
            p.text = value["text"]
            style = value.get("style")
            if style is not None:
                if not isinstance(style, str):
                    raise HandlerError("paragraph 'style' must be a string or null")
                try:
                    p.style = doc.styles[style]
                except KeyError as exc:
                    raise HandlerError(f"unknown style: {style!r}") from exc
        elif key == "core_props":
            if not isinstance(value, dict):
                raise HandlerError("core_props value must be a dict")
            cp = doc.core_properties
            for prop_name, prop_value in value.items():
                if not hasattr(cp, prop_name):
                    raise HandlerError(f"unknown core property: {prop_name!r}")
                try:
                    setattr(cp, prop_name, prop_value)
                except (AttributeError, TypeError, ValueError) as exc:
                    raise HandlerError(f"failed to set core property {prop_name!r}: {exc}") from exc
        else:
            raise HandlerError(f"unsupported structured_set expression: {expr!r}")

        try:
            doc.save(str(path))
        except Exception as exc:  # noqa: BLE001
            raise HandlerError(f"failed to save docx: {path} ({exc})") from exc

    def structured_delete(self, path: Path, expr: str) -> None:
        doc = _open(path)
        key = expr.strip()

        if key.startswith("paragraph:"):
            idx = _parse_index_expr(key, "paragraph:")
            paragraphs = doc.paragraphs
            if idx < 0 or idx >= len(paragraphs):
                raise HandlerError(f"paragraph index out of range: {idx}")
            p = paragraphs[idx]
            parent = p._element.getparent()
            if parent is None:
                raise HandlerError(f"paragraph {idx} has no parent element; cannot remove")
            parent.remove(p._element)
        elif key.startswith("table:"):
            idx = _parse_index_expr(key, "table:")
            tables = doc.tables
            if idx < 0 or idx >= len(tables):
                raise HandlerError(f"table index out of range: {idx}")
            t = tables[idx]
            parent = t._element.getparent()
            if parent is None:
                raise HandlerError(f"table {idx} has no parent element; cannot remove")
            parent.remove(t._element)
        else:
            raise HandlerError(f"unsupported structured_delete expression: {expr!r}")

        try:
            doc.save(str(path))
        except Exception as exc:  # noqa: BLE001
            raise HandlerError(f"failed to save docx: {path} ({exc})") from exc


    # ---------- metadata write ----------

    _WRITABLE_CORE_PROPS: tuple[str, ...] = (
        "author",
        "category",
        "comments",
        "content_status",
        "created",
        "identifier",
        "keywords",
        "language",
        "last_modified_by",
        "last_printed",
        "modified",
        "revision",
        "subject",
        "title",
        "version",
    )

    def write_meta(self, path: Path, updates: dict[str, Any]) -> dict[str, Any]:
        """Merge ``updates`` into the document's core properties.

        Unknown keys raise :class:`HandlerError`. Returns the dict that was
        actually applied (matching the requested updates after normalization).
        """
        doc = _open(path)
        cp = doc.core_properties
        applied: dict[str, Any] = {}
        for key, value in updates.items():
            if key not in self._WRITABLE_CORE_PROPS:
                raise HandlerError(
                    f"unknown core property: {key!r} "
                    f"(allowed: {list(self._WRITABLE_CORE_PROPS)})"
                )
            try:
                setattr(cp, key, value)
            except (AttributeError, TypeError, ValueError) as exc:
                raise HandlerError(
                    f"failed to set core property {key!r}: {exc}"
                ) from exc
            applied[key] = value
        try:
            doc.save(str(path))
        except Exception as exc:  # noqa: BLE001
            raise HandlerError(f"failed to save docx: {exc}") from exc
        return applied

    def strip_meta(self, path: Path) -> dict[str, Any]:
        """Clear every writable core property to its empty default."""
        doc = _open(path)
        cp = doc.core_properties
        cleared: list[str] = []
        for key in self._WRITABLE_CORE_PROPS:
            try:
                # Strings reset to "", datetimes/objects reset to None.
                current = getattr(cp, key, None)
                if isinstance(current, str):
                    setattr(cp, key, "")
                else:
                    setattr(cp, key, None)
                cleared.append(key)
            except (AttributeError, TypeError, ValueError):  # pragma: no cover
                # Some properties (e.g. "revision") may be read-only in
                # certain python-docx versions; silently skip those rather
                # than fail the whole strip. This branch is defensive and
                # not reachable in the version tested here.
                continue
        try:
            doc.save(str(path))
        except Exception as exc:  # noqa: BLE001
            raise HandlerError(f"failed to save docx: {exc}") from exc
        return {"stripped": cleared}


# Runtime-checkable protocol sanity assertion.
_handler: FormatHandler = DocxHandler()
default_registry.register(_handler)
